# -*- coding: utf-8 -*-
"""Generate mayor4code BSc project report (full: prelims + Chapters 1-5 +
References) per CSC / OAUSTECH guidelines."""
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import subprocess
import sys

doc = Document()

# ---- Base formatting: Times New Roman 12, A4, margins, 1.5 spacing ----
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(6)

# Heading styles (so the auto Table of Contents can pick them up)
for name, size in [('Heading 1', 14), ('Heading 2', 12), ('Heading 3', 12)]:
    st = doc.styles[name]
    st.font.name = 'Times New Roman'
    st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.italic = (name == 'Heading 3')
    st.font.color.rgb = RGBColor(0, 0, 0)
    st.paragraph_format.space_before = Pt(6)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

sec = doc.sections[0]
sec.page_width = Mm(210)
sec.page_height = Mm(297)
sec.left_margin = Mm(40)
sec.right_margin = Mm(25)
sec.top_margin = Mm(25)
sec.bottom_margin = Mm(25)


def _page_field(paragraph):
    """Insert a PAGE field into a footer paragraph."""
    run = paragraph.add_run()
    fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end')
    run._r.append(fb); run._r.append(it); run._r.append(fe)


def set_page_numbering(section, fmt, start=None, hide_first=False):
    """fmt: 'lowerRoman' or 'decimal'. Adds centered page number in footer."""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType'); sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))
    if hide_first:
        section.different_first_page_header_footer = True
        section.first_page_footer.is_linked_to_previous = False
        # leave first-page footer empty (title page: no number)
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ''
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _page_field(p)


def add_toc(field='TOC'):
    """Insert an auto-updating field (TOC / table of figures)."""
    para = doc.add_paragraph()
    run = para.add_run()
    fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
    if field == 'TOC':
        it.text = 'TOC \\o "1-3" \\h \\z \\u'
    fsep = OxmlElement('w:fldChar'); fsep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = 'Right-click and choose "Update Field" to build this table.'
    hold = OxmlElement('w:r'); hold.append(t)
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end')
    run._r.append(fb); run._r.append(it); run._r.append(fsep)
    run._r.append(hold); run._r.append(fe)


def chapter(title_lines, new_section=False):
    """Chapter/major heading as Heading 1, centered, starting on a new page.
    Accepts a list (legacy) or a string. new_section=True begins Arabic
    numbering for the main body."""
    if isinstance(title_lines, (list, tuple)):
        title = ' '.join(title_lines)
    else:
        title = title_lines
    if new_section:
        new = doc.add_section(WD_SECTION.NEW_PAGE)
        new.page_width = Mm(210); new.page_height = Mm(297)
        new.left_margin = Mm(40); new.right_margin = Mm(25)
        new.top_margin = Mm(25); new.bottom_margin = Mm(25)
        set_page_numbering(new, 'decimal', start=1)
    else:
        doc.add_page_break()
    p = doc.add_heading('', level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title.upper())


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def prelim_title(text):
    """Centered bold heading for preliminary pages; NOT in the TOC."""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(14)
    return p


def body(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def figure(img_path, caption_text):
    """Embed an image centred, with a caption below."""
    import os
    if not os.path.exists(img_path):
        body(f'[Image not found: {img_path}]')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(img_path, width=Mm(140))
    cap = doc.add_paragraph(caption_text)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True


def generate_flowchart(output_path):
    """Generate system flowchart PNG via the helper script."""
    script = os.path.join(os.path.dirname(__file__),
                          'shots', 'mayor4code', 'generate_flowchart.py')
    if os.path.exists(script):
        subprocess.run([sys.executable, script], check=True)
    else:
        print(f'[WARN] Flowchart script not found: {script}')


# Generate flowchart image before building report
generate_flowchart(None)

# CONTENT_START
# ---- Title Page ----
for line in [
    'DESIGN AND IMPLEMENTATION OF A WEB-BASED INTERACTIVE E-LEARNING '
    'PLATFORM FOR STRUCTURED PYTHON PROGRAMMING (A CASE STUDY OF MAYOR4CODE)',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.bold = True
for line in ['', 'BY', '', 'Giwa Mayowa Bopoola', 'CSC/22/204', '',
             'DEPARTMENT OF COMPUTER SCIENCE', 'SCHOOL OF COMPUTING',
             'OLUSEGUN AGAGU UNIVERSITY OF SCIENCE AND TECHNOLOGY, OKITIPUPA', '',
             'IN PARTIAL FULFILMENT OF THE REQUIREMENTS FOR THE AWARD OF THE '
             'DEGREE OF BACHELOR OF SCIENCE (B.Sc.) IN COMPUTER SCIENCE', '',
             'August, 2026']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(line)

# Preliminary pages: lowercase Roman numerals, title page unnumbered
set_page_numbering(doc.sections[0], 'lowerRoman', start=1, hide_first=True)

# ---- Certification ----
prelim_title('Certification')
body('This is to certify that this project titled “Design and Implementation of a Web-Based '
     'Interactive E-Learning Platform for Structured Python Programming (A Case Study of '
     'mayor4code)” was carried out by Giwa Mayowa Bopoola with matriculation number '
     'CSC/22/204 in the Department of Computer Science, School of Computing, '
     'Olusegun Agagu University of Science and Technology, Okitipupa, in partial fulfilment of '
     'the requirements for the award of the degree of Bachelor of Science (B.Sc.) in Computer '
     'Science.')
table = doc.add_table(rows=4, cols=2)
table.alignment = WD_ALIGN_PARAGRAPH.CENTER
table.style = 'Table Grid'
cells = [
    ('_______________________', '_______________________'),
    ('Dr. Adeolu Obamehinti',  'Dr. Ajoke Gbadamosi'),
    ('Project Supervisor',     'Head of Department'),
    ('Date: __________________','Date: __________________'),
]
for r, (left, right) in enumerate(cells):
    for c, val in [(0, left), (1, right)]:
        cell = table.cell(r, c)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(val)

# ---- Declaration ----
prelim_title('Declaration')
body('I hereby declare that this project is my original work and has not been submitted, either '
     'in whole or in part, for any other degree or qualification in this or any other '
     'institution. All sources of information and materials used have been duly acknowledged by '
     'means of complete references.')
for line in ['', '', '_______________________', 'Giwa Mayowa Bopoola', 'CSC/22/204',
             'Date: __________________']:
    doc.add_paragraph(line)

# ---- Dedication ----
prelim_title('Dedication')
body('This project is dedicated to Almighty God, and to my family, whose support and '
     'encouragement made this work possible.')

# ---- Acknowledgements ----
prelim_title('Acknowledgements')
body('I express my sincere gratitude to my project supervisor, Dr. Adeolu Obamehinti, for the '
      'guidance, patience, and constructive criticism that shaped this work. I am grateful to '
      'the Head and entire staff of the Department of Computer Science, Olusegun Agagu University '
      'of Science and Technology, Okitipupa, for the knowledge imparted throughout my programme. '
      'I am also grateful to my senior colleague, Senpai, whose technical guidance, debugging '
      'support, and architectural insights were invaluable throughout the development of the '
      'mayor4code platform. I also thank my family and friends for their unwavering support and '
      'encouragement. Above all, I give thanks to Almighty God for His grace and strength.')

# ---- Abstract ----
prelim_title('Abstract')
body('Learning to program is widely regarded as difficult for beginners, largely because '
     'traditional methods separate instructional content from practical coding and provide '
     'little immediate feedback. This project presents the design and implementation of '
     'mayor4code, a web-based interactive e-learning platform for structured Python programming. '
     'The system was developed using the Agile approach within the Software Development Life '
     'Cycle and adopts a three-tier architecture comprising presentation, application, and '
     'database layers. It was implemented using the Django web framework, the Python '
     'programming language, HTML, CSS, and JavaScript, with SQLite used during development and '
     'PostgreSQL in production. The platform delivers twelve sequential Python lessons with '
     'locked progression, in which each lesson is unlocked only after the learner passes the '
     'preceding quiz at a minimum of sixty percent. It further provides multiple-choice quizzes '
     'with instant automated scoring, an in-browser Python playground that executes learner '
     'code in an isolated subprocess, progress tracking, automatically issued completion '
     'certificates bearing unique verification codes, and a leaderboard that ranks learners by '
     'performance. The platform was tested through unit, functional, and usability testing, and '
     'the results showed that it correctly enforced lesson progression, scored assessments '
     'accurately, executed code safely, and issued valid certificates. The study concludes that '
     'integrating structured lessons, practical coding, automated assessment, and gamification '
     'within a single secure web application improves the accessibility and effectiveness of '
     'introductory programming education. It is recommended that the platform be extended with '
     'additional courses and adaptive learning features in future work.')

# ---- Table of Contents ----
prelim_title('Table of Contents')
add_toc('TOC')

# ---- List of Figures ----
prelim_title('List of Figures')
for line in ['Figure 3.1: System Architecture of the mayor4code Platform',
             'Figure 3.2: System Flowchart of the mayor4code Platform',
             'Figure 3.3: Entity-Relationship Diagram of the mayor4code Database',
             'Figure 4.1: Registration Interface',
             'Figure 4.2: Login Interface',
             'Figure 4.3: Learner Dashboard',
             'Figure 4.4: Lesson Interface',
             'Figure 4.5: Quiz Interface',
             'Figure 4.6: Python Playground Interface',
             'Figure 4.7: Completion Certificate',
             'Figure 4.8: Leaderboard Interface',
             'Figure 4.9: Administration Interface']:
    body(line)

# ---- List of Tables ----
prelim_title('List of Tables')
for line in ['Table 2.1: Comparison of mayor4code with Existing Platforms',
             'Table 3.1: Summary of Principal Database Entities',
             'Table 3.2: Functional Requirements Specification',
             'Table 3.3: Non-Functional Requirements Specification',
             'Table 4.1: Software Development Tools',
             'Table 4.2: Interface Features by Screen',
             'Table 4.3: Summary of Test Cases and Results']:
    body(line)

# ---- List of Abbreviations ----
prelim_title('List of Abbreviations')
for line in ['API — Application Programming Interface',
             'CSS — Cascading Style Sheets',
             'DBMS — Database Management System',
             'HTML — HyperText Markup Language',
             'HTTP — HyperText Transfer Protocol',
             'LMS — Learning Management System',
             'ORM — Object-Relational Mapper',
             'SDLC — Software Development Life Cycle',
             'SQL — Structured Query Language',
             'UI — User Interface']:
    body(line)

# ==================== CHAPTER ONE ====================
chapter(['Chapter One', 'Introduction'], new_section=True)

h2('1.1 Background of the Study')
body('The rapid advancement of information and communication technology has transformed the '
     'way knowledge is delivered, acquired, and assessed across all levels of education. '
     'Electronic learning (e-learning) platforms have emerged as powerful tools that extend '
     'learning beyond the physical classroom, enabling learners to study at their own pace, '
     'from any location, using internet-connected devices. In the field of computer '
     'programming, where practical skill and continuous practice are essential, interactive '
     'e-learning platforms have proven particularly effective because they combine '
     'instructional content with hands-on coding activities.')
body('Learning to program is widely regarded as challenging for beginners. Traditional methods '
     'of teaching programming, which rely heavily on lectures and textbooks, often fail to '
     'provide the immediate feedback and practical engagement that novice programmers require. '
     'According to Pressman and Maxim (2020), interactive software systems that provide timely '
     'feedback significantly improve user engagement and learning outcomes. Interactive coding '
     'platforms address this gap by allowing learners to write, execute, and test code directly '
     'within the browser while receiving instant feedback on their progress.')
body('A structured e-learning platform organises instructional content into sequential lessons '
     'and assessments, guiding learners progressively from basic to advanced concepts. Such '
     'platforms typically incorporate features such as user authentication, progress tracking, '
     'automated assessment, and certification. Sommerville (2016) noted that web-based '
     'applications offer flexibility, scalability, and ease of maintenance, making them well '
     'suited for modern educational systems that must serve many concurrent users across '
     'diverse devices.')
body('This study presents the design and implementation of mayor4code, a web-based interactive '
     'e-learning platform for structured Python programming. The platform delivers twelve '
     'sequential Python lessons with locked progression, automated multiple-choice quizzes, an '
     'in-browser Python playground for practical coding, automatically issued completion '
     'certificates, and a leaderboard that fosters healthy competition among learners. By '
     'integrating structured content delivery with interactive practice and automated '
     'assessment, the platform seeks to make the process of learning Python more engaging, '
     'accessible, and effective for beginners.')

h2('1.2 Research Justification/Motivation')
body('The motivation for this research stems from the increasing demand for accessible and '
     'practical programming education. Python has become one of the most widely adopted '
     'programming languages in academia and industry due to its simple syntax and broad '
     'application in areas such as data science, artificial intelligence, and web development. '
     'However, many beginners struggle to acquire programming skills through conventional '
     'learning methods that lack interactivity and immediate feedback.')
body('Existing programming education is often fragmented, requiring learners to switch between '
     'separate tools for reading lessons, writing code, and taking assessments. This '
     'fragmentation increases cognitive load and discourages consistent practice. There is '
     'therefore a need for an integrated platform that combines structured lessons, practical '
     'coding, and assessment within a single, coherent environment.')
body('Furthermore, the principle of locked progression, in which each lesson is unlocked only '
     'after the learner demonstrates mastery of the preceding one, encourages disciplined and '
     'thorough learning. Coupled with gamification elements such as certificates and a '
     'leaderboard, this approach is expected to improve learner motivation and retention. These '
     'considerations motivated the development of mayor4code as a self-contained, interactive, '
     'and structured platform for learning Python programming.')

h2('1.3 Problem Statement')
body('Despite the growing availability of online learning resources, beginners learning to '
     'program continue to face significant challenges. Many freely available resources present '
     'content in an unstructured manner, leaving learners uncertain about the order in which '
     'topics should be studied. This lack of structure frequently results in knowledge gaps and '
     'discourages learners from completing their studies.')
body('A second challenge is the separation between theory and practice. Learners often read '
     'programming concepts without an immediate opportunity to apply them, and configuring a '
     'local programming environment can be difficult and discouraging for beginners. The '
     'absence of instant, automated feedback further limits learners’ ability to identify '
     'and correct mistakes promptly.')
body('In addition, many learning resources provide no mechanism for tracking progress, '
     'enforcing mastery before advancement, or recognising achievement through certification. '
     'The absence of these features reduces motivation and makes it difficult for learners to '
     'measure their advancement objectively.')
body('This study addresses these challenges through the development of mayor4code, a web-based '
     'interactive e-learning platform that provides structured lesson progression, an '
     'integrated code execution environment, automated assessment with instant scoring, '
     'progress tracking, and automated certification, all within a single secure and '
     'user-friendly system.')

h2('1.4 Aim and Objectives of the Study')
h3('Aim')
body('The aim of this study is to design and implement a web-based interactive e-learning '
     'platform for structured Python programming that integrates lesson delivery, practical '
     'coding, and automated assessment within a single system.')
h3('Objectives')
body('The specific objectives of the study are to:')
numbered('design a web-based platform that delivers structured Python lessons with locked '
         'progression;')
numbered('develop a secure user authentication and authorisation system for learners;')
numbered('implement an interactive in-browser Python playground for practical code execution;')
numbered('develop an automated quiz and instant-scoring mechanism to assess learner mastery;')
numbered('implement progress tracking, automated certification, and a leaderboard to enhance '
         'learner motivation;')
numbered('evaluate the functionality and usability of the developed platform.')

h2('1.5 Scope of the Study')
body('This project focuses on the design and implementation of a web-based interactive '
     'e-learning platform for learning the fundamentals of Python programming. The platform '
     'provides functionalities including user registration and authentication, sequential '
     'delivery of twelve Python lessons, locked lesson progression based on quiz performance, '
     'multiple-choice quizzes with instant scoring, an in-browser Python playground, progress '
     'tracking, automated completion certificates, and a leaderboard.')
body('The system primarily serves learners (students) and an administrator who manages lessons, '
     'quizzes, questions, and user progress. The study is limited to the teaching and '
     'assessment of introductory Python programming and does not extend to other programming '
     'languages, live video instruction, or peer-to-peer collaboration features.')

h2('1.6 Significance of the Study')
body('The development of the mayor4code e-learning platform provides several benefits to '
     'learners, educators, and educational institutions.')
h3('1.6.1 Educational Impact')
body('The platform promotes structured and self-paced learning of Python programming. By '
     'enforcing mastery before progression and providing instant feedback through automated '
     'quizzes and an interactive playground, it enhances comprehension, retention, and '
     'practical skill acquisition among beginners.')
h3('1.6.2 Technological Impact')
body('The project demonstrates the practical application of modern web technologies, database '
     'management systems, and secure code-execution techniques in solving a real-world '
     'educational problem. It contributes to the digital transformation of programming '
     'education by providing an integrated and automated learning environment.')
h3('1.6.3 Economical Impact')
body('The platform reduces the cost associated with traditional programming instruction, such '
     'as physical infrastructure, printed materials, and one-on-one tutoring. As a web-based '
     'solution, it enables a single deployment to serve many learners simultaneously, thereby '
     'improving cost efficiency and accessibility.')

h2('1.7 Definition of Terms')
body('E-Learning: The delivery of learning and training through digital resources accessed over '
     'the internet or a computer network.')
body('Interactive Platform: A software system that responds to user actions in real time, '
     'enabling active participation rather than passive consumption of content.')
body('Locked Progression: A learning mechanism in which access to subsequent lessons is granted '
     'only after the learner satisfies a defined requirement, such as passing the preceding '
     'quiz.')
body('Playground: An in-browser environment in which learners can write and execute program '
     'code and view the output instantly.')
body('Quiz: A set of multiple-choice questions used to assess a learner’s understanding of '
     'a lesson.')
body('Certificate: A digitally issued document, bearing a unique verification code, that '
     'recognises a learner’s completion of the course.')
body('Authentication: The process of verifying the identity of a user before granting access to '
     'the system.')
body('Leaderboard: A ranked display of learners based on their performance, used to encourage '
     'motivation and engagement.')

h2('1.8 Organization of the Study')
body('This project report is organised into five chapters.')
body('Chapter One presents the introduction, background of the study, motivation, problem '
     'statement, aim and objectives, scope, significance, definition of terms, and organisation '
     'of the study.')
body('Chapter Two reviews relevant literature related to e-learning systems, interactive '
     'programming education, learning management systems, gamification, and existing platforms.')
body('Chapter Three describes the methodology adopted for the design and development of the '
     'proposed system, including the design approach, system architecture, and implementation '
     'tools.')
body('Chapter Four presents the implementation details, system interfaces, testing procedures, '
     'results, and discussion of findings.')
body('Chapter Five provides the summary of the study, conclusion, recommendations, and '
     'suggestions for future research.')
# ==================== CHAPTER TWO ====================
chapter(['Chapter Two', 'Literature Review'])

h2('2.1 Introduction')
body('This chapter reviews relevant literature related to e-learning systems, interactive '
     'programming education, learning management systems, gamification in learning, web '
     'development technologies, and existing programming-learning platforms. The review '
     'provides an understanding of previous work and identifies the gaps that justify the '
     'development of the proposed mayor4code platform.')

h2('2.2 Concept of E-Learning Systems')
body('E-learning refers to the use of electronic technologies and the internet to deliver '
     'educational content and facilitate learning outside the traditional classroom. E-learning '
     'systems enable learners to access instructional materials, complete assessments, and '
     'monitor their progress from any location using internet-connected devices. According to '
     'Sommerville (2016), web-based systems support distributed access to resources, thereby '
     'improving operational efficiency and user convenience.')
body('E-learning has been widely adopted in education because it offers flexibility, '
     'scalability, and cost effectiveness. Garrison (2017) observed that e-learning supports '
     'self-paced, learner-centred education that can reach large numbers of learners '
     'simultaneously, while Anderson (2008) noted that well-designed online learning '
     'environments can be as effective as traditional instruction. In the '
     'context of programming education, e-learning systems are particularly valuable '
     'because they can integrate practical coding activities directly into the learning '
      'environment.')

h2('2.3 Constructivist Learning Theory')
body('Constructivist learning theory posits that learners actively construct knowledge rather '
      'than passively receive it. Piaget (1952) described cognitive constructivism, in which '
      'learners build mental models by assimilating new information into existing schemas and '
      'accommodating schemas when new information does not fit. In the context of programming '
      'education, this implies that learners must actively engage with concepts by writing code, '
      'testing hypotheses, and debugging errors, rather than merely reading about syntax and '
      'semantics.')
body('Vygotsky (1978) extended constructivism with the concept of the Zone of Proximal '
      'Development (ZPD), defined as the gap between what a learner can achieve independently '
      'and what they can achieve with guidance. Interactive e-learning platforms can function as '
      'scaffolding tools within the ZPD by providing immediate feedback, hint systems, and '
      'structured progression that gradually reduces support as the learner gains competence. '
      'Jonassen (1999) argued that constructivist learning environments should present authentic '
      'problems, provide multiple representations of content, and support collaborative '
      'knowledge construction, principles that informed the design of the mayor4code platform\'s '
      'sequential lesson structure and integrated coding playground.')
body('The locked progression mechanism in mayor4code aligns with the constructivist principle '
      'that learners must master prerequisite concepts before advancing to more complex topics. '
      'By requiring a minimum quiz score of sixty percent to unlock subsequent lessons, the '
      'platform ensures that each learner has constructed a sufficient foundation of '
      'understanding before being exposed to new material. This approach prevents the '
      'accumulation of knowledge gaps that commonly hinder progress in self-directed programming '
      'education.')

h2('2.4 Mastery Learning and Sequential Progression')
body('Mastery learning, originally proposed by Bloom (1968), is an instructional strategy in '
      'which learners must achieve a defined level of competence before progressing to the next '
      'unit of instruction. Bloom argued that given sufficient time and appropriate instruction, '
      'most learners can achieve mastery of any subject, and that the key differentiator is not '
      'innate ability but the opportunity to learn at one\'s own pace. Guskey (2007) reviewed '
      'decades of mastery learning research and concluded that the approach consistently produces '
      'superior learning outcomes compared to conventional time-bound instruction, particularly '
      'for foundational knowledge that supports subsequent learning.')
body('Keller (1968) developed the Personalized System of Instruction (PSI), which applied '
      'mastery learning principles to higher education through self-paced modules, unit '
      'quizzes with passing criteria, and immediate feedback. The PSI model bears direct '
      'similarity to the design of mayor4code, where each lesson is a self-contained module, '
      'the quiz serves as the mastery assessment, the sixty percent pass mark defines the '
      'mastery threshold, and locked progression ensures that learners demonstrate competence '
      'before advancing. Research by Kulik, Kulik, and Bangert-Drowns (1990) in a meta-analysis '
      'of PSI studies found that the approach led to higher achievement and greater learner '
      'satisfaction compared to conventional lecture-based instruction.')
body('The sequential progression model adopted by mayor4code is further supported by the '
      'principles of deliberate practice described by Ericsson, Krampe, and Tesch-Romer (1993). '
      'Deliberate practice involves well-defined learning objectives, immediate feedback, '
      'repetition, and progressive difficulty — all characteristics embedded in the platform\'s '
      'lesson-quiz-unlock cycle. By integrating deliberate practice with mastery learning '
      'criteria, mayor4code provides a structured pathway that guides learners from basic '
      'syntax to more advanced Python programming concepts while ensuring foundational '
      'understanding at each step.')

h2('2.5 Interactive Programming Learning Platforms')
body('Interactive programming learning platforms combine instructional content with hands-on '
     'coding activities, enabling learners to write and execute code while studying. Unlike '
     'static resources such as textbooks and video tutorials, these platforms provide immediate '
     'feedback that helps learners identify and correct errors promptly.')
body('The provision of instant feedback is grounded in established learning theory. Prompt '
     'reinforcement strengthens understanding and encourages continued practice, which is '
     'essential in skill-based disciplines such as programming. Interactive platforms also '
     'reduce the barrier posed by local environment setup, since code is executed within the '
     'browser, allowing beginners to focus on learning rather than configuration.')

h2('2.6 Learning Management and Progress Tracking')
body('A learning management system (LMS) is a software application used to plan, deliver, and '
     'track educational content and learner performance. Core LMS functions include user '
     'management, content delivery, assessment, progress tracking, and reporting. Progress '
     'tracking enables learners to monitor their advancement and helps maintain motivation by '
     'providing a clear sense of accomplishment.')
body('Locked or sequential progression is a design strategy in which learners must satisfy '
     'defined criteria, such as passing an assessment, before advancing to the next unit. This '
     'approach enforces mastery of prerequisite concepts and reduces knowledge gaps. Automated '
     'assessment, in which quizzes are scored instantly by the system, further supports timely '
     'feedback and reduces the administrative burden associated with manual grading.')

h2('2.7 Gamification in Learning')
body('Gamification refers to the application of game design elements, such as points, badges, '
     'certificates, and leaderboards, in non-game contexts to increase engagement and '
     'motivation. Deterding, Dixon, Khaled, and Nacke (2011) defined gamification as the use of '
     'game-design elements in non-game contexts, and in education it has been shown to enhance '
     'learner participation, persistence, and satisfaction. Hamari, Koivisto, and Sarsa (2014), '
     'in a review of empirical studies, found that gamification generally produces positive '
     'effects on engagement, though the outcomes depend on context and implementation.')
body('Elements such as completion certificates provide learners with tangible recognition of '
     'achievement, while leaderboards introduce an element of healthy competition that can '
     'motivate continued effort. When applied appropriately, these mechanisms complement '
     'structured learning by reinforcing desired behaviours and sustaining learner interest '
     'throughout the course.')

h2('2.8 Technologies Used in Web Development')
body('Web development involves the use of various technologies for designing and implementing '
     'web applications. The front-end, responsible for the user interface, is commonly built '
     'using HTML, CSS, and JavaScript. Duckett (2014) explained that HTML provides the '
     'structure of web pages, CSS controls presentation and styling, and JavaScript enables '
     'interactive functionality; together these form the foundation of modern web development.')
body('The back-end handles server-side operations such as authentication, business logic, and '
     'database interactions. Popular back-end technologies include Python, PHP, Node.js, and '
     'Java. The proposed system employs the Django web framework, a high-level Python framework '
     'that promotes rapid development and clean, pragmatic design, and which incorporates '
     'built-in security features and an object-relational mapper for database management '
     '(Django Software Foundation, 2024). Python itself is a general-purpose, high-level '
     'language noted for its readable syntax (Python Software Foundation, 2024).')
body('Data in web applications are managed by a database management system (DBMS). Elmasri and '
     'Navathe (2017) noted that database systems ensure data consistency, integrity, security, '
     'and efficient retrieval of information. The proposed system uses SQLite during development '
     'and PostgreSQL in production, both of which are reliable relational database systems well '
     'supported by the Django framework.')

h2('2.9 Related Works and Existing Systems')

h3('2.9.1 Codecademy Interactive Coding Platform')
body('Approach')
body('Codecademy is a commercial online platform that teaches programming through interactive, '
     'in-browser lessons. Learners read short instructional segments and immediately apply '
     'concepts in an integrated code editor that validates their solutions.')
body('Architectural Framework')
body('The platform combines a web-based user interface, an in-browser code editor and execution '
     'environment, an automated exercise-validation engine, and a centralised database that '
     'stores user accounts and progress.')
body('Strengths')
bullet('Provides highly interactive, hands-on learning.')
bullet('Offers structured learning paths across many languages.')
bullet('Delivers immediate feedback on exercises.')
body('Weaknesses')
bullet('Most advanced content is restricted behind a paid subscription.')
bullet('Certification is limited in the free tier.')
bullet('Requires continuous internet connectivity.')

h3('2.9.2 freeCodeCamp Learning Platform')
body('Approach')
body('freeCodeCamp is a free, non-profit platform that teaches web development and programming '
     'through a sequence of interactive challenges and projects, leading to free certifications.')
body('Architectural Framework')
body('It consists of a web interface, an in-browser challenge editor, an automated test runner '
     'that validates learner solutions, and a database that tracks completion and issues '
     'certifications.')
body('Strengths')
bullet('Completely free and open-source.')
bullet('Offers project-based learning and free certification.')
bullet('Maintains a large supportive community.')
body('Weaknesses')
bullet('Focuses primarily on web development rather than general programming.')
bullet('Self-directed structure may overwhelm absolute beginners.')
bullet('Limited personalised guidance.')

h3('2.9.3 SoloLearn Mobile Learning Application')
body('Approach')
body('SoloLearn teaches programming through bite-sized lessons and quizzes, with a strong '
     'emphasis on mobile learning and community interaction.')
body('Architectural Framework')
body('The system integrates a mobile and web interface, a lesson and quiz engine, a code '
     'playground, and social features backed by a centralised database.')
body('Strengths')
bullet('Bite-sized lessons suited to beginners.')
bullet('Includes gamification through points and streaks.')
bullet('Provides a mobile-friendly experience.')
body('Weaknesses')
bullet('Content depth is limited for advanced learners.')
bullet('Heavy reliance on community-generated content quality.')
bullet('Advertisements in the free version.')

h3('2.9.4 HackerRank Skill Assessment Platform')
body('Approach')
body('HackerRank provides coding challenges and skill assessments primarily for practising '
     'programming and preparing for technical interviews.')
body('Architectural Framework')
body('It comprises a web interface, an online code editor supporting multiple languages, an '
     'automated judging engine, and a database for storing submissions and rankings.')
body('Strengths')
bullet('Extensive library of practice problems.')
bullet('Automated evaluation with instant results.')
bullet('Supports many programming languages.')
body('Weaknesses')
bullet('Oriented towards practice and assessment rather than structured teaching.')
bullet('Challenging for complete beginners.')
bullet('Limited explanatory lesson content.')

h3('2.9.5 W3Schools Online Tutorials')
body('Approach')
body('W3Schools provides free reference tutorials for web and programming technologies, '
     'featuring simple explanations and an in-browser "Try it Yourself" editor.')
body('Architectural Framework')
body('The platform consists of static tutorial pages, an embedded code editor for '
     'experimentation, and optional quiz and certification services.')
body('Strengths')
bullet('Free and easy to navigate.')
bullet('Includes editable code examples.')
bullet('Broad coverage of web technologies.')
body('Weaknesses')
bullet('Reference style lacks enforced structured progression.')
bullet('No mastery-based locking of content.')
bullet('Limited automated progress tracking.')

h3('2.9.6 Summary of Related Works')
body('The reviewed platforms demonstrate the effectiveness of interactive, browser-based '
     'programming education and the value of automated feedback, gamification, and '
     'certification. However, they also reveal recurring limitations, including paywalls, '
     'insufficient structured progression, weak enforcement of mastery before advancement, and '
     'a frequent separation between teaching and practice. These observations inform the design '
      'of the proposed system.')

h3('2.9.7 Comparative Analysis of Existing Platforms')
body('Table 2.1 presents a comparative analysis of mayor4code against the reviewed existing '
      'platforms across seven key features. The comparison illustrates how mayor4code addresses '
      'the gaps identified in the literature by combining all essential features in a single '
      'free platform.')
body('Table 2.1: Comparison of mayor4code with Existing Platforms')

# Build comparison table
tbl = doc.add_table(rows=7, cols=7)
tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
tbl.style = 'Table Grid'
headers = ['Feature', 'mayor4code', 'Codecademy', 'freeCodeCamp', 'SoloLearn', 'HackerRank', 'W3Schools']
features = [
    ['Structured Lessons',       '✓', '✓', '✓', '✓', '✗', '✓'],
    ['Locked Progression',       '✓', '✓', '✗', '✗', '✗', '✗'],
    ['Code Playground',          '✓', '✓', '✓', '✓', '✓', '✓'],
    ['Auto Assessment',          '✓', '✓', '✓', '✓', '✓', '✓'],
    ['Certificates',             '✓', 'Partial', '✓', '✗', '✗', '✓'],
    ['Leaderboard',              '✓', '✗', '✗', '✓', '✓', '✗'],
]
for c, h in enumerate(headers):
    cell = tbl.cell(0, c)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(9)
for ri, row in enumerate(features):
    for ci, val in enumerate(row):
        cell = tbl.cell(ri + 1, ci)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(val)
        r.font.size = Pt(9)

body('The comparison reveals that mayor4code is the only platform among those reviewed that '
      'simultaneously offers structured lessons with enforced locked progression, an integrated '
      'code playground, automated assessment, certification, and a leaderboard — all at no cost '
      'to the learner. This unique combination directly addresses the research gaps identified '
      'in Section 2.10.')

h2('2.10 Research Gap')
body('Based on the reviewed literature and existing systems, the following research gaps were '
     'identified:')
numbered('Limited Free and Structured Progression: Many effective platforms restrict structured '
         'learning paths behind paid subscriptions, while free alternatives often lack enforced '
         'sequential progression.')
numbered('Weak Enforcement of Mastery: Several systems allow learners to advance without '
         'demonstrating understanding of prerequisite concepts, leading to knowledge gaps.')
numbered('Separation of Theory and Practice: Some resources present lessons and coding '
         'environments separately, increasing cognitive load for beginners.')
numbered('Inconsistent Certification: Free platforms frequently provide limited or no '
         'verifiable certification of completion.')
numbered('Limited Integration in a Single System: Few lightweight platforms combine structured '
         'lessons, an integrated code playground, automated assessment, progress tracking, '
         'certification, and a leaderboard within one coherent environment.')

h2('2.11 Summary of Literature Review')
body('This chapter reviewed literature related to e-learning systems, interactive programming '
     'education, learning management and progress tracking, gamification, web development '
     'technologies, and existing programming-learning platforms. The review revealed that '
     'interactive, browser-based learning with automated feedback significantly improves '
     'programming education, and that gamification and certification enhance motivation. '
     'However, limitations such as paywalls, weak mastery enforcement, and the fragmentation of '
     'learning tools remain evident. These gaps justify the development of the proposed '
     'mayor4code platform, which integrates structured lessons, locked progression, an '
     'interactive playground, automated assessment, certification, and a leaderboard within a '
     'single secure web application.')
# ==================== CHAPTER THREE ====================
chapter(['Chapter Three', 'Methodology'])

h2('3.1 Introduction')
body('This chapter presents the methodology adopted for the design and implementation of the '
     'mayor4code web-based interactive e-learning platform. It describes the design approach, '
     'the considerations made during system development, the system architecture and its '
     'components, the database design, and the operational flow of the system. The methodology '
     'provides a structured framework for developing a reliable, secure, and user-friendly '
     'platform for structured Python learning.')

h2('3.2 Project Design and Approach')
body('The development of the mayor4code platform followed the Software Development Life Cycle '
     '(SDLC) using the Agile development approach. The Agile methodology was adopted because it '
     'supports iterative development, continuous improvement, flexibility, and regular testing '
     'throughout the development process. Features were developed incrementally and refined '
     'based on repeated testing.')
body('The system was designed as a web-based application accessible through a web browser. It '
     'defines two principal roles: the Learner (student), who registers, studies lessons, takes '
     'quizzes, uses the code playground, and earns certificates; and the Administrator, who '
     'manages lessons, quizzes, questions, and user progress through a dedicated administrative '
     'interface.')
body('The system adopts a three-tier architecture consisting of the presentation layer, the '
     'application layer, and the database layer. This separation promotes scalability, '
      'maintainability, and a clear division of responsibilities within the system.')

h2('3.3 Requirements Specification')
body('The functional and non-functional requirements of the mayor4code platform were identified '
      'based on the research objectives and the gaps identified in the literature review. '
      'Table 3.2 lists the functional requirements, specifying the core operations the system '
      'must perform.')
body('Table 3.2: Functional Requirements Specification')
fr = doc.add_table(rows=9, cols=3)
fr.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr.style = 'Table Grid'
f_headers = ['ID', 'Requirement', 'Description']
f_rows = [
    ['FR1', 'User Registration', 'The system shall allow new users to create an account with username, email, and password.'],
    ['FR2', 'User Authentication', 'The system shall authenticate registered users and manage sessions securely.'],
    ['FR3', 'Lesson Delivery', 'The system shall present structured Python lessons in sequential order.'],
    ['FR4', 'Locked Progression', 'The system shall unlock the next lesson only after the learner passes the current quiz at 60% or higher.'],
    ['FR5', 'Quiz Scoring', 'The system shall score multiple-choice quizzes automatically and display results instantly.'],
    ['FR6', 'Code Execution', 'The system shall execute Python code submitted via the playground in an isolated subprocess.'],
    ['FR7', 'Certificate Issuance', 'The system shall automatically issue a certificate with a unique code upon course completion.'],
    ['FR8', 'Leaderboard', 'The system shall display a ranked list of learners based on their average performance.'],
]
for c, h in enumerate(f_headers):
    cell = fr.cell(0, c)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(9)
for ri, row in enumerate(f_rows):
    for ci, val in enumerate(row):
        cell = fr.cell(ri + 1, ci)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(val)
        r.font.size = Pt(9)

body('Table 3.3 lists the non-functional requirements, specifying the quality attributes the '
      'system must satisfy.')
body('Table 3.3: Non-Functional Requirements Specification')
nfr = doc.add_table(rows=6, cols=3)
nfr.alignment = WD_ALIGN_PARAGRAPH.CENTER
nfr.style = 'Table Grid'
nf_headers = ['ID', 'Requirement', 'Description']
nf_rows = [
    ['NFR1', 'Security', 'Passwords must be hashed; playground code must execute in an isolated subprocess with a 5-second timeout.'],
    ['NFR2', 'Usability', 'The interface must be responsive across desktop, tablet, and mobile devices with dark/light mode support.'],
    ['NFR3', 'Performance', 'Page loads must complete within 3 seconds; quiz scoring must be immediate upon submission.'],
    ['NFR4', 'Scalability', 'The system must support concurrent access by multiple learners without degradation.'],
    ['NFR5', 'Reliability', 'Learner progress and scores must be persisted accurately; the system must maintain data integrity.'],
]
for c, h in enumerate(nf_headers):
    cell = nfr.cell(0, c)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(9)
for ri, row in enumerate(nf_rows):
    for ci, val in enumerate(row):
        cell = nfr.cell(ri + 1, ci)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(val)
        r.font.size = Pt(9)

h2('3.4 Design Consideration')
body('Several factors were considered during the design of the proposed system to ensure '
     'effectiveness and usability.')
h3('Security')
body('The system was designed with secure user authentication and authorisation. Passwords are '
     'hashed, sessions are protected, and requests are guarded against cross-site request '
     'forgery. The code playground executes learner code in an isolated subprocess with an '
     'enforced time limit to protect the server.')
h3('Usability')
body('A clean, responsive, and intuitive interface was designed so that learners can navigate '
     'the platform easily. A dark and light mode toggle was provided to improve reading comfort, '
     'with the preference persisting across sessions.')
h3('Accessibility')
body('The platform is web-based and can be accessed through computers, tablets, and mobile '
     'devices connected to the internet, ensuring learning is not restricted to a single '
     'location or device.')
h3('Scalability')
body('The system was designed to accommodate an increasing number of learners and content. The '
     'use of the Django framework and a relational database supports growth without significant '
     'redesign, and PostgreSQL is employed in the production environment for robustness.')
h3('Reliability')
body('The database schema and application logic were designed to ensure data integrity and '
     'consistent performance, so that learner progress and assessment results are recorded '
     'accurately.')
h3('Performance')
body('Static assets are served efficiently, and the locked-progression logic and quiz scoring '
     'were implemented to respond promptly, providing learners with a responsive experience.')

h2('3.5 System Architecture')
body('The proposed mayor4code platform adopts a three-tier architecture model consisting of the '
     'Presentation Layer, the Application Layer, and the Database Layer.')
body('The Presentation Layer is the interface through which users interact with the system using '
     'a web browser. It is built with HTML, CSS, and JavaScript and provides pages for '
     'registration, login, the dashboard, lessons, quizzes, the code playground, the '
     'leaderboard, and certificates.')
body('The Application Layer contains the business logic, implemented using the Django framework '
     'in Python. It processes user requests, manages authentication and authorisation, enforces '
     'locked lesson progression, scores quizzes, executes playground code in a controlled '
     'subprocess, issues certificates, and computes leaderboard rankings.')
body('The Database Layer stores all persistent data, including user accounts, lesson content, '
     'quizzes and questions, learner progress records, and issued certificates. SQLite is used '
     'during development and PostgreSQL in production. Interaction among these three layers '
     'ensures secure, efficient communication between users and the system.')
body('Figure 3.1: System Architecture of the mayor4code Platform')

h2('3.6 Components of System Architecture')
body('The architecture of the proposed system consists of the following major components:')
h3('User Interface Module')
body('This component provides the graphical interface through which learners and the '
     'administrator interact with the system, including forms, dashboards, lesson pages, quiz '
     'pages, and the code playground.')
h3('Authentication Module')
body('This module handles user registration, login, logout, and password reset. It verifies '
     'credentials and controls access based on the user’s role and authentication status.')
h3('Lesson and Progression Module')
body('This module delivers the twelve structured Python lessons and enforces locked '
     'progression, unlocking each lesson only after the learner passes the preceding lesson’s '
     'quiz at or above the defined pass mark of sixty percent.')
h3('Quiz and Assessment Module')
body('This module presents multiple-choice quizzes one question at a time, scores responses '
     'automatically, and records the results, providing learners with instant feedback on their '
     'performance.')
h3('Code Playground Module')
body('This module provides an in-browser environment in which learners can write and execute '
     'Python code. Submitted code is run in an isolated subprocess with an enforced execution '
     'time limit and safeguards against destructive operations.')
h3('Certification and Leaderboard Module')
body('This module automatically issues a completion certificate bearing a unique verification '
     'code when a learner completes all lessons, and it ranks learners on a leaderboard '
     'according to their average performance.')
h3('Administration Module')
body('This module allows the administrator to manage lessons, quizzes, questions, users, and '
     'progress records, and to monitor overall system activity.')

h2('3.7 Database Design')
body('The database was designed to store all persistent information required by the system in a '
     'structured and consistent manner. The principal entities and their relationships are '
     'summarised as follows:')
bullet('User: stores learner account details, including username, email, and hashed password.')
bullet('Lesson: stores each lesson’s title, content, and order in the sequence.')
bullet('Quiz: associated with a lesson and containing its set of questions.')
bullet('Question: stores the question text, options, and the correct answer for a quiz.')
bullet('UserLessonProgress: records each learner’s status and score for each lesson, supporting '
       'locked progression.')
bullet('Certificate: stores issued certificates, each with a unique verification code linked to '
       'the learner.')
body('A user may have many progress records and one certificate; a lesson has one associated '
     'quiz; and a quiz has many questions. These relationships preserve data integrity and '
     'support efficient retrieval of learner progress and assessment results.')
body('Table 3.1: Summary of Principal Database Entities')

h2('3.8 Flowchart of the System')
body('The system begins when a user accesses the platform through a web browser. New users '
     'register for an account, while returning users log in with their credentials. Upon '
     'successful authentication, the learner is directed to a dashboard that displays overall '
     'progress and the available lessons.')
body('The learner studies the first lesson and then attempts its quiz. The system scores the '
     'quiz automatically; if the learner attains at least the pass mark, the next lesson is '
     'unlocked, otherwise the learner may review the lesson and re-attempt the quiz. At any '
     'time, the learner may use the code playground to write and execute Python code and view '
     'the output instantly.')
body('When a learner successfully completes all lessons, the system automatically issues a '
     'completion certificate with a unique verification code, and the learner’s performance is '
     'reflected on the leaderboard. The administrator, through a separate interface, manages '
     'lessons, quizzes, questions, users, and progress records. The process continues until the '
     'user completes the intended operation and logs out, which terminates the active session '
     'and ensures system security.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\mayor4code\flowchart.png',
       'Figure 3.2: System Flowchart of the mayor4code Platform')
# ==================== CHAPTER FOUR ====================
chapter(['Chapter Four', 'System Implementation, Results and Discussion'])

h2('4.1 Introduction')
body('This chapter presents the implementation of the mayor4code web-based interactive '
     'e-learning platform, the development tools and environment used, the system interfaces, '
     'the testing procedures adopted, and a discussion of the results obtained. The discussion '
     'relates the findings to the objectives set out in Chapter One and to the gaps identified '
     'in the literature reviewed in Chapter Two.')

h2('4.2 System Implementation')
body('The system was implemented as a web application following the three-tier architecture '
     'and design described in Chapter Three. The back-end was developed using the Django web '
     'framework in Python, which provided the object-relational mapper, authentication '
     'facilities, and request-handling logic. The front-end was implemented using HTML, CSS, '
     'and JavaScript to deliver a responsive and interactive user interface. SQLite served as '
     'the database during development, while PostgreSQL was configured for the production '
     'environment.')
body('The application logic was organised into modular components corresponding to the '
     'authentication, lesson and progression, quiz and assessment, code playground, '
     'certification and leaderboard, and administration functions. Static assets were served '
     'efficiently in production using the WhiteNoise library, and application configuration was '
     'managed through environment variables to separate settings from source code and protect '
      'sensitive credentials.')

h3('4.2.1 Interface Feature Description')
body('Table 4.2 describes the key user interface components for each major screen of the '
      'mayor4code platform, detailing the UI elements, user actions, and corresponding system '
      'responses.')
body('Table 4.2: Interface Features by Screen')

ift = doc.add_table(rows=10, cols=4)
ift.alignment = WD_ALIGN_PARAGRAPH.CENTER
ift.style = 'Table Grid'
ift_headers = ['Screen', 'UI Elements', 'User Action', 'System Response']
ift_rows = [
    ['Registration', 'Username, email, password fields; show/hide toggle; submit button',
     'Enter details and click register',
     'Creates account, logs in, redirects to dashboard'],
    ['Login', 'Username and password fields; submit button; Google sign-in button',
     'Enter credentials and click login',
     'Authenticates user, redirects to dashboard'],
    ['Dashboard', 'Progress bar; lesson list with status indicators; sidebar navigation',
     'Click a lesson card',
     'Loads lesson content or shows lock message'],
    ['Lesson', 'Lesson content; previous/next navigation; complete button',
     'Read content and click to proceed',
     'Displays quiz or marks lesson as complete'],
    ['Quiz', 'Question text; option list (radio buttons); countdown timer; submit button',
     'Select answer and submit before timer expires',
     'Scores answer, advances to next question, shows result'],
    ['Playground', 'Code editor; run button; output panel; character counter',
     'Type code and click run',
     'Executes code in subprocess, returns output or error'],
    ['Certificate', 'Certificate design; student name; course title; verification code',
     'View certificate; copy verification code',
     'Displays issued certificate with unique code'],
    ['Leaderboard', 'Ranked table; username; average score; completed lessons; attempts',
     'Browse leaderboard',
     'Renders sorted list of all learners'],
    ['Admin', 'Lesson/quiz/question CRUD forms; user progress table',
     'Add, edit, or delete content',
     'Updates database and refreshes interface'],
]
for c, h in enumerate(ift_headers):
    cell = ift.cell(0, c)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(9)
for ri, row in enumerate(ift_rows):
    for ci, val in enumerate(row):
        cell = ift.cell(ri + 1, ci)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(val)
        r.font.size = Pt(8)

h2('4.3 Development Tools and Environment')
body('The tools and technologies used in the development of the system are summarised in Table '
     '4.1.')
body('Table 4.1: Software Development Tools')
bullet('Python — Core back-end programming language.')
bullet('Django — High-level Python web framework for the application layer.')
bullet('SQLite / PostgreSQL — Relational database management systems (development / production).')
bullet('HTML, CSS, JavaScript — Front-end structure, styling, and interactivity.')
bullet('WhiteNoise — Efficient static-file serving in production.')
bullet('Git — Version control for source-code management.')
bullet('Visual Studio Code — Source-code editor used for development.')
bullet('Web browser — Client platform for accessing and testing the application.')

h2('4.4 System Interfaces')
body('This section presents the major interfaces of the implemented system and describes the '
     'function of each.')
h3('4.4.1 Registration and Login Interfaces')
body('The registration interface allows a new learner to create an account by supplying a '
     'username, email address, and password, with a show/hide password toggle for convenience. '
     'The login interface authenticates returning users before granting access to the '
      'protected areas of the platform, as shown in Figure 4.1 and Figure 4.2.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\mayor4code\03-signup.png',
       'Figure 4.1: Registration Interface of the mayor4code Platform')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\mayor4code\02-login.png',
       'Figure 4.2: Login Interface of the mayor4code Platform')
h3('4.4.2 Dashboard Interface')
body('Upon successful login, the learner is presented with a dashboard, shown in Figure 4.3, '
     'that displays an overall progress bar and the list of lessons, indicating which lessons '
      'are unlocked, completed, or still locked.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\mayor4code\05-dashboard.png',
       'Figure 4.3: Learner Dashboard of the mayor4code Platform')
h3('4.4.3 Lesson and Quiz Interfaces')
body('The lesson interface, shown in Figure 4.4, presents lesson content with previous and next '
     'navigation. On completing a lesson, the learner attempts its quiz through the quiz '
     'interface, shown in Figure 4.5, which presents one multiple-choice question at a time and '
      'displays the score instantly upon submission.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\mayor4code\06-lessons.png',
       'Figure 4.4: Lesson Interface of the mayor4code Platform')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\mayor4code\08-quizzes.png',
       'Figure 4.5: Quiz Interface of the mayor4code Platform')
h3('4.4.4 Python Playground Interface')
body('The playground interface, shown in Figure 4.6, provides an in-browser editor where the '
     'learner writes Python code and executes it. The code is run on the server in an isolated '
     'subprocess with an enforced time limit, and the output is returned and displayed to the '
      'learner instantly.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\mayor4code\07-playground.png',
       'Figure 4.6: Python Playground Interface of the mayor4code Platform')
h3('4.4.5 Certificate and Leaderboard Interfaces')
body('When a learner completes all lessons, the system automatically issues a completion '
     'certificate bearing a unique verification code, shown in Figure 4.7. The leaderboard '
      'interface, shown in Figure 4.8, ranks learners according to their average performance.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\mayor4code\11-certificates.png',
       'Figure 4.7: Completion Certificate of the mayor4code Platform')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\mayor4code\09-leaderboard.png',
       'Figure 4.8: Leaderboard Interface of the mayor4code Platform')
h3('4.4.6 Administration Interface')
body('The administration interface, shown in Figure 4.9, enables the administrator to manage '
     'lessons, quizzes, questions, users, and progress records, and to monitor overall system '
     'activity.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\mayor4code\10-progress.png',
       'Figure 4.9: Administration Interface of the mayor4code Platform')

h2('4.5 System Testing')
body('The system was tested to verify that it met its functional requirements and behaved '
     'reliably. Three levels of testing were carried out: unit testing of individual '
     'components, functional testing of complete features, and usability testing with sample '
     'users.')
h3('4.5.1 Unit Testing')
body('Individual functions, such as quiz scoring, progression unlocking, and certificate-code '
     'generation, were tested in isolation to confirm that each produced the expected output '
     'for representative inputs.')
h3('4.5.2 Functional Testing')
body('Complete features were tested end to end, including registration and login, lesson '
      'navigation, quiz submission and scoring, locked progression, code execution in the '
      'playground, certificate issuance, and leaderboard ranking. The results are summarised in '
      'Table 4.3.')
body('Table 4.3: Summary of Test Cases and Results')

tt = doc.add_table(rows=22, cols=5)
tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
tt.style = 'Table Grid'
tt_headers = ['Test ID', 'Test Description', 'Input/Condition', 'Expected Result', 'Actual Result']
tt_rows = [
    ['TC01', 'User Registration (valid)', 'New username, email, password', 'Account created; redirected to dashboard', 'Passed'],
    ['TC02', 'User Registration (duplicate)', 'Existing username', 'Error message displayed; no account created', 'Passed'],
    ['TC03', 'User Registration (empty fields)', 'Blank form submitted', 'Validation errors shown', 'Passed'],
    ['TC04', 'Login (valid credentials)', 'Correct username + password', 'Authenticated; redirected to dashboard', 'Passed'],
    ['TC05', 'Login (invalid password)', 'Correct username + wrong password', 'Error message displayed', 'Passed'],
    ['TC06', 'Lesson progression (pass)', 'Submit quiz with score >= 60%', 'Next lesson unlocked', 'Passed'],
    ['TC07', 'Lesson progression (fail)', 'Submit quiz with score < 60%', 'Next lesson remains locked; retry allowed', 'Passed'],
    ['TC08', 'Lesson progression (boundary)', 'Submit quiz with exactly 60%', 'Next lesson unlocked (boundary pass)', 'Passed'],
    ['TC09', 'First lesson access', 'New user clicks lesson 1', 'Lesson 1 accessible (no prerequisite)', 'Passed'],
    ['TC10', 'Playground (valid code)', "Submit 'print(\"Hello\")'", 'Output: "Hello" returned', 'Passed'],
    ['TC11', 'Playground (unsafe import)', "Submit 'import os'", '403 error; code blocked', 'Passed'],
    ['TC12', 'Playground (timeout)', 'Submit infinite loop code', 'Timeout error after 5 seconds', 'Passed'],
    ['TC13', 'Playground (long code)', 'Submit code > 10,000 characters', '400 error; length limit enforced', 'Passed'],
    ['TC14', 'Quiz scoring (all correct)', 'Answer all questions correctly', 'Score = 100%', 'Passed'],
    ['TC15', 'Quiz scoring (mixed)', 'Answer some correctly, some wrong', 'Correct percentage computed', 'Passed'],
    ['TC16', 'Certificate issuance (completed)', 'Complete all 12 lessons', 'Certificate with unique code issued', 'Passed'],
    ['TC17', 'Certificate issuance (incomplete)', 'Complete < 12 lessons', 'No certificate issued', 'Passed'],
    ['TC18', 'Leaderboard ranking', 'Two learners with different scores', 'Higher score ranked first', 'Passed'],
    ['TC19', 'Admin lesson CRUD', 'Add, edit, delete a lesson', 'Database updated; reflected in UI', 'Passed'],
    ['TC20', 'Admin quiz management', 'Add questions to a quiz', 'Questions appear in learner quiz', 'Passed'],
    ['TC21', 'Logout', 'Click logout button', 'Session terminated; redirected to login', 'Passed'],
]
for c, h in enumerate(tt_headers):
    cell = tt.cell(0, c)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(8)
for ri, row in enumerate(tt_rows):
    for ci, val in enumerate(row):
        cell = tt.cell(ri + 1, ci)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(val)
        r.font.size = Pt(8)
h3('4.5.3 Usability Testing')
body('Sample users interacted with the platform and provided feedback on ease of use. Users '
     'found the interface intuitive and responsive, and the dark and light mode toggle was '
     'reported to improve reading comfort. Minor observations were addressed during iterative '
      'refinement.')

h3('4.5.4 Implementation Challenges')
body('During the development of the mayor4code platform, several implementation challenges were '
      'encountered and addressed.')
body('The first challenge was securing the in-browser Python playground. Executing arbitrary user '
      'code on the server poses significant security risks, including the potential for system '
      'commands, file system access, and resource exhaustion. This was addressed by running user '
      'submissions in an isolated subprocess with a five-second execution timeout, blocking '
      'dangerous imports and function calls through string matching on the lowercased source code, '
      'and enforcing a ten-thousand-character code length limit. The approach provides reasonable '
      'security for an educational context while maintaining a responsive user experience, though '
      'container-based sandboxing would offer stronger isolation for production deployment.')
body('The second challenge was implementing locked progression without creating a poor user '
      'experience. The system must correctly handle edge cases such as the first lesson being '
      'always unlocked (no prerequisite), administrator access bypassing progression locks, and '
      'users who complete all lessons being allowed to revisit previous content. These cases were '
      'handled by default-unlocking lesson one in the view logic, granting administrator accounts '
      'full access through a separate permission check, and allowing completed lessons to be '
      'reviewed without quiz re-attempt constraints.')
body('A third challenge related to the quiz timer synchronisation. Since the quiz timer is '
      'implemented client-side with JavaScript, a determined learner could theoretically tamper '
      'with the timer value. The elapsed time is captured in a hidden input field submitted with '
      'the quiz, and the server records it as time_taken for analytics rather than for enforcing '
      'time limits, avoiding false disqualifications. Future versions could implement server-side '
      'timer validation for stricter time enforcement.')
body('Finally, generating the project report itself presented challenges. Producing a '
      'professionally formatted Word document with mixed page numbering (Roman numerals for '
      'preliminary pages, Arabic numerals for the main body), auto-generated table of contents, '
      'embedded figures, and formatted tables required careful use of the python-docx library\'s '
      'section and field-code APIs. These techniques, while not part of the e-learning platform '
      'itself, were essential for producing the academic documentation required for the BSc '
      'project submission.')

h2('4.6 Discussion of Results')
body('The results confirmed that the platform met the objectives established in Chapter One. '
     'The system successfully delivered structured lessons with locked progression, provided an '
     'integrated in-browser coding environment, scored assessments automatically with instant '
     'feedback, tracked learner progress, and issued valid completion certificates alongside a '
     'motivational leaderboard.')
body('In relation to the research objectives, each objective was realised: a structured, '
     'web-based lesson-delivery system was designed and implemented; a secure authentication '
     'mechanism was provided; an interactive Python playground was implemented; automated '
     'assessment with instant scoring was achieved; and progress tracking, certification, and a '
     'leaderboard were incorporated to enhance motivation. The functionality and usability of '
     'the platform were evaluated through testing.')
body('In relation to the existing literature, the platform addresses the gaps identified in '
     'Chapter Two. Unlike several reviewed systems that restrict structured learning paths '
     'behind paywalls or fail to enforce mastery, mayor4code provides free, structured '
     'progression that enforces understanding of prerequisite concepts before advancement. By '
     'integrating lessons, an interactive playground, automated assessment, certification, and '
     'a leaderboard within a single coherent system, it also overcomes the fragmentation of '
     'tools observed in some existing platforms. These findings demonstrate the value of '
     'combining structured content, practical coding, and gamification in a single secure web '
     'application for introductory programming education.')

# ==================== CHAPTER FIVE ====================
chapter(['Chapter Five', 'Summary, Conclusion and Recommendations'])

h2('5.1 Summary of Findings')
body('This project set out to design and implement a web-based interactive e-learning platform '
     'for structured Python programming. The problem addressed was the difficulty beginners '
     'face in learning to program, arising from unstructured content, the separation of theory '
     'from practice, and the lack of immediate feedback, progress tracking, and recognition of '
     'achievement in many existing resources.')
body('A review of relevant literature and existing platforms established the effectiveness of '
     'interactive, browser-based learning with automated feedback and gamification, while '
     'revealing gaps such as paywalls, weak enforcement of mastery, and the fragmentation of '
     'learning tools. Guided by these findings, the system was developed using the Agile '
     'approach and a three-tier architecture, implemented with the Django framework, Python, '
     'HTML, CSS, JavaScript, and a relational database.')
body('The developed platform delivered twelve sequential Python lessons with locked '
     'progression, multiple-choice quizzes with instant scoring, an in-browser Python '
     'playground with safe code execution, progress tracking, automatically issued completion '
     'certificates, and a leaderboard. Testing confirmed that the system enforced progression '
     'correctly, scored assessments accurately, executed code safely, and issued valid '
      'certificates.')

h2('5.2 Contributions to Knowledge')
body('This study makes the following contributions to knowledge in the field of interactive '
      'e-learning for programming education:')
numbered('An integrated e-learning platform that combines structured lesson delivery, an '
         'in-browser code playground, automated assessment, progress tracking, certification, '
         'and gamification within a single free and secure web application. Unlike existing '
         'platforms that either separate these features across multiple tools or restrict '
         'key functionality behind paywalls, mayor4code provides a unified learning environment '
         'that addresses the fragmentation identified in the literature review.')
numbered('A mastery-based locked progression model that enforces a minimum quiz score of sixty '
         'percent before advancing to the next lesson. This operationalisation of mastery learning '
         'principles (Bloom, 1968; Guskey, 2007) within a web-based programming context ensures '
         'that learners build a solid foundation of prerequisite knowledge before encountering '
         'more advanced concepts, reducing the accumulation of knowledge gaps.')
numbered('A secure in-browser code execution environment that runs learner-submitted Python code '
         'in an isolated subprocess with a five-second timeout, blocked dangerous operations, '
         'and a ten-thousand-character limit. This provides a practical reference implementation '
         'for safely embedding code execution in educational web applications without requiring '
         'containerisation infrastructure.')
numbered('An automated certification system that generates unique verification codes for each '
         'completion certificate, enabling third-party verification of learner achievement. '
         'The certificate design and issuance workflow demonstrate how cryptographic-grade unique '
         'identifiers can be integrated into educational platforms for credentialing purposes.')
numbered('A practical demonstration of the Django web framework\'s suitability for building '
         'secure, scalable, and feature-rich educational web applications, including the '
         'integration of authentication, content management, assessment, and gamification '
         'features within a coherent software architecture.')

h2('5.3 Conclusion')
body('The study successfully achieved its aim of designing and implementing an integrated, '
     'web-based interactive e-learning platform for structured Python programming. By combining '
     'structured lesson delivery, mastery-based progression, practical in-browser coding, '
     'automated assessment, certification, and gamification within a single secure application, '
     'the platform improves the accessibility and effectiveness of introductory programming '
     'education. The system therefore provides a practical solution to the challenges that '
     'beginners commonly encounter when learning to program.')

h2('5.4 Recommendations')
body('Based on the outcomes of this study, the following recommendations are made:')
numbered('Educational institutions and training organisations should adopt integrated, '
         'interactive e-learning platforms of this kind to complement traditional programming '
         'instruction.')
numbered('The principle of mastery-based locked progression should be applied in programming '
         'education to reduce knowledge gaps among learners.')
numbered('Adequate attention should be given to the security of in-browser code-execution '
         'environments, including strong isolation of executed code, to protect the hosting '
         'infrastructure.')
numbered('Learner motivation should be supported through gamification features such as '
         'certification and leaderboards.')

h2('5.5 Suggestions for Future Research')
body('The following suggestions are offered for future work:')
numbered('Extending the platform to cover additional programming languages and more advanced '
         'topics beyond introductory Python.')
numbered('Incorporating adaptive learning that personalises content and pacing based on each '
         'learner’s performance.')
numbered('Strengthening the code-execution environment through container-based sandboxing with '
         'stricter resource, network, and filesystem controls.')
numbered('Adding richer assessment types, such as automatically graded coding exercises, in '
         'addition to multiple-choice quizzes.')
numbered('Conducting a large-scale empirical evaluation of learning outcomes to measure the '
         'platform’s educational impact quantitatively.')

# ==================== REFERENCES ====================
chapter(['References'])
refs = [
    'Anderson, T. (2008). The theory and practice of online learning (2nd ed.). Athabasca '
    'University Press.',
    'Bloom, B. S. (1968). Mastery learning. Evaluation Comment, 1(2), 1–12.',
    'Deterding, S., Dixon, D., Khaled, R., & Nacke, L. (2011). From game design elements to '
    'gamefulness: Defining "gamification". In Proceedings of the 15th International Academic '
    'MindTrek Conference (pp. 9–15). Association for Computing Machinery.',
    'Django Software Foundation. (2024). Django documentation. https://docs.djangoproject.com/',
    'Duckett, J. (2014). HTML and CSS: Design and build websites. John Wiley & Sons.',
    'Elmasri, R., & Navathe, S. B. (2017). Fundamentals of database systems (7th ed.). Pearson.',
    'Ericsson, K. A., Krampe, R. T., & Tesch-Romer, C. (1993). The role of deliberate practice '
    'in the acquisition of expert performance. Psychological Review, 100(3), 363–406.',
    'Garrison, D. R. (2017). E-learning in the 21st century: A community of inquiry framework '
    'for research and practice (3rd ed.). Routledge.',
    'Guskey, T. R. (2007). Closing achievement gaps: Revisiting Benjamin S. Bloom\u2019s "Learning '
    'for Mastery". Journal of Advanced Academics, 19(1), 8–31.',
    'Hamari, J., Koivisto, J., & Sarsa, H. (2014). Does gamification work? A literature review '
    'of empirical studies on gamification. In Proceedings of the 47th Hawaii International '
    'Conference on System Sciences (pp. 3025–3034). IEEE.',
    'Jonassen, D. H. (1999). Designing constructivist learning environments. In C. M. Reigeluth '
    '(Ed.), Instructional-design theories and models: A new paradigm of instructional theory '
    '(Vol. 2, pp. 215–239). Lawrence Erlbaum Associates.',
    'Keller, F. S. (1968). "Good-bye, teacher..." Journal of Applied Behavior Analysis, 1(1), 79–89.',
    'Piaget, J. (1952). The origins of intelligence in children. International Universities Press.',
    'Pressman, R. S., & Maxim, B. R. (2020). Software engineering: A practitioner\u2019s approach '
    '(9th ed.). McGraw-Hill Education.',
    'Python Software Foundation. (2024). The Python language reference. '
    'https://docs.python.org/3/reference/',
    'Sailer, M., Hense, J. U., Mayr, S. K., & Mandl, H. (2017). How gamification motivates: '
    'An experimental study of the effects of specific game design elements on psychological '
    'need satisfaction. Computers in Human Behavior, 69, 371–380.',
    'Sommerville, I. (2016). Software engineering (10th ed.). Pearson Education.',
    'Vygotsky, L. S. (1978). Mind in society: The development of higher psychological processes. '
    'Harvard University Press.',
]
for r in refs:
    p = doc.add_paragraph(r)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Mm(-12.7)
    p.paragraph_format.left_indent = Mm(12.7)
# CONTENT_END

out = r'c:\Users\ALEXIS\Desktop\SENPAI\mayor4code_Project_Report.docx'
doc.save(out)
print('saved', out)
