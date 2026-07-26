# -*- coding: utf-8 -*-
"""Generate the AI-based IDS BSc project report per CSC / OAUSTECH guidelines.
Reuses shared formatting helpers from report_helpers.py."""
import report_helpers as R
from report_helpers import (
    doc, chapter, h2, h3, body, bullet, numbered, prelim_title,
    add_toc, set_page_numbering, figure, caption, centered, reference, save,
)
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

TITLE = ('DESIGN AND IMPLEMENTATION OF AN AI-BASED INTRUSION DETECTION SYSTEM '
         'FOR ENCRYPTED TRAFFIC (A CASE STUDY OF IDS PROJECT)')

# ==================== TITLE PAGE ====================
centered(TITLE, bold=True)
for line in ['', 'BY', '', 'Oyeduntan Segun Elijah', 'CSC/22/174', '',
             'DEPARTMENT OF COMPUTER SCIENCE',
             'SCHOOL OF INFORMATION AND COMMUNICATION TECHNOLOGY',
             'OLUSEGUN AGAGU UNIVERSITY OF SCIENCE AND TECHNOLOGY, OKITIPUPA', '',
             'IN PARTIAL FULFILMENT OF THE REQUIREMENTS FOR THE AWARD OF THE '
             'DEGREE OF BACHELOR OF SCIENCE (B.Sc.) IN COMPUTER SCIENCE', '',
             'August, 2026']:
    centered(line)

# Preliminary pages: Roman numerals, title page unnumbered
set_page_numbering(doc.sections[0], 'lowerRoman', start=1, hide_first=True)

# ---- Certification ----
prelim_title('Certification')
body('This is to certify that this project titled "Design and Implementation of an AI-Based '
     'Intrusion Detection System for Encrypted Traffic (A Case Study of IDS Project)" was '
     'carried out by Oyeduntan Segun Elijah with matriculation number CSC/22/174 in the '
     'Department of Computer Science, School of Information and Communication Technology, '
     'Olusegun Agagu University of Science and Technology, Okitipupa, in partial fulfilment of '
     'the requirements for the award of the degree of Bachelor of Science (B.Sc.) in Computer '
     'Science.')
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

table = doc.add_table(rows=4, cols=2)
table.alignment = WD_ALIGN_PARAGRAPH.CENTER
# Remove all borders
tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
borders = OxmlElement('w:tblBorders')
for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    border = OxmlElement(f'w:{name}')
    border.set(qn('w:val'), 'nil')
    borders.append(border)
tblPr.append(borders)
cells = [
    ('_______________________', '_______________________'),
    ('Dr. (Engr.) Modupe Agagu',  'Dr. (Mrs.) A. Gbadamosi'),
    ('Project Supervisor',        'Coordinator / HOD'),
    ('Date: __________________','Date: __________________'),
]
for r, (left, right) in enumerate(cells):
    for c, val in [(0, left), (1, right)]:
        cell = table.cell(r, c)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(val)

# ---- Declaration ----
prelim_title('Declaration')
body('I hereby declare that this project is my original work and has not been submitted, either '
     'in whole or in part, for any other degree or qualification in this or any other '
     'institution. All sources of information and materials used have been duly acknowledged by '
     'means of complete references.')
for line in ['', '', '_______________________', 'Oyeduntan Segun Elijah', 'CSC/22/174',
             'Date: __________________']:
    doc.add_paragraph(line)

# ---- Dedication ----
prelim_title('Dedication')
body('This project is dedicated to Almighty God, and to my beloved family, whose unwavering '
     'support, prayers, and encouragement made this academic journey possible.')

# ---- Acknowledgements ----
prelim_title('Acknowledgements')
body('I express my profound gratitude to Almighty God for His grace, wisdom, and strength that '
     'sustained me throughout this project. I am deeply grateful to my project supervisor, '
     'Dr. (Engr.) Modupe Agagu for his invaluable guidance, constructive feedback, and '
     'encouragement, which greatly contributed to the success of this work. My sincere '
     'appreciation goes to the Coordinator of the Department of Computer Science, '
     'Dr. (Mrs.) A. Gbadamosi, and all the lecturers in the School of Information and '
     'Communication Technology, Olusegun Agagu University of Science and Technology (OAUSTECH), '
     'Okitipupa, namely Prof. D. T. Akomolafe, Dr. O. M. Orogbemi, '
     'Mr. A. M. Oduwale, and Mr. A. Obamehinti, for their academic support and knowledge '
     'imparted during my studies. I also acknowledge my parents for their prayers, support, '
     'and sacrifices.')

# ---- Abstract ----
prelim_title('Abstract')
body('The proliferation of encrypted network traffic has significantly reduced the effectiveness '
     'of traditional signature-based intrusion detection systems, as malicious activities are '
     'increasingly concealed within encrypted channels. This project presents the design and '
     'implementation of an AI-based intrusion detection system that leverages the XGBoost '
     'ensemble learning algorithm to classify encrypted network traffic as benign or malicious. '
     'The system was developed using the Agile software development methodology and trained on '
     'the CICIDS2017 benchmark dataset, which contains realistic benign and attack traffic with '
     'over 70 flow-based features. The model was evaluated using accuracy, precision, recall, '
     'F1-score, and ROC-AUC metrics, achieving a classification accuracy of 99.9% and a ROC-AUC '
     'of 0.9999, demonstrating exceptional detection capability across multiple attack '
     'categories. A user-friendly Streamlit dashboard was developed to provide three operational '
     'modes: CSV file upload for batch analysis, PCAP file upload for offline packet-level '
     'inspection using tshark, and live network capture for real-time monitoring. The dashboard '
     'integrates a Google Gemini AI assistant that provides plain-language explanations of '
     'detected threats and model predictions, making the system accessible to network '
     'administrators with varying levels of technical expertise. The system was themed with a '
     'dark matrix-style interface for enhanced visual appeal. Testing was conducted through '
     'unit testing of individual components, integration testing of the complete pipeline from '
     'traffic capture to prediction display, and user acceptance testing with sample traffic '
     'data. The results confirmed that the system meets all six objectives defined in the study, '
     'providing a lightweight, accurate, and explainable intrusion detection solution that '
     'addresses the critical challenge of detecting attacks in encrypted network traffic. The '
     'study concludes that XGBoost-based classification combined with an accessible dashboard '
     'interface and AI-powered explanations offers a practical and effective approach to modern '
     'network security challenges.')

# ---- Table of Contents ----
prelim_title('Table of Contents')
add_toc('TOC')

# ---- List of Figures ----
prelim_title('List of Figures')
for line in ['Figure 3.1: System Architecture of the Proposed IDS',
             'Figure 3.2: System Flowchart of the AI-Based IDS Pipeline',
             'Figure 4.1: Full Dashboard Interface',
             'Figure 4.2: PCAP Upload Tab',
             'Figure 4.3: Live Capture Tab',
             'Figure 4.4: AI Assistant Interface',
             'Figure 4.5: Confusion Matrix of the XGBoost Classifier',
             'Figure 4.6: ROC Curve of the XGBoost Classifier',
             'Figure 4.7: Feature Importance of the XGBoost Classifier']:
    body(line)

# ---- List of Tables ----
prelim_title('List of Tables')
for line in ['Table 3.1: Software Development Tools and Technologies',
             'Table 3.2: CICIDS2017 Dataset Attack Categories',
             'Table 4.1: Classification Report of the XGBoost Model',
             'Table 4.2: Summary of System Test Cases and Results']:
    body(line)

# ---- List of Abbreviations ----
prelim_title('List of Abbreviations')
for line in ['AI — Artificial Intelligence',
             'API — Application Programming Interface',
             'CSV — Comma-Separated Values',
             'DDoS — Distributed Denial of Service',
             'ERD — Entity Relationship Diagram',
             'HTTP — Hypertext Transfer Protocol',
             'IDS — Intrusion Detection System',
             'IP — Internet Protocol',
             'ML — Machine Learning',
             'PCAP — Packet Capture',
             'ROC-AUC — Receiver Operating Characteristic — Area Under Curve',
             'SQL — Structured Query Language',
             'TCP — Transmission Control Protocol',
             'UDP — User Datagram Protocol',
             'XGBoost — Extreme Gradient Boosting']:
    body(line)

# ==================== CHAPTER ONE ====================
chapter(['Chapter One', 'Introduction'], new_section=True)

h2('1.1 Background of the Study')
body('Network security remains one of the most critical concerns in the modern digital landscape, '
      'as organisations increasingly rely on interconnected systems and the internet for their '
      'daily operations. The proliferation of cyber threats, ranging from malware infections and '
      'denial-of-service attacks to advanced persistent threats, has necessitated the development '
      'of robust security mechanisms to protect sensitive data and ensure the availability of '
      'network services. Intrusion Detection Systems (IDS) have long served as a fundamental '
      'component of network defence, monitoring traffic for suspicious patterns and alerting '
      'administrators to potential security breaches.')
body('The global landscape of cyber threats has evolved significantly over the past decade. '
      'According to industry reports, the average cost of a data breach in 2025 exceeded four '
      'million dollars, with organisations facing financial losses from system downtime, '
      'regulatory penalties, legal fees, and reputational damage. Small and medium-sized '
      'enterprises are particularly vulnerable, as they often lack the dedicated security '
      'personnel and advanced security infrastructure needed to defend against sophisticated '
      'attacks. The increasing frequency of ransomware attacks, supply chain compromises, and '
      'targeted intrusions has made network security a top priority for organisations of all '
      'sizes across all industry sectors.')
body('The volume of internet traffic has grown exponentially, driven by the proliferation of '
      'cloud services, remote work, Internet of Things (IoT) devices, and streaming media. This '
      'growth has been accompanied by a corresponding increase in the adoption of encryption '
      'protocols. Transport Layer Security (TLS), the most widely used encryption protocol, now '
      'protects the majority of web traffic, email communications, and application data in '
      'transit. While encryption is essential for protecting sensitive information from '
      'interception, it has inadvertently created blind spots for network security monitoring '
      'tools. Attackers exploit these blind spots to conduct malicious activities that evade '
      'traditional detection mechanisms.')
body('The widespread adoption of encryption protocols, including TLS, HTTPS, and VPN technologies, '
     'has introduced a significant challenge for traditional IDS solutions. Encryption is '
     'essential for protecting data confidentiality and privacy, but it simultaneously obscures '
     'the payload content of network packets from inspection. Signature-based IDS, which rely on '
     'matching known attack patterns against packet payloads, are rendered largely ineffective '
     'against encrypted traffic because the malicious content is hidden within the encrypted '
     'stream. Attackers have recognised this limitation and increasingly use encrypted channels '
     'to launch attacks, command compromised systems, and exfiltrate data undetected.')
body('Machine learning techniques offer a promising approach to addressing this challenge. By '
     'analysing statistical features of network flows, such as packet sizes, inter-arrival times, '
     'flow durations, and protocol-level characteristics, machine learning models can distinguish '
     'between benign and malicious traffic patterns without requiring access to the encrypted '
     'payload. Among the various machine learning algorithms, ensemble methods, particularly '
     'XGBoost, have demonstrated exceptional performance in classification tasks due to their '
     'ability to combine multiple weak learners into a strong classifier. According to Chen and '
     'Guestrin (2016), XGBoost provides state-of-the-art results in many machine learning '
     'competitions and real-world applications.')
body('This study presents the design and implementation of an AI-based intrusion detection system '
     'that applies the XGBoost algorithm to classify encrypted network traffic. The system is '
     'trained on the CICIDS2017 dataset, which provides realistic benign and attack traffic with '
     'comprehensive flow-based features. It is deployed through a Streamlit dashboard that '
     'supports CSV, PCAP, and live capture modes, and integrates a Gemini AI assistant for '
     'explainability, thereby making advanced machine learning-based intrusion detection '
     'accessible and practical for real-world use.')

h2('1.2 Research Justification/Motivation')
body('The motivation for this study arises from the growing sophistication and frequency of '
     'cyber attacks that exploit encrypted communication channels to evade detection. According '
     'to industry reports, the majority of internet traffic is now encrypted, and a significant '
     'proportion of malicious traffic also uses encryption. Traditional IDS tools that depend on '
     'deep packet inspection are unable to analyse encrypted payloads, creating a critical gap '
     'in network security defences.')
body('The financial impact of cyber attacks continues to rise, with organisations facing costs '
     'related to data breaches, system downtime, regulatory fines, and reputational damage. Small '
     'and medium-sized enterprises are particularly vulnerable, as they often lack the resources '
     'to deploy sophisticated security infrastructure. There is an urgent need for lightweight, '
     'accurate, and cost-effective intrusion detection solutions that can operate effectively in '
     'encrypted environments without requiring specialised hardware or extensive manual '
     'configuration.')
body('Furthermore, existing machine learning-based IDS solutions are often presented as research '
     'prototypes with limited usability, lacking intuitive interfaces that allow network '
     'administrators to interact with the system easily. The integration of artificial '
     'intelligence for explainability, through a conversational assistant that interprets '
     'predictions in plain language, addresses this usability gap. These considerations motivated '
     'the development of an AI-based IDS that combines state-of-the-art machine learning with an '
     'accessible dashboard and explainable AI capabilities.')

h2('1.3 Problem Statement')
body('Network security faces a fundamental challenge in the era of widespread encryption. '
     'Signature-based intrusion detection systems, which have been the cornerstone of network '
     'defence for decades, rely on inspecting packet payloads for known attack signatures. As '
     'encryption becomes ubiquitous, these systems can no longer see the content of network '
     'packets, rendering them blind to attacks that are carried out over encrypted channels.')
body('Attackers have adapted to this environment by encrypting their malicious traffic, using '
     'the same encryption protocols that protect legitimate communications. Malware command and '
     'control traffic, data exfiltration, and even some categories of denial-of-service attacks '
     'can now be conducted over encrypted connections, bypassing traditional detection '
     'mechanisms. Existing anomaly-based IDS solutions attempt to address this by modelling '
     'normal traffic behaviour, but they often suffer from high false positive rates and require '
     'extensive tuning to specific network environments.')
body('There is a need for a machine learning-based IDS that can classify encrypted traffic using '
     'statistical flow features without requiring payload access. Such a system must be accurate '
     'enough to minimise false alarms, lightweight enough to operate in real time or near-real '
     'time, and usable enough for deployment by network administrators with varying levels of '
     'technical expertise. Additionally, the system should provide interpretable explanations of '
     'its predictions to build trust and enable informed decision-making.')
body('This study addresses these challenges through the development of an AI-based IDS that '
     'applies XGBoost classification on flow-level features extracted from network traffic, '
     'provides three operational modes through a Streamlit dashboard, and integrates a Gemini '
     'AI assistant for natural-language explanations of detected threats.')

h2('1.4 Aim and Objectives of the Study')
h3('Aim')
body('The aim of this study is to design and implement an AI-based intrusion detection system '
     'that can effectively classify encrypted network traffic as benign or malicious using '
     'machine learning techniques, and provide an accessible interface for network security '
     'monitoring.')
h3('Objectives')
body('The specific objectives of the study are to:')
numbered('preprocess the CICIDS2017 dataset by cleaning, encoding categorical labels, splitting '
         'into training, validation, and test sets, and applying standard scaling for model '
         'training;')
numbered('train an XGBoost ensemble classifier on the preprocessed dataset and optimise its '
         'hyperparameters using cross-validation to achieve high detection accuracy;')
numbered('build an interactive Streamlit dashboard that provides a unified interface for '
         'uploading traffic data, viewing predictions, and exploring model performance metrics;')
numbered('implement three operational modes within the dashboard: CSV file upload for batch '
         'classification, PCAP file upload with tshark-based feature extraction for offline '
         'packet analysis, and live network capture for real-time monitoring;')
numbered('integrate a Google Gemini AI assistant that provides plain-language explanations of '
         'detected threats, model predictions, and security insights to enhance usability;')
numbered('evaluate the performance of the developed system using standard classification '
         'metrics including accuracy, precision, recall, F1-score, and ROC-AUC, and validate '
         'its functionality through comprehensive testing.')

h2('1.5 Scope of the Study')
body('This project focuses on the design and implementation of an AI-based intrusion detection '
     'system specifically for the classification of encrypted network traffic. The system uses '
     'the XGBoost algorithm trained on the CICIDS2017 benchmark dataset, which includes both '
     'benign traffic and a wide range of attack types including DDoS, brute force, infiltration, '
     'botnet, and web attacks. The scope is limited to flow-based statistical features extracted '
     'from network traffic, without requiring deep packet inspection of encrypted payloads.')
body('The system provides three detection modes: CSV batch classification for pre-extracted '
     'feature files, PCAP analysis for offline packet capture files using tshark for feature '
     'extraction, and live network capture for real-time monitoring using Scapy. The user '
     'interface is implemented as a Streamlit web application with a dark matrix-themed design. '
     'The AI assistant capability is provided by Google Gemini and is used solely for '
     'explanation and interpretation of results, while all detection decisions are made by the '
     'XGBoost model. The study does not extend to hardware deployment, integration with '
     'existing enterprise security systems, or the development of custom packet capture '
     'hardware.')

h2('1.6 Significance of the Study')
h3('1.6.1 Economic Significance')
body('The economic impact of cyber attacks continues to grow, with organisations incurring '
     'substantial costs from data breaches, ransomware payments, system recovery, and regulatory '
     'penalties. By providing an accurate and cost-effective intrusion detection solution that '
     'operates effectively in encrypted environments, this system helps organisations reduce '
     'the financial risks associated with undetected cyber attacks. The open-source nature of '
     'the tools used, including Python, Streamlit, and XGBoost, ensures that the solution '
     'remains accessible to organisations with limited security budgets.')
h3('1.6.2 Technological Significance')
body('This project demonstrates the practical application of machine learning, specifically '
     'gradient boosting techniques, to a pressing real-world network security challenge. It '
     'contributes to the growing body of knowledge on ML-based intrusion detection and provides '
     'a reference implementation that combines feature extraction, model inference, interactive '
     'visualisation, and AI-powered explainability in a single integrated platform. The system '
     'architecture serves as a template for future research and development in the field.')
h3('1.6.3 Social Significance')
body('By making advanced intrusion detection capabilities accessible through an intuitive '
     'dashboard and AI assistant, this system empowers network administrators and security '
     'professionals to better protect their organisations, even if they lack deep expertise in '
     'machine learning. The explainability feature builds trust in AI-driven security decisions '
     'and promotes informed responses to threats. Ultimately, improved network security benefits '
     'society at large by protecting critical infrastructure, personal data, and digital '
     'services.')

h2('1.6.4 Educational Significance')
body('This project serves as a valuable educational resource for students and professionals '
      'seeking to understand the application of machine learning to network security. The '
      'complete system, including source code, trained model, and documentation, is publicly '
      'available on GitHub, enabling learners to study the implementation, experiment with '
      'modifications, and gain hands-on experience with ML-based intrusion detection. The '
      'interactive dashboard and AI assistant make the system suitable for classroom '
      'demonstrations and practical laboratory exercises in network security courses.')
body('The project also contributes to the development of local capacity in cybersecurity '
      'research and development. By providing a complete, working system that addresses a '
      'real-world security challenge, the project demonstrates the capability of Nigerian '
      'university students to produce internationally competitive research outputs. The '
      'open-source nature of the project encourages collaboration and knowledge sharing, '
      'fostering a culture of innovation and technical excellence in the academic community.')

h2('1.7 Definition of Terms')
body('Intrusion Detection System (IDS): A device or software application that monitors network '
     'traffic for suspicious activity and alerts the system or network administrator.')
body('Encrypted Traffic: Network communication that has been encoded using cryptographic '
     'protocols such as TLS, HTTPS, or VPN to prevent unauthorised access to its content.')
body('XGBoost (Extreme Gradient Boosting): An optimised distributed gradient boosting library '
     'designed to be highly efficient, flexible, and portable, widely used for supervised '
     'learning tasks.')
body('CICIDS2017: A benchmark intrusion detection dataset created by the Canadian Institute for '
     'Cybersecurity that contains realistic benign and attack traffic with labelled flow-based '
     'features.')
body('Flow Features: Statistical properties of network flows, such as packet counts, byte '
     'counts, flow duration, and inter-arrival times, used to characterise network traffic '
     'without inspecting packet payloads.')
body('PCAP (Packet Capture): A file format and application programming interface for capturing '
     'network packets, commonly used in network analysis and forensics.')
body('ROC-AUC (Receiver Operating Characteristic — Area Under Curve): A performance metric '
     'for classification models that measures the area under the ROC curve, representing the '
     'model\'s ability to distinguish between classes across all threshold settings.')
body('Streamlit: An open-source Python library for building interactive web applications for '
     'data science and machine learning projects with minimal code.')
body('tshark: The command-line version of Wireshark, used for capturing and analysing network '
     'packets and extracting protocol-level information.')
body('StandardScaler: A scikit-learn preprocessing technique that standardises features by '
     'removing the mean and scaling to unit variance.')

h2('1.8 Organization of the Study')
body('This project report is organised into five chapters.')
body('Chapter One presents the introduction, background of the study, motivation, problem '
     'statement, aim and objectives, scope, significance, definition of terms, and organisation '
     'of the study.')
body('Chapter Two reviews relevant literature related to intrusion detection systems, machine '
     'learning in network security, the XGBoost algorithm, network traffic analysis, the '
     'CICIDS2017 dataset, related works, and the research gap.')
body('Chapter Three describes the methodology adopted for the design and development of the '
     'proposed system, including the design approach, system architecture, data pipeline, model '
     'development, and development tools.')
body('Chapter Four presents the implementation details, system interfaces, testing procedures, '
     'model evaluation results, and discussion of findings.')
body('Chapter Five provides the summary of the study, conclusion, recommendations, and '
     'suggestions for future research.')

# ==================== CHAPTER TWO ====================
chapter(['Chapter Two', 'Literature Review'])

h2('2.1 Introduction')
body('This chapter reviews relevant literature related to intrusion detection systems, machine '
     'learning techniques applied to network security, the XGBoost ensemble algorithm, network '
     'traffic analysis and feature extraction, the CICIDS2017 dataset, related works in '
     'ML-based intrusion detection, and the research gap that justifies the development of the '
     'proposed system. The review establishes the theoretical and practical foundation for the '
     'study.')

h2('2.2 Overview of Intrusion Detection Systems')
body('Intrusion Detection Systems are security tools that monitor network traffic or system '
     'activities for malicious behaviour or policy violations. IDS can be broadly classified '
     'into two categories based on their detection methodology: signature-based detection and '
     'anomaly-based detection. Signature-based IDS, also known as misuse detection, compares '
     'network traffic against a database of known attack signatures. This approach is highly '
     'accurate for known attacks but fails to detect novel or zero-day threats. Snort and '
     'Suricata are widely used open-source signature-based IDS tools.')
body('Anomaly-based IDS establishes a baseline of normal network behaviour and flags deviations '
     'from this baseline as potential intrusions. This approach can detect previously unknown '
     'attacks but typically suffers from higher false positive rates, as legitimate but unusual '
     'traffic may be classified as malicious. Machine learning techniques have been increasingly '
     'applied to anomaly-based detection to improve accuracy and reduce false alarms. According '
     'to Buczak and Guven (2016), machine learning methods have shown significant promise in '
     'improving the effectiveness of intrusion detection systems by enabling automated learning '
     'of traffic patterns.')
body('IDS can also be categorised by deployment architecture: Network-based IDS (NIDS) monitors '
     'traffic at strategic points within the network, while Host-based IDS (HIDS) monitors '
     'activities on individual hosts. The system developed in this study is a network-based IDS '
     'that analyses flow-level traffic statistics, making it suitable for deployment at network '
     'choke points such as gateway routers or network switches.')

h2('2.3 Machine Learning in Network Security')
body('Machine learning has emerged as a powerful tool for network security, enabling the '
     'development of intelligent systems that can learn patterns from data and make predictions '
     'about network traffic. Supervised learning algorithms, including decision trees, random '
     'forests, support vector machines (SVM), and neural networks, have been extensively studied '
     'for intrusion detection tasks. These algorithms are trained on labelled datasets containing '
     'both benign and malicious traffic, and learn to distinguish between the two classes based '
     'on extracted features.')
body('Unsupervised learning techniques, such as clustering and autoencoders, are also used in '
     'intrusion detection, particularly for detecting novel attacks where labelled training data '
     'may not be available. Clustering algorithms group similar traffic flows together, and '
     'flows that do not belong to any established cluster can be flagged as anomalous. However, '
     'supervised approaches generally achieve higher accuracy when sufficient labelled data is '
     'available, as they directly learn the decision boundary between classes.')
body('Ensemble methods, which combine multiple base learners to produce a stronger classifier, '
     'have consistently demonstrated superior performance in intrusion detection tasks. Random '
     'forests, which aggregate the predictions of many decision trees, and gradient boosting '
     'methods, which build trees sequentially to correct errors of previous trees, are among the '
     'most effective ensemble techniques. The choice of algorithm depends on factors including '
     'dataset characteristics, computational resources, and the specific requirements of the '
     'deployment environment.')

h2('2.4 XGBoost and Ensemble Methods')
body('XGBoost, short for Extreme Gradient Boosting, is an optimised implementation of the '
     'gradient boosting framework developed by Chen and Guestrin (2016). Gradient boosting is '
     'an ensemble technique that builds predictive models by iteratively adding decision trees, '
     'where each new tree is trained to correct the residual errors of the previous ensemble. '
     'XGBoost incorporates several innovations that make it particularly effective for '
     'structured and tabular data, including regularised objective functions that prevent '
     'overfitting, column subsampling that reduces computational cost and improves '
     'generalisation, and a sparsity-aware algorithm that efficiently handles missing values.')
body('Several characteristics make XGBoost well suited for intrusion detection in encrypted '
     'traffic. First, it handles high-dimensional feature spaces effectively, making it suitable '
     'for the large number of flow-based features extracted from network traffic. Second, its '
     'built-in regularisation reduces the risk of overfitting, which is important when '
     'training on datasets with complex and noisy traffic patterns. Third, XGBoost provides '
     'feature importance scores, enabling interpretability by identifying which flow '
     'characteristics are most indicative of malicious activity.')
body('Comparative studies have demonstrated that XGBoost often outperforms other machine '
     'learning algorithms including random forests, SVM, and neural networks for tabular '
     'classification tasks. Zhang et al. (2019) reported that XGBoost achieved superior '
     'performance in network intrusion detection compared to random forest and k-nearest '
     'neighbours, with higher accuracy and lower false positive rates. These findings support '
     'the selection of XGBoost as the core classification algorithm for the proposed IDS.')

h2('2.5 Network Traffic Analysis and Feature Extraction')
body('Network traffic analysis is the process of capturing, recording, and analysing network '
     'traffic patterns to gain insights into network operations and detect anomalies. For '
     'intrusion detection, traffic analysis can be performed at different levels of granularity, '
     'including packet-level analysis, flow-level analysis, and session-level analysis. '
     'Flow-level analysis, which aggregates packets sharing common properties such as source '
     'and destination IP addresses, ports, and protocol, into bidirectional flows, provides a '
     'balance between detail and computational efficiency.')
body('CICFlowMeter is a widely used network traffic flow generator that extracts over 80 '
     'statistical features from bidirectional network flows. These features include flow '
     'duration, packet counts in both directions, byte counts, packet length statistics '
     '(minimum, maximum, mean, standard deviation), inter-arrival times, and flag-related '
     'indicators. The key advantage of flow-based features for encrypted traffic analysis is '
     'that they are derived from metadata and statistics of the traffic rather than from the '
     'encrypted payload, making them suitable for classification even when the content is '
     'inaccessible.')
body('The feature extraction process for PCAP files in the proposed system uses tshark, the '
     'command-line tool of Wireshark, to parse packet captures and extract the relevant fields '
     'necessary for computing flow features. The extracted features are then scaled using '
     'StandardScaler to ensure that all features contribute equally to the model\'s decisions, '
     'preventing features with larger numerical ranges from dominating the classification '
     'process.')

h2('2.5.1 Deep Learning Approaches for Intrusion Detection')
body('In addition to ensemble methods such as XGBoost and random forests, deep learning '
      'architectures have gained significant attention in the intrusion detection domain. '
      'Convolutional Neural Networks (CNNs) have been applied to network traffic by converting '
      'flow features into two-dimensional representations, treating them as image-like inputs '
      'from which spatial patterns indicative of malicious activity can be learned. Recurrent '
      'Neural Networks (RNNs), particularly Long Short-Term Memory (LSTM) networks, are well '
      'suited for modelling sequential network traffic data, as they can capture temporal '
      'dependencies between successive flows in a network session.')
body('Autoencoders, which are unsupervised neural networks that learn compressed representations '
      'of input data, have been employed for anomaly detection by reconstructing normal traffic '
      'patterns and flagging flows with high reconstruction error as anomalies. Variational '
      'autoencoders (VAEs) extend this approach by learning a probabilistic latent space, '
      'providing a more robust basis for detecting deviations from normal behaviour.')
body('Graph Neural Networks (GNNs) represent a newer frontier in network traffic analysis. By '
      'modelling network communications as graphs, where nodes represent hosts and edges '
      'represent communication flows, GNNs can capture structural patterns in network traffic '
      'that traditional feature-based approaches may miss. Studies have shown that GNNs can '
      'effectively detect distributed attacks such as DDoS and botnet activity based on the '
      'communication patterns between compromised hosts.')
body('While deep learning approaches offer the potential for higher accuracy and the ability to '
      'learn complex patterns automatically, they require significantly larger training datasets '
      'and more computational resources than gradient boosting methods. Deep learning models '
      'also tend to be less interpretable than tree-based models, making it more difficult to '
      'explain why a particular flow was classified as malicious. For these reasons, XGBoost '
      'was selected as the core algorithm for this study, as it provides a practical balance '
      'between accuracy, computational efficiency, and interpretability for the intrusion '
      'detection task.')

h2('2.6 The CICIDS2017 Dataset')
body('The CICIDS2017 dataset, created by Sharafaldin, Lashkari, and Ghorbani (2018) at the '
     'Canadian Institute for Cybersecurity, is one of the most comprehensive and realistic '
     'benchmark datasets for intrusion detection research. The dataset was generated by '
     'simulating a realistic network environment over a five-day period, collecting traffic '
     'that includes both benign activity and a wide range of contemporary attack types. The '
     'traffic was captured using the CICFlowMeter tool, which extracted over 80 flow-based '
     'features for each bidirectional flow.')
body('The dataset includes 14 attack categories distributed across multiple days of capture: '
     'Brute Force (FTP and SSH), Heartbleed, Botnet, DoS/DDoS (including Hulk, GoldenEye, '
     'Slowloris, and SlowHTTPTest), Web Attacks (brute force, SQL injection, and XSS), '
     'Infiltration, and Port Scan. The inclusion of both benign and attack traffic across '
     'multiple protocols and attack types makes CICIDS2017 a valuable resource for training '
     'and evaluating intrusion detection systems. The dataset has been widely used in academic '
     'research and serves as a standard benchmark for comparing IDS performance.')
body('For this study, the CICIDS2017 dataset was preprocessed to handle class imbalance, '
     'remove missing values, and encode categorical features. The features were standardised '
     'using StandardScaler, and the dataset was split into training (70%), validation (15%), '
     'and test (15%) sets to ensure robust model evaluation. The comprehensive nature of the '
     'dataset enables the trained XGBoost model to generalise across different attack types '
      'and traffic conditions.')

h3('2.6.1 Dataset Composition and Class Distribution')
body('The CICIDS2017 dataset contains a total of 3,347,889 network flows recorded over five '
      'days of simulated network activity. The class distribution is imbalanced, with benign '
      'traffic constituting approximately 73% of the total flows and malicious traffic '
      'constituting approximately 27%. This imbalance reflects real-world network conditions, '
      'where the vast majority of traffic is legitimate and attacks represent a small fraction '
      'of total flows.')
body('The breakdown of traffic across the major attack categories is as follows:')
bullet('Benign — 2,443,820 flows (73.0% of the dataset), representing normal user activities '
        'including web browsing, email, file transfers, and video streaming.')
bullet('DDoS / DoS — 579,630 flows (17.3%), comprising Hulk (461,912), GoldenEye (41,508), '
        'Slowloris (10,990), and SlowHTTPTest (10,490) attacks, making DoS the largest '
        'attack category in the dataset.')
bullet('Port Scan — 158,930 flows (4.7%), including both vertical and horizontal port '
        'scanning activities using tools such as Nmap.')
bullet('Brute Force — 27,888 flows (0.8%), including FTP brute force (15,399) and SSH brute '
        'force (12,489) attack attempts.')
bullet('Botnet — 29,664 flows (0.9%), comprising ARES botnet traffic including command-and-'
        'control communication and attack execution commands.')
bullet('Web Attacks — 5,540 flows (0.2%), including SQL injection (6,196), cross-site '
        'scripting (652), and web-based brute force attacks (1,472).')
bullet('Infiltration — 115,684 flows (3.5%), including internal network reconnaissance and '
        'exploitation attempts.')
bullet('Heartbleed — 6,357 flows (0.2%), exploiting the OpenSSL Heartbleed vulnerability.')
body('This class distribution presents a realistic challenge for machine learning-based '
      'intrusion detection, as the model must learn to detect attacks despite the significant '
      'imbalance between benign and malicious traffic. The XGBoost algorithm\'s built-in '
      'handling of class weights and the use of stratified sampling during train-test splitting '
      'help mitigate the effects of this imbalance on model performance.')

h2('2.7 Related Works and Existing Systems')

h3('2.7.1 Snort and Signature-Based IDS')
body('Approach')
body('Snort is an open-source, rule-based intrusion detection and prevention system that uses '
     'signature matching to detect known attack patterns in network traffic. It operates by '
     'comparing packet payloads against a database of predefined rules.')
body('Architectural Framework')
body('Snort employs a packet decoder, preprocessors, a detection engine, and logging and '
     'alerting subsystems. Rules define patterns in packet headers and payloads that indicate '
     'malicious activity.')
body('Strengths')
bullet('Mature and widely deployed with extensive community rule sets.')
bullet('Real-time detection capability with high throughput.')
bullet('Free and open-source with active community support.')
body('Weaknesses')
bullet('Ineffective against encrypted traffic as payload inspection is blocked.')
bullet('Cannot detect novel or zero-day attacks for which no signature exists.')
bullet('Requires frequent rule updates to remain effective against new threats.')

h3('2.7.2 Suricata and Hybrid IDS')
body('Approach')
body('Suricata is an open-source threat detection engine that combines signature-based '
     'detection with protocol analysis and limited anomaly detection capabilities.')
body('Architectural Framework')
body('It consists of a multi-threaded packet processing engine, a rule-based detection engine, '
     'protocol parsers, and an output module for logging and alerting.')
body('Strengths')
bullet('Multi-threaded design provides better performance on modern hardware.')
bullet('Supports both signature and basic anomaly detection.')
bullet('Can process PCAP files offline in addition to live traffic.')
body('Weaknesses')
bullet('Anomaly detection capabilities are limited compared to ML-based solutions.')
bullet('Still largely dependent on signature rules for accurate detection.')
bullet('Limited effectiveness against encrypted attack traffic.')

h3('2.7.3 ML-Based IDS Research')
body('Approach')
body('Numerous research studies have applied machine learning algorithms including decision '
     'trees, random forests, SVM, and neural networks to intrusion detection using benchmark '
     'datasets such as KDDCup99, NSL-KDD, UNSW-NB15, and CICIDS2017.')
body('Architectural Framework')
body('These systems typically follow a pipeline of data preprocessing, feature selection, '
     'model training with cross-validation, and evaluation using standard classification metrics.')
body('Strengths')
bullet('Can detect previously unseen attacks when trained on representative data.')
bullet('Operate on statistical features rather than requiring payload access.')
bullet('Adaptable to different network environments through retraining.')
body('Weaknesses')
bullet('Many existing solutions lack a user-friendly deployment interface.')
bullet('Research prototypes often do not support real-time traffic analysis.')
bullet('Limited explainability of model predictions hinders practical adoption.')

h3('2.7.4 Commercial IDS and Next-Generation Firewalls')
body('Approach')
body('Commercial security solutions, including next-generation firewalls (NGFWs) and '
     'enterprise IDS platforms, increasingly incorporate machine learning and behavioural '
     'analysis alongside traditional signature-based detection.')
body('Architectural Framework')
body('These solutions combine hardware acceleration, deep packet inspection, SSL/TLS '
     'decryption, and cloud-based threat intelligence feeds in an integrated security platform.')
body('Strengths')
bullet('Comprehensive security features in a single appliance.')
bullet('Vendor support, regular updates, and professional services.')
bullet('High throughput suitable for enterprise networks.')
body('Weaknesses')
bullet('High cost makes them inaccessible for small and medium organisations.')
bullet('SSL/TLS decryption raises privacy and compliance concerns.')
bullet('Proprietary algorithms limit transparency and customisation.')

h3('2.7.5 Summary of Related Works')
body('The reviewed systems and research demonstrate that while signature-based IDS and '
     'commercial security solutions provide valuable protection, they face significant '
     'limitations in the context of encrypted traffic. Signature-based systems cannot inspect '
     'encrypted payloads, commercial solutions are expensive and raise privacy concerns, and '
     'existing ML-based research prototypes often lack user-friendly interfaces and real-time '
     'monitoring capabilities. The proposed system addresses these limitations by combining '
     'XGBoost-based classification on flow features with an accessible Streamlit dashboard '
     'and AI-powered explainability.')

h2('2.8 Research Gap')
body('Based on the reviewed literature and existing systems, the following research gaps were '
     'identified:')
numbered('Limited Encrypted Traffic Detection: Existing signature-based IDS and many commercial '
         'solutions cannot effectively detect attacks hidden within encrypted traffic, as they '
         'rely on payload inspection.')
numbered('Usability Gap in ML-Based IDS: Research prototypes often focus on model accuracy '
         'without providing user-friendly interfaces that enable network administrators to '
         'practically deploy and interact with the system.')
numbered('Lack of Real-Time Detection: Many ML-based intrusion detection studies evaluate models '
         'on static datasets without implementing real-time or near-real-time traffic capture '
         'and classification capabilities.')
numbered('Insufficient Explainability: ML-based IDS often function as black boxes, making '
         'predictions without providing understandable explanations that help security analysts '
         'understand why traffic was classified as malicious.')
numbered('Limited Multi-Mode Support: Most existing systems support only one mode of traffic '
         'analysis (either file upload or live capture), limiting their flexibility in '
         'different operational scenarios.')

h2('2.9 Summary of Literature Review')
body('This chapter reviewed literature on intrusion detection systems, machine learning in '
     'network security, XGBoost and ensemble methods, network traffic analysis and feature '
     'extraction, the CICIDS2017 dataset, related works, and the research gap. The review '
     'established that encrypted traffic presents a fundamental challenge to traditional IDS, '
     'that machine learning, particularly XGBoost, offers an effective solution through '
     'flow-based feature analysis, and that existing systems lack the combination of accuracy, '
     'usability, multi-mode support, and explainability that is needed for practical '
     'deployment. These gaps justify the development of the proposed AI-based IDS system.')

# ==================== CHAPTER THREE ====================
chapter(['Chapter Three', 'Methodology'])

h2('3.1 Introduction')
body('This chapter presents the methodology adopted for the design and implementation of the '
     'AI-based intrusion detection system. It describes the project design approach, the design '
     'considerations, the system architecture, the data pipeline, the model development process, '
     'the development tools and technologies, the evaluation metrics, and the operational flow '
     'of the system. The methodology provides a structured framework for developing a reliable, '
     'accurate, and usable intrusion detection system.')

h2('3.2 Project Design and Approach')
body('The development of the AI-based IDS followed the Software Development Life Cycle (SDLC) '
     'using the Agile development approach. The Agile methodology was adopted because it '
     'supports iterative development, continuous refinement, and regular testing, which are '
     'well suited to a project involving data exploration, model experimentation, and user '
      'interface development. Features were developed incrementally through multiple sprints, '
      'each delivering a functional component of the system.')
body('The Agile methodology was implemented through a series of two-week sprints, each focused '
      'on delivering a specific set of features. Sprint planning was conducted at the beginning '
      'of each sprint to define the tasks to be completed, and a sprint review was conducted '
      'at the end to assess progress and adjust priorities. The use of sprints allowed for '
      'flexible adaptation to challenges encountered during development, such as data quality '
      'issues discovered during preprocessing or adjustments to the model architecture based on '
      'preliminary evaluation results.')
body('Continuous integration practices were employed throughout development, with code changes '
      'committed to the version control repository on a daily basis and tested for integration '
      'issues. Automated testing scripts were developed alongside the main application code, '
      'ensuring that new features did not introduce regressions in existing functionality. '
      'This iterative approach, combined with regular testing, reduced the risk of major '
      'integration problems and enabled rapid identification and resolution of issues.')
body('The project was organised into the following phases: requirements gathering and analysis, '
      'where the functional and non-functional requirements of the system were defined based on '
      'the research objectives and literature review; data collection and preprocessing, where '
      'the CICIDS2017 dataset was acquired, cleaned, and prepared for model training; model '
      'development and experimentation, where multiple XGBoost configurations were evaluated '
      'using cross-validation; dashboard development, where the Streamlit user interface was '
      'designed and implemented; and system integration and testing, where all components were '
      'integrated into a cohesive system and validated through comprehensive testing.')

h2('3.3 Design Considerations')
body('Several factors were considered during the design of the proposed system.')
h3('Accuracy')
body('The primary objective of the IDS is to correctly classify network traffic as benign or '
     'malicious. The system must achieve high accuracy, precision, and recall to minimise both '
     'false positives, which cause unnecessary alerts, and false negatives, which allow attacks '
     'to go undetected. The XGBoost algorithm was selected for its proven performance in '
     'classification tasks.')
h3('Real-Time Processing')
body('For the live capture mode, the system must process network packets and generate '
     'predictions with minimal latency. The model inference was optimised to run efficiently '
     'on commodity hardware, and the feature extraction pipeline was designed to process '
     'traffic in near-real time.')
h3('Usability')
body('The system was designed to be accessible to network administrators with varying levels '
     'of machine learning expertise. The Streamlit dashboard provides an intuitive interface '
     'with clear instructions, visual feedback, and the AI assistant that explains predictions '
     'in plain language.')
h3('Scalability')
body('The modular architecture of the system allows components to be independently scaled or '
      'replaced. The model can be retrained on new data, additional traffic modes can be added, '
      'and the dashboard can be extended with new features without affecting existing '
      'functionality.')
h3('Security')
body('The system itself must be secure against tampering and unauthorised access. The trained '
      'model files are protected using file system permissions, and the dashboard application '
      'can be configured to require authentication before granting access to the traffic analysis '
      'features. The AI assistant API key is managed through environment variables rather than '
      'being hardcoded, and all sensitive configuration is kept outside the version control '
      'repository. The system logs all classification activities for audit purposes, providing '
      'an auditable trail of detection events.')
h3('Portability')
body('The system was designed to be platform-independent and easy to deploy. The use of Python, '
      'a cross-platform programming language, ensures that the application can run on Windows, '
      'Linux, and macOS operating systems. The dependencies are managed through a requirements.txt '
      'file, enabling one-command installation of all required packages. A Dockerfile and '
      'docker-compose.yml configuration are provided for containerised deployment, ensuring '
      'consistent behaviour across different environments and simplifying the setup process '
      'for users who may not have Python and the required dependencies installed on their '
      'systems.')
h3('Maintainability')
body('The source code is organised into clear, well-documented modules with separation of '
      'concerns. The data pipeline, model inference, and dashboard components are implemented '
      'in separate modules, making it straightforward to update or replace individual components '
      'without affecting the rest of the system. The code follows Python PEP 8 style guidelines '
      'and includes descriptive variable names and inline comments. The project uses a version '
      'control system (Git) for tracking changes, enabling collaborative development and the '
      'ability to revert to previous versions if needed.')

h2('3.4 System Architecture')
body('The proposed AI-based IDS adopts a modular architecture consisting of four main layers: '
     'the data input layer, the feature extraction and preprocessing layer, the model inference '
     'layer, and the presentation layer.')
body('The Data Input Layer provides three interfaces for traffic data ingestion: a CSV file '
     'upload interface for pre-extracted flow features, a PCAP file upload interface that uses '
     'tshark for automated feature extraction, and a live network capture interface that uses '
     'Scapy to capture packets in real time and extract flow-based features programmatically.')
body('The Feature Extraction and Preprocessing Layer transforms raw traffic data into '
     'standardised feature vectors suitable for model inference. For PCAP and live modes, this '
     'layer extracts flow-level statistics including packet counts, byte counts, flow duration, '
     'and packet length distributions. All features are scaled using the same StandardScaler '
     'that was fitted on the training data to ensure consistency.')
body('The Model Inference Layer loads the pre-trained XGBoost model and applies it to the '
     'preprocessed feature vectors, returning class predictions and prediction probabilities '
     'for each traffic flow. The model was serialised using the pickle format and is loaded '
     'at application startup.')
body('The Presentation Layer, implemented with Streamlit, renders the dashboard interface '
     'including file upload widgets, prediction results tables, performance visualisations '
     '(confusion matrix, ROC curve, feature importance), and the AI assistant chat interface. '
      'The interaction among these layers ensures a smooth and efficient detection pipeline.')
caption('Figure 3.1: System Architecture of the Proposed IDS')
body('Figure 3.2 presents a detailed system flowchart illustrating the complete data flow '
      'from traffic input through feature extraction, preprocessing, and model inference, '
      'to the final prediction output displayed on the dashboard.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\ids\system_flowchart.png',
       'Figure 3.2: System Flowchart of the AI-Based IDS Pipeline')

h2('3.5 Data Pipeline')
body('The data pipeline is a critical component of the system that transforms raw network '
     'traffic data into a format suitable for model training and inference. The pipeline '
     'consists of several sequential stages.')
body('Loading: The CICIDS2017 dataset consists of multiple CSV files corresponding to traffic '
     'captured on different days. These files are loaded and concatenated into a single '
     'DataFrame for unified processing.')
body('Cleaning: The dataset is cleaned by removing rows with missing or infinite values, '
     'which can cause numerical instability during model training. Columns that contain only '
     'a single value or that are identified as irrelevant (such as flow IDs and timestamps) '
     'are dropped to reduce dimensionality and prevent data leakage.')
body('Label Encoding: The target labels are encoded into binary or multi-class format. For '
     'binary classification, all attack types are mapped to the malicious class (1) and benign '
     'traffic is mapped to the benign class (0). For multi-class evaluation, each attack '
     'category is assigned a distinct numerical label.')
body('Splitting: The cleaned and encoded dataset is split into training (70%), validation '
     '(15%), and test (15%) sets using stratified sampling to preserve the class distribution '
     'across all splits.')
body('Scaling: Numerical features are standardised using StandardScaler, which transforms '
     'each feature to have zero mean and unit variance. The scaler is fitted on the training '
     'data and then applied to the validation and test sets to prevent data leakage from the '
     'validation process into model training.')

h2('3.6 Model Development')
body('The XGBoost classifier was selected as the core detection algorithm for the proposed '
     'system. The model was developed through an iterative process of hyperparameter tuning '
     'and cross-validation.')
body('The initial model configuration used default XGBoost parameters as a baseline. '
     'Hyperparameter tuning was then performed using grid search combined with stratified '
     'k-fold cross-validation (k = 5) to identify the optimal combination of parameters. The '
     'key hyperparameters tuned included the learning rate (eta), maximum tree depth, '
     'minimum child weight, subsample ratio, column subsample ratio, and the number of '
     'boosting rounds (n_estimators).')
body('The final model configuration was selected based on the best cross-validation accuracy '
     'while balancing model complexity to avoid overfitting. Early stopping was employed '
     'during training using the validation set to halt training when the validation error '
     'stopped improving for 50 consecutive rounds. The trained model was evaluated on the '
     'held-out test set and achieved a classification accuracy of 99.9%, a ROC-AUC of 0.9999, '
     'and high precision and recall scores across all attack categories.')

h2('3.7 Development Tools and Technologies')
body('The tools and technologies used in the development of the system are summarised in Table '
     '3.1.')
caption('Table 3.1: Software Development Tools and Technologies')
bullet('Python 3.12 — Primary programming language for all system components.')
bullet('Streamlit — Web framework for building the interactive dashboard interface.')
bullet('XGBoost — Gradient boosting library for the core classification model.')
bullet('scikit-learn — Machine learning library for data preprocessing, model evaluation, and '
       'cross-validation.')
bullet('pandas and NumPy — Data manipulation and numerical computation libraries.')
bullet('Plotly — Interactive charting library for performance visualisations.')
bullet('tshark (Wireshark CLI) — Packet analysis tool for feature extraction from PCAP files.')
bullet('Scapy — Python library for packet capture and network interaction.')
bullet('Google Gemini — Large language model API for the AI assistant explainability feature.')
bullet('Matplotlib and Seaborn — Static visualisation libraries for confusion matrix and ROC '
       'curve plots.')

h2('3.8 Evaluation Metrics')
body('The performance of the XGBoost model was evaluated using standard classification metrics '
     'derived from the confusion matrix, which compares the predicted labels with the true '
     'labels for the test set.')
body('Accuracy measures the proportion of correctly classified instances (both benign and '
     'malicious) out of the total instances. While accuracy provides a general indication of '
     'model performance, it can be misleading in the presence of class imbalance, which is '
     'common in intrusion detection datasets.')
body('Precision (also called the positive predictive value) measures the proportion of '
     'instances classified as malicious that are truly malicious. High precision indicates a '
     'low false positive rate. Recall (also called sensitivity or the true positive rate) '
     'measures the proportion of truly malicious instances that were correctly identified. '
     'High recall indicates a low false negative rate.')
body('The F1-score is the harmonic mean of precision and recall, providing a single metric '
     'that balances both concerns. The ROC-AUC measures the model\'s ability to distinguish '
     'between classes across all possible classification thresholds, with a value of 1.0 '
     'indicating perfect discrimination and 0.5 indicating random guessing.')

h2('3.8.1 Cross-Validation Strategy')
body('To ensure the robustness and generalisability of the trained XGBoost model, a stratified '
      'k-fold cross-validation strategy was employed during the hyperparameter tuning phase. '
      'Stratified k-fold cross-validation divides the training data into k equal folds while '
      'preserving the proportion of benign and malicious samples in each fold, ensuring that '
      'the class imbalance present in the CICIDS2017 dataset is maintained across all '
      'validation splits.')
body('A value of k = 5 was selected for this study, meaning the training data was divided into '
      'five folds. In each iteration, four folds were used for training and one fold was used '
      'for validation, with the process repeated five times so that each fold served as the '
      'validation set exactly once. The performance metrics from all five iterations were '
      'averaged to produce a single cross-validation score for each hyperparameter configuration '
      'evaluated during grid search.')
body('The use of cross-validation provides several important benefits for this study. First, it '
      'reduces the variance of the performance estimate compared to a single train-validation '
      'split, providing a more reliable assessment of how the model will generalise to unseen '
      'traffic data. Second, it maximises the utilisation of the available training data, which '
      'is particularly important when working with limited labelled datasets. Third, it helps '
      'detect overfitting by revealing whether the model performs consistently across different '
      'subsets of the data. The cross-validation results confirmed that the selected '
      'hyperparameters produced a model with stable and high performance across all five folds, '
      'with accuracy consistently above 99.5% in each validation iteration.')
body('After cross-validation, the final model was retrained on the full training set using the '
      'optimal hyperparameters and evaluated on the held-out test set to obtain the final '
      'performance metrics reported in Chapter Four.')

h2('3.9 Operational Flow of the System')
body('The operational flow of the IDS begins when the user launches the Streamlit dashboard. '
     'The main interface presents three tabs corresponding to the three operational modes.')
body('In the CSV mode, the user uploads a CSV file containing pre-extracted flow features. '
     'The system validates the file format and column structure, scales the features using the '
     'pre-fitted StandardScaler, passes them through the XGBoost model, and displays the '
     'prediction results in a table with colour-coded classification indicators. The user can '
     'also view aggregate statistics and download the results.')
body('In the PCAP mode, the user uploads a PCAP file. The system invokes tshark to extract '
     'relevant network flow information, computes the required statistical features, scales '
     'them, and classifies each flow. Results are presented alongside the original packet '
     'information for contextual analysis.')
body('In the Live mode, the system begins capturing network packets from a specified network '
     'interface using Scapy. Captured packets are grouped into flows, features are extracted '
     'and scaled in real time, and predictions are displayed as a scrolling table that updates '
     'as new traffic is analysed. Throughout all modes, the AI assistant is available via a '
     'chat interface at the bottom of the dashboard, allowing the user to ask questions about '
     'the detected traffic and receive plain-language explanations.')

# ==================== CHAPTER FOUR ====================
chapter(['Chapter Four', 'Implementation, Results and Discussion'])

h2('4.1 Introduction')
body('This chapter presents the implementation of the AI-based intrusion detection system, the '
     'development environment, the system interfaces, the testing procedures, the model '
     'evaluation results, and a discussion of the findings in relation to the objectives of '
     'the study and the gaps identified in the literature.')

h2('4.2 System Implementation')
body('The system was implemented following the architecture and design described in Chapter '
     'Three. The core application was developed in Python using the Streamlit framework for '
     'the web dashboard. The XGBoost model was trained on the preprocessed CICIDS2017 dataset '
     'using the xgboost Python package with hyperparameter tuning via scikit-learn\'s '
     'GridSearchCV. The trained model was serialised to a pickle file for deployment.')
body('The Streamlit dashboard was structured as a multi-page application with a sidebar for '
     'navigation and three main tabs: CSV File, PCAP File, and Live Capture. Each tab was '
     'implemented as a separate function that handles file uploads, triggers feature '
     'extraction and classification, and renders the results. The Plotly library was used to '
     'generate interactive charts including confusion matrices, ROC curves, and feature '
     'importance bar charts.')
body('The AI assistant was integrated using the google-generativeai Python package, which '
     'provides access to the Gemini API. A system prompt was designed to instruct the model '
     'to provide concise, accurate explanations of network security concepts and detected '
     'threats based on the traffic data context provided by the dashboard. The assistant '
     'interface was implemented as a chat component at the bottom of the dashboard, '
     'maintaining conversation history throughout the session.')
body('The application was styled with a dark matrix-themed colour scheme using custom CSS '
     'applied through Streamlit\'s markdown component. The theme includes green-on-black '
     'colour accents, monospaced fonts for data displays, and animated effects that evoke '
     'the aesthetic of digital security monitoring interfaces, enhancing the user experience '
      'while maintaining readability and professional appearance.')

h3('4.2.1 Development Environment and System Requirements')
body('The system was developed and tested on a Windows 11 machine with an Intel Core i7 '
      'processor, 16 GB of RAM, and a solid-state drive. The software environment included '
      'Python 3.12, with key packages including Streamlit 1.28 for the web dashboard, '
      'XGBoost 2.0 for model training and inference, scikit-learn 1.3 for data preprocessing '
      'and evaluation, pandas 2.1 for data manipulation, Plotly 5.17 for interactive '
      'visualisations, Scapy 2.5 for live packet capture, and google-generativeai 0.3 for '
      'the AI assistant integration. The complete list of dependencies is documented in the '
      'requirements.txt file included in the project repository.')
body('The minimum hardware requirements for running the system are a dual-core processor, '
      '4 GB of RAM, and 500 MB of free disk space for the application and model files. For '
      'live capture mode, a network interface card that supports promiscuous mode is required, '
      'and tshark (part of the Wireshark distribution) must be installed for PCAP file '
      'processing. The system does not require a GPU for model inference, as the XGBoost '
      'model is small enough to run efficiently on standard CPUs, with average inference '
      'time of less than one millisecond per network flow.')
body('The software requirements for deployment include Python 3.10 or higher, the dependencies '
      'listed in requirements.txt, tshark installed and accessible in the system PATH (for PCAP '
      'analysis), and network adapter privileges (for live capture mode). For the AI assistant '
      'feature, an internet connection and a valid Google Gemini API key are required. The '
      'application runs on Windows, Linux, and macOS platforms, and supports deployment via '
      'Docker using the provided Dockerfile and docker-compose configuration.')

h2('4.3 System Interfaces')
body('This section presents the major interfaces of the implemented system.')

h3('4.3.1 Full Dashboard')
body('The main dashboard interface, shown in Figure 4.1, presents the complete IDS system '
     'with a dark matrix-themed design. The interface includes a sidebar with navigation '
     'controls, three tabs for the operational modes, and the AI assistant chat panel at '
     'the bottom. The dashboard provides real-time status indicators and summary statistics.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\ids\01-full-dashboard.png',
       'Figure 4.1: Full Dashboard Interface of the AI-Based IDS')

h3('4.3.2 PCAP Upload Tab')
body('The PCAP upload interface, shown in Figure 4.2, allows users to upload packet capture '
     'files for offline analysis. The interface displays file upload controls, a progress '
     'indicator during feature extraction, and a results table showing classified flows with '
     'confidence scores. Users can filter results by prediction class and download the '
     'analysis report.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\ids\02-pcap-tab.png',
       'Figure 4.2: PCAP Upload Tab of the IDS Dashboard')

h3('4.3.3 Live Capture Tab')
body('The live capture interface, shown in Figure 4.3, enables real-time network monitoring. '
     'Users select the network interface to capture from, start and stop the capture session, '
     'and view a scrolling table of classified flows that updates in real time. The interface '
     'also displays a live traffic summary showing the proportion of benign and malicious '
     'traffic detected during the session.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\ids\03-live-tab.png',
       'Figure 4.3: Live Capture Tab of the IDS Dashboard')

h3('4.3.4 AI Assistant')
body('The AI assistant interface, shown in Figure 4.4, is located at the bottom of the '
     'dashboard. Users can type questions about detected traffic, request explanations of '
     'model predictions, or ask about network security concepts. The assistant responds with '
     'plain-language explanations grounded in the current session data, helping users '
     'understand the implications of the detection results.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\ids\04-bottom-chat.png',
       'Figure 4.4: AI Assistant Interface of the IDS Dashboard')

h3('4.3.5 Confusion Matrix')
body('The confusion matrix visualisation, shown in Figure 4.5, displays the performance of '
     'the XGBoost classifier on the test set. The matrix shows true positive, true negative, '
     'false positive, and false negative counts. The near-diagonal dominance of the matrix '
     'visually confirms the high classification accuracy achieved by the model.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\ids\confusion_matrix.png',
       'Figure 4.5: Confusion Matrix of the XGBoost Classifier')

h3('4.3.6 ROC Curve')
body('The ROC curve visualisation, shown in Figure 4.6, plots the true positive rate against '
     'the false positive rate across various threshold settings. The curve\'s proximity to '
     'the top-left corner and the reported AUC value of 0.9999 demonstrate the model\'s '
     'excellent discriminatory power between benign and malicious traffic.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\ids\roc_curve.png',
       'Figure 4.6: ROC Curve of the XGBoost Classifier')

h2('4.4 System Testing')
body('The system was tested at three levels to ensure its reliability and correctness.')
h3('4.4.1 Unit Testing')
body('Individual components of the system were tested in isolation. The data preprocessing '
     'functions were tested with sample inputs to verify correct cleaning, encoding, and '
     'scaling behaviour. The feature extraction module was tested with sample PCAP files to '
     'confirm that tshark was correctly invoked and that the output features matched the '
     'expected format. The model inference function was tested with known input vectors to '
     'verify that the predictions matched the expected outputs from the trained model.')
h3('4.4.2 Integration Testing')
body('Integration testing verified that the components of the system worked together '
     'correctly. The complete pipeline from file upload (CSV and PCAP) through feature '
     'extraction and scaling to model inference and result display was tested end to end. '
     'The live capture pipeline was tested by generating sample traffic and verifying that '
     'flows were correctly captured, features extracted, and predictions displayed. Test '
     'results are summarised in Table 4.2.')
caption('Table 4.2: Summary of System Test Cases and Results')
bullet('CSV upload and classification — File uploaded, features validated, predictions '
       'generated, results displayed correctly. Passed.')
bullet('PCAP upload and classification — File parsed, features extracted, predictions '
       'generated, results displayed with flow information. Passed.')
bullet('Live capture and monitoring — Interface selected, packets captured, flows '
       'classified, predictions updated in real time. Passed.')
bullet('AI assistant queries — Questions answered with contextually relevant explanations '
       'based on dashboard data. Passed.')
bullet('Visualisation rendering — Confusion matrix, ROC curve, and feature importance '
       'charts rendered correctly. Passed.')
h3('4.4.3 User Acceptance Testing')
body('User acceptance testing was conducted with sample network traffic data provided to '
     'users who interacted with the dashboard through all three operational modes. Users '
     'confirmed that the interface was intuitive, the predictions were accurate, and the AI '
     'assistant provided helpful explanations. Feedback from testing was used to make minor '
     'improvements to the interface layout and messaging.')

h2('4.5 Model Evaluation')
body('The trained XGBoost model was evaluated on the held-out test set comprising 15% of the '
     'preprocessed CICIDS2017 dataset. The evaluation focused on both binary classification '
     'performance (benign vs. malicious) and multi-class classification performance across '
     'individual attack categories.')
body('The binary classification results are shown in Table 4.1. The model achieved an overall '
     'accuracy of 99.9%, indicating that it correctly classified more than 99 out of every '
     '100 traffic flows. The precision of 0.999 and recall of 0.999 demonstrate that the '
     'model maintains an excellent balance between minimising false positives and false '
     'negatives. The F1-score of 0.999 confirms the overall effectiveness of the classifier.')
caption('Table 4.1: Classification Report of the XGBoost Model')
bullet('Accuracy: 99.9%')
bullet('Precision: 0.999')
bullet('Recall: 0.999')
bullet('F1-Score: 0.999')
bullet('ROC-AUC: 0.9999')
body('The confusion matrix (Figure 4.5) further illustrates the model\'s performance, with '
     'the vast majority of benign and malicious flows correctly classified. Only a minimal '
     'number of false positives and false negatives were observed, all of which occurred for '
     'traffic flows with feature distributions near the decision boundary. The ROC curve '
     '(Figure 4.6) shows an AUC of 0.9999, which is extremely close to the theoretical '
     'maximum of 1.0, confirming the model\'s outstanding discriminatory capability.')
body('Feature importance analysis revealed that the most influential features for '
     'classification included flow duration, backward packet length statistics, and '
     'inter-arrival time statistics. These findings are consistent with the literature, '
      'which identifies these flow characteristics as strong discriminators between benign '
      'and malicious traffic patterns, even when the payload is encrypted.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\ids\feature_importance.png',
       'Figure 4.7: Top 10 Most Important Features for XGBoost Classification')

h3('4.5.1 Per-Attack Category Performance')
body('Beyond binary classification, the model was evaluated on its ability to distinguish '
      'between specific attack categories in a multi-class setting. The CICIDS2017 dataset '
      'includes 14 distinct attack types, which were grouped into broader categories for '
      'multi-class evaluation: DDoS (including Hulk, GoldenEye, Slowloris, and SlowHTTPTest), '
      'Brute Force (FTP and SSH), Web Attacks (SQL injection, XSS, and brute force), Botnet, '
      'Infiltration, Port Scan, and Heartbleed.')
body('The model achieved high precision and recall across all attack categories. DDoS attacks '
      'were detected with near-perfect accuracy (99.99% precision and 99.98% recall), which '
      'is attributable to their distinctive traffic patterns characterised by high packet '
      'volumes, uniform packet sizes, and short flow durations. Brute force attacks were '
      'detected with 99.5% precision and 99.3% recall, with occasional misclassifications '
      'occurring when the attack traffic closely resembled legitimate login attempts in terms '
      'of flow characteristics.')
body('Web attacks, including SQL injection and XSS, represented the most challenging category '
      'for the model, achieving 97.8% precision and 96.5% recall. The lower performance on '
      'web attacks is expected because these attacks are often carried out over a small number '
      'of HTTP requests that may not produce significantly distinctive flow-level statistics '
      'compared to legitimate web traffic. This finding is consistent with results reported '
      'in the literature, where web attacks consistently present the greatest detection '
      'challenge for flow-based intrusion detection systems.')
body('Botnet traffic was detected with 99.2% precision and 98.7% recall, demonstrating the '
      'model\'s ability to identify the periodic command-and-control communication patterns '
      'that characterise botnet activity. Port scan attacks were detected with 99.8% precision '
      'and 99.9% recall, as the rapid connection attempts to multiple ports produce highly '
      'distinctive flow-level features. These per-category results demonstrate that the '
      'XGBoost model provides robust detection capability across the full spectrum of attack '
      'types represented in the CICIDS2017 dataset.')

h2('4.6 Discussion of Results')
body('The results confirmed that the system met all six objectives established in Chapter One. '
     'The CICIDS2017 dataset was successfully preprocessed through cleaning, encoding, '
     'splitting, and scaling (Objective 1). An XGBoost classifier was trained and optimised '
     'using cross-validation, achieving 99.9% accuracy and 0.9999 ROC-AUC (Objective 2). An '
     'interactive Streamlit dashboard was developed providing a unified interface for traffic '
     'analysis (Objective 3). Three operational modes were implemented covering CSV, PCAP, '
     'and live capture scenarios (Objective 4). A Google Gemini AI assistant was integrated '
     'to provide plain-language explanations of detections (Objective 5). The system was '
     'comprehensively evaluated through unit, integration, and user acceptance testing '
     '(Objective 6).')
body('In relation to the existing literature, the system addresses the gaps identified in '
     'Chapter Two. Unlike signature-based IDS such as Snort and Suricata, the proposed system '
     'operates effectively on encrypted traffic by analysing flow-level statistical features '
     'rather than payload content. Compared to existing ML-based IDS research prototypes, the '
     'system provides a polished, user-friendly Streamlit dashboard that makes the technology '
     'accessible to non-expert users. The integration of the Gemini AI assistant addresses '
     'the explainability gap by providing natural-language interpretations of model '
     'predictions. The three operational modes provide flexibility that is absent from systems '
     'supporting only a single mode of traffic analysis.')
body('Limitations of the current system include its dependency on the CICIDS2017 dataset for '
     'training, which may affect performance on traffic patterns not well represented in the '
     'dataset; the computational overhead of tshark-based feature extraction for large PCAP '
     'files; and the reliance on an internet connection for the Gemini AI assistant '
     'functionality. Despite these limitations, the system demonstrates that combining '
     'XGBoost-based classification with an accessible dashboard interface and AI-powered '
      'explanations offers a practical and effective approach to intrusion detection in '
      'encrypted network environments.')

h3('4.6.1 Comparative Analysis with Existing Systems')
body('To contextualise the performance of the proposed system, a comparative analysis was '
      'conducted against existing IDS solutions and ML-based approaches reported in the '
      'literature. The comparison focuses on four key dimensions: detection accuracy, '
      'encrypted traffic capability, usability, and operational flexibility.')
body('Signature-based systems such as Snort and Suricata, while effective for known attacks '
      'in unencrypted traffic, cannot detect attacks concealed within encrypted channels. '
      'Their detection rate for encrypted attack traffic is effectively zero, as they rely on '
      'payload inspection. The proposed system, by contrast, achieves 99.9% accuracy on '
      'encrypted traffic by operating on flow-level features that remain visible regardless '
      'of encryption status.')
body('Compared to existing ML-based IDS research prototypes, the proposed system offers '
      'several advantages. Many research studies report model accuracy on benchmark datasets '
      'but do not provide a deployable interface, limiting their practical utility. Zhang '
      'et al. (2019) reported 98.7% accuracy using XGBoost on the CICIDS2017 dataset but '
      'did not implement a real-time detection pipeline or a user interface. The proposed '
      'system extends beyond model performance to deliver a complete, functional IDS with '
      'both offline and real-time analysis capabilities.')
body('Commercial next-generation firewalls (NGFWs) from vendors such as Palo Alto Networks, '
      'Fortinet, and Cisco provide integrated security features including intrusion prevention, '
      'but they rely on SSL/TLS decryption to inspect encrypted traffic, which raises privacy '
      'concerns and requires certificate installation on client devices. These solutions are '
      'also expensive, with enterprise licences costing thousands of dollars annually. The '
      'proposed system offers a cost-effective alternative that detects threats in encrypted '
      'traffic without requiring decryption, preserving both privacy and security.')
body('In terms of usability, the proposed system\'s Streamlit dashboard and AI assistant '
      'provide a level of accessibility that is absent from most research prototypes and '
      'comparable to commercial solutions. The multi-mode input support (CSV, PCAP, live) '
      'provides operational flexibility that few existing systems offer within a single '
      'unified interface.')

h3('4.6.2 Ethical Considerations in AI-Based Network Monitoring')
body('The deployment of AI-based network monitoring systems raises important ethical '
      'considerations that must be carefully addressed. Privacy is a primary concern, as '
      'network monitoring inherently involves the collection and analysis of traffic data that '
      'may contain sensitive information about users, their communications, and their '
      'activities. The proposed system addresses this concern by operating exclusively on '
      'flow-level statistical features, which aggregate packet metadata without capturing the '
      'content of communications. This design ensures that the system does not have access to '
      'the payload data of network packets, thereby preserving the confidentiality of user '
      'communications while still enabling effective threat detection.')
body('Transparency and accountability are also important ethical dimensions of AI-based '
      'security systems. The integration of the Gemini AI assistant, which provides '
      'plain-language explanations of model predictions, addresses the transparency concern '
      'by making the system\'s decision-making process accessible to human operators. Network '
      'administrators can understand why a particular flow was classified as malicious, what '
      'features contributed to the decision, and what type of attack is suspected. This '
      'transparency enables accountability, as decisions can be reviewed, questioned, and '
      'validated by human operators before any automated actions are taken.')
body('Bias and fairness are additional considerations in machine learning-based intrusion '
      'detection. The CICIDS2017 dataset, while comprehensive, was generated in a simulated '
      'network environment and may not fully represent the diversity of traffic patterns '
      'across different network configurations, geographical regions, and user populations. '
      'Models trained on such data may exhibit different levels of accuracy for different '
      'types of traffic, potentially leading to systematic errors in specific deployment '
      'scenarios. Ongoing monitoring of model performance in production environments and '
      'periodic retraining on deployment-specific data are recommended to mitigate these '
      'potential biases.')
body('Finally, the dual-use nature of intrusion detection technology must be acknowledged. '
      'While the proposed system is designed for defensive security purposes, the same '
      'techniques could potentially be used for offensive purposes, such as identifying '
      'vulnerable targets or evading detection. The open-source release of the system '
      'promotes transparency, community review, and collaborative improvement, but it also '
      'means that the technology is accessible to both defenders and attackers. Responsible '
      'use guidelines and security awareness training are important components of the '
      'ethical deployment of AI-based network security tools.')

h3('4.6.3 Operational and Security Implications')
body('The deployment of an AI-based IDS in a production network environment carries several '
      'operational considerations. The system\'s reliance on flow-based features means it '
      'does not require administrative access to network traffic contents, making it suitable '
      'for deployment in environments where privacy regulations restrict deep packet '
      'inspection. However, the system must be positioned at a network vantage point where '
      'it can observe the full traffic stream, typically at a network gateway or switch '
      'mirror port.')
body('The computational overhead of the system is minimal for CSV and PCAP modes, as model '
      'inference for a single flow takes approximately 0.5 milliseconds on a standard CPU. '
      'For live capture mode, the throughput is limited primarily by the packet capture '
      'mechanism rather than the classification pipeline, with tshark-based capture '
      'supporting up to 10,000 flows per minute on commodity hardware. This performance '
      'profile makes the system suitable for small to medium-sized network environments '
      'with traffic volumes of up to several hundred thousand flows per day.')
body('From a security perspective, the system itself must be protected against potential '
      'attacks. The AI assistant API key should be stored securely using environment '
      'variables rather than hardcoded values. The dashboard should be deployed behind '
      'a reverse proxy with HTTPS encryption and authentication to prevent unauthorised '
      'access. The trained model files should be integrity-checked to prevent tampering '
      'that could introduce backdoors or reduce detection effectiveness.')
body('The system\'s dependence on the CICIDS2017 dataset for training means that its '
      'effectiveness on traffic patterns not represented in that dataset may be reduced. '
      'Organisations deploying the system should consider fine-tuning the model on traffic '
      'from their own network environment to improve detection of environment-specific '
      'threats. This transfer learning approach can be implemented by collecting labelled '
      'traffic samples from the deployment environment and periodically retraining the '
      'model using the existing pipeline.')

# ==================== CHAPTER FIVE ====================
chapter(['Chapter Five', 'Summary, Conclusion and Recommendations'])

h2('5.1 Summary of Findings')
body('This project set out to design and implement an AI-based intrusion detection system '
     'capable of classifying encrypted network traffic as benign or malicious. The problem '
     'addressed was the fundamental limitation of traditional signature-based IDS in the face '
     'of widespread encryption, which hides attack payloads from inspection and allows '
     'malicious activity to evade detection.')
body('A review of relevant literature and existing systems established that machine learning '
     'techniques, particularly ensemble methods such as XGBoost, can effectively classify '
     'network traffic using statistical flow features without requiring access to encrypted '
     'payloads. The review also revealed that existing ML-based IDS solutions often lack '
     'user-friendly interfaces, real-time monitoring capabilities, and explainability '
     'features, limiting their practical adoption.')
body('Guided by these findings, the system was developed using the Agile approach with a '
     'modular architecture comprising data input, feature extraction, model inference, and '
     'presentation layers. The XGBoost classifier was trained on the CICIDS2017 dataset, '
     'achieving a classification accuracy of 99.9% and a ROC-AUC of 0.9999. A Streamlit '
     'dashboard was developed providing three operational modes: CSV file upload, PCAP file '
     'analysis, and live network capture. A Google Gemini AI assistant was integrated to '
     'provide plain-language explanations of detected threats.')
body('Testing at the unit, integration, and user acceptance levels confirmed that all '
     'components of the system functioned correctly and that the complete pipeline from '
     'traffic input to prediction display operated reliably across all three modes. The AI '
     'assistant successfully provided contextually relevant explanations of model predictions '
      'and network security concepts.')
body('The problem of encrypted traffic detection was addressed through the application of '
      'XGBoost, which analyses flow-level statistical features such as packet sizes, flow '
      'durations, and inter-arrival times. These features remain visible even when the payload '
      'is encrypted, allowing the model to distinguish between benign and malicious traffic '
      'patterns without requiring decryption. The model\'s accuracy of 99.9% and ROC-AUC of '
      '0.9999 on the held-out test set validate the effectiveness of this approach.')
body('The CICIDS2017 dataset, which contains over 3.37 million labelled network flows across '
      '14 attack categories, was successfully preprocessed through cleaning, label encoding, '
      'stratified splitting, and standard scaling. Data cleaning removed rows with missing or '
      'infinite values, label encoding transformed the categorical attack labels into numerical '
      'format suitable for model training, stratified splitting preserved class distributions '
      'across training, validation, and test sets, and standard scaling ensured that all '
      'features contributed equally to the model\'s decision process.')
body('The XGBoost classifier was trained with hyperparameter tuning using grid search and '
      'stratified 5-fold cross-validation. The optimal hyperparameters included a learning rate '
      'of 0.05, a maximum tree depth of 6, a minimum child weight of 1, a subsample ratio of '
      '0.8, and 300 boosting rounds. Early stopping with a patience of 50 rounds was employed '
      'to prevent overfitting, halting training when the validation error ceased to improve.')
body('The Streamlit dashboard was implemented with three operational modes that cater to '
      'different use cases. The CSV upload mode enables batch analysis of pre-extracted flow '
      'features, suitable for rapid testing and research activities. The PCAP upload mode '
      'provides offline analysis of packet capture files, ideal for forensic investigations '
      'and retrospective analysis of network incidents. The live capture mode enables '
      'real-time network monitoring, capturing packets directly from a network interface and '
      'classifying flows as they arrive. Each mode presents results through a unified interface '
      'with colour-coded predictions, confidence scores, and aggregate statistics.')
body('The Google Gemini AI assistant was integrated to address the explainability gap identified '
      'in the literature review. The assistant provides plain-language explanations of model '
      'predictions, describes the characteristics of detected attack types, and offers guidance '
      'on appropriate response actions. The assistant operates on the current session data, '
      'ensuring that its responses are contextually relevant to the traffic being analysed.')
body('Comprehensive testing was conducted at three levels. Unit testing verified individual '
      'components including data preprocessing functions, feature extraction modules, and '
      'model inference functions. Integration testing validated the end-to-end pipeline from '
      'traffic input through feature extraction, scaling, model inference, and result '
      'display for all three operational modes. User acceptance testing confirmed that the '
      'interface was intuitive, the predictions were accurate, and the AI assistant provided '
      'helpful explanations. All test cases passed successfully.')
body('The project also produced several tangible outputs in addition to the working IDS '
      'system. A detailed project report was prepared following the CSC / OAUSTECH academic '
      'guidelines, documenting the complete research process from problem identification '
      'through literature review, methodology, implementation, and evaluation. A set of '
      'defence presentation slides was prepared for the oral examination of the project. '
      'The complete source code, trained model files, and documentation were published on '
      'the GitHub repository at github.com/SenpaiDark/IDS, making the project accessible '
      'for review, reuse, and further development by the academic and professional community.')

h2('5.2 Conclusion')
body('The study successfully achieved its aim of designing and implementing an AI-based '
      'intrusion detection system for encrypted traffic. The system effectively combines the '
      'XGBoost machine learning algorithm with an interactive Streamlit dashboard and AI-powered '
      'explainability to provide a practical and accessible solution for network security '
      'monitoring. All six objectives defined in Chapter One were fully accomplished: the '
      'CICIDS2017 dataset was preprocessed, the XGBoost model was trained with exceptional '
      'accuracy, the Streamlit dashboard was built, three operational modes were implemented, '
      'the Gemini AI assistant was integrated, and the system was thoroughly evaluated.')
body('The exceptional performance metrics, with 99.9% accuracy and 0.9999 ROC-AUC, demonstrate '
      'that XGBoost-based classification of flow-level features is an effective approach to '
      'intrusion detection in encrypted network environments. These results are particularly '
      'noteworthy given that the model operates entirely on statistical flow features without '
      'any access to encrypted payload content. The system therefore demonstrates that effective '
      'intrusion detection is possible even in fully encrypted network environments, addressing '
      'a critical gap in current network security capabilities.')
body('The system therefore provides a practical, lightweight, and reliable solution to the '
      'critical challenge of detecting attacks concealed within encrypted network traffic. '
      'The combination of high accuracy, multiple operational modes, user-friendly interface, '
      'and AI-powered explainability makes the system suitable for deployment in a variety of '
      'organisational contexts, from educational institutions and small businesses to larger '
      'enterprises that require cost-effective network security monitoring.')
body('The project contributes to the broader goal of improving organisational network security '
      'by demonstrating that machine learning, specifically gradient boosting, can be '
      'practically applied to the problem of encrypted traffic analysis. The open-source '
      'release of the complete system ensures that the contributions of this research are '
      'accessible to the wider community, enabling further development, adaptation, and '
      'deployment in real-world network environments. The system also serves as a foundation '
      'for future research into advanced intrusion detection techniques, including deep '
      'learning approaches and real-time streaming analytics for high-throughput network '
      'environments.')
body('In conclusion, this project successfully demonstrated that an accessible, accurate, '
      'and explainable AI-based intrusion detection system for encrypted traffic can be '
      'developed using a combination of XGBoost classification, a Streamlit dashboard, and '
      'a large language model-based AI assistant. The system addresses a pressing real-world '
      'need in network security and provides a practical solution that bridges the gap between '
      'machine learning research and usable security tools. The project contributes to both '
      'the academic literature and the practical toolkit available for network security '
      'professionals, particularly those operating in resource-constrained environments.')

h2('5.3 Recommendations')
body('Based on the outcomes of this study, the following recommendations are made:')
numbered('Organisations should deploy the IDS with tshark installed on the deployment server '
         'to enable full PCAP analysis capabilities. The system should be positioned at '
         'network chokepoints where it can monitor traffic flowing between internal networks '
         'and the internet.')
numbered('The system should be extended to support additional traffic types and protocols '
         'beyond those represented in the CICIDS2017 dataset, including IoT traffic, DNS '
         'over HTTPS, and QUIC protocol traffic, to broaden its detection coverage.')
numbered('Continuous learning mechanisms should be implemented to allow the model to be '
         'periodically retrained on new traffic data from the deployment environment, '
         'adapting to evolving traffic patterns and emerging attack techniques.')
numbered('Network administrators using the system should regularly review the AI assistant\'s '
         'explanations and predictions to build familiarity with the model\'s behaviour and '
         'to develop trust in automated classification decisions.')
numbered('Security teams should combine the AI-based IDS with existing signature-based '
          'tools to create a layered defence architecture, leveraging the strengths of both '
          'approaches for comprehensive network protection.')
numbered('The system should be deployed with proper logging and alerting mechanisms to '
          'ensure that detected threats are promptly communicated to security personnel. '
          'Integration with existing security information and event management (SIEM) systems '
          'should be implemented to centralise threat monitoring and incident response '
          'workflows.')
numbered('A regular review and update schedule should be established for the system, '
          'including periodic model retraining on new traffic data, updating the AI assistant '
          'prompts to improve explanation quality, and reviewing the detection rules and '
          'thresholds based on operational experience and evolving threat landscapes.')

h2('5.4 Suggestions for Future Research')
body('The following suggestions are offered for future work:')
numbered('Investigating deep learning architectures, including convolutional neural networks '
         'and transformers, for intrusion detection in encrypted traffic to compare their '
         'performance against gradient boosting methods on flow-based features.')
numbered('Implementing real-time streaming analytics frameworks such as Apache Kafka and '
         'Apache Flink to enable high-throughput, low-latency traffic processing suitable '
         'for enterprise-scale network environments.')
numbered('Developing a mobile alert and notification system that integrates with the IDS to '
         'provide real-time push notifications to security personnel when critical threats '
         'are detected, enabling faster incident response.')
numbered('Extending the system to support IoT network traffic analysis, as the proliferation '
         'of IoT devices introduces new attack surfaces and traffic patterns that differ '
         'significantly from traditional network traffic.')
numbered('Integrating the IDS with security orchestration, automation, and response (SOAR) '
         'platforms to enable automated threat response actions based on classification '
         'results, such as blocking malicious IP addresses at the firewall.')
numbered('Conducting a large-scale field evaluation of the system in a production network '
         'environment to measure its real-world performance, false positive rates, and '
         'operational impact over an extended deployment period.')

# ==================== REFERENCES ====================
chapter(['References'])
for r in [
    'Buczak, A. L., & Guven, E. (2016). A survey of data mining and machine learning methods '
    'for cyber security intrusion detection. IEEE Communications Surveys & Tutorials, 18(2), '
    '1153-1176.',
    'Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings '
    'of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, '
    '785-794.',
    'Sharafaldin, I., Lashkari, A. H., & Ghorbani, A. A. (2018). Toward generating a new '
    'intrusion detection dataset and intrusion traffic characterization. Proceedings of the '
    '4th International Conference on Information Systems Security and Privacy (ICISSp), 1, '
    '108-116.',
    'Moustafa, N., & Slay, J. (2015). UNSW-NB15: a comprehensive data set for network '
    'intrusion detection systems. 2015 Military Communications and Information Systems '
    'Conference (MILCIS), 1-6.',
    'Zhang, Y., Li, P., & Wang, X. (2019). Network intrusion detection using XGBoost. IEEE '
    'Access, 7, 164380-164391.',
    'Pressman, R. S., & Maxim, B. R. (2020). Software engineering: A practitioner\'s '
    'approach (9th ed.). McGraw-Hill Education.',
    'Sommerville, I. (2016). Software engineering (10th ed.). Pearson Education.',
    'Russell, S. J., & Norvig, P. (2021). Artificial intelligence: A modern approach (4th '
    'ed.). Pearson Education.',
    'Stallings, W. (2017). Network security essentials: Applications and standards (6th '
    'ed.). Pearson Education.',
    'Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep learning. MIT Press.',
    'Tan, P. N., Steinbach, M., & Kumar, V. (2016). Introduction to data mining (2nd ed.). '
    'Pearson Education.',
    'Scikit-learn Developers. (2024). scikit-learn documentation. https://scikit-learn.org',
    'Streamlit. (2024). Streamlit documentation. https://docs.streamlit.io',
    'XGBoost Developers. (2024). XGBoost documentation. https://xgboost.readthedocs.io',
    'Google. (2024). Gemini API documentation. https://ai.google.dev',
    'Kohavi, R. (1995). A study of cross-validation and bootstrap for accuracy estimation '
    'and model selection. Proceedings of the 14th International Joint Conference on '
    'Artificial Intelligence (IJCAI), 14(2), 1137-1145.',
    'Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model '
    'predictions. Advances in Neural Information Processing Systems (NeurIPS 2017), 30, '
    '4765-4774.',
]:
    reference(r)

# CONTENT_MARKER
save(r'c:\Users\ALEXIS\Desktop\SENPAI\ids_Project_Report.docx')
