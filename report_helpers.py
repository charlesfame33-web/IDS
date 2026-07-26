# -*- coding: utf-8 -*-
"""Generate mayor4code BSc project report (full: prelims + Chapters 1-5 +
References) per CSC / OAUSTECH guidelines."""
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- Base formatting: Times New Roman 12, A4, margins, 1.5 spacing ----
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(12)

# Heading styles (so the auto Table of Contents can pick them up)
for name, size in [('Heading 1', 14), ('Heading 2', 12), ('Heading 3', 12)]:
    st = doc.styles[name]
    st.font.name = 'Times New Roman'
    st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.italic = (name == 'Heading 3')
    st.font.color.rgb = RGBColor(0, 0, 0)
    st.paragraph_format.space_before = Pt(6)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

sec = doc.sections[0]
sec.page_width = Mm(210)
sec.page_height = Mm(297)
sec.left_margin = Mm(40)
sec.right_margin = Mm(25)
sec.top_margin = Mm(25)
sec.bottom_margin = Mm(25)


def _page_field(paragraph):
    """Insert a PAGE field into a footer paragraph."""
    run = paragraph.add_run()
    fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end')
    run._r.append(fb); run._r.append(it); run._r.append(fe)


def set_page_numbering(section, fmt, start=None, hide_first=False):
    """fmt: 'lowerRoman' or 'decimal'. Adds centered page number in footer."""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType'); sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))
    if hide_first:
        section.different_first_page_header_footer = True
        section.first_page_footer.is_linked_to_previous = False
        # leave first-page footer empty (title page: no number)
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ''
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _page_field(p)


def add_toc(field='TOC'):
    """Insert an auto-updating field (TOC / table of figures)."""
    para = doc.add_paragraph()
    run = para.add_run()
    fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
    if field == 'TOC':
        it.text = 'TOC \\o "1-3" \\h \\z \\u'
    fsep = OxmlElement('w:fldChar'); fsep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = 'Right-click and choose "Update Field" to build this table.'
    hold = OxmlElement('w:r'); hold.append(t)
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end')
    run._r.append(fb); run._r.append(it); run._r.append(fsep)
    run._r.append(hold); run._r.append(fe)


def chapter(title_lines, new_section=False):
    """Chapter/major heading as Heading 1, centered, starting on a new page.
    Accepts a list (legacy) or a string. new_section=True begins Arabic
    numbering for the main body."""
    if isinstance(title_lines, (list, tuple)):
        title = ' '.join(title_lines)
    else:
        title = title_lines
    if new_section:
        new = doc.add_section(WD_SECTION.NEW_PAGE)
        new.page_width = Mm(210); new.page_height = Mm(297)
        new.left_margin = Mm(40); new.right_margin = Mm(25)
        new.top_margin = Mm(25); new.bottom_margin = Mm(25)
        set_page_numbering(new, 'decimal', start=1)
    else:
        doc.add_page_break()
    p = doc.add_heading('', level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title.upper())


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def prelim_title(text):
    """Centered bold heading for preliminary pages; NOT in the TOC."""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(14)
    return p


def body(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


# ---- Convenience: caption + title-page helpers shared by both reports ----
def figure(img_path, caption_text):
    """Embed an image centred, with a caption below."""
    import os
    if not os.path.exists(img_path):
        body(f'[Image not found: {img_path}]')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run()
    r.add_picture(img_path, width=Mm(140))
    cap = doc.add_paragraph(caption_text)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(12)
    cap.runs[0].italic = True


def caption(text):
    """Figure/Table caption: centered, italic."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    return p


def centered(text, bold=False, size=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    return p


def reference(text):
    """APA hanging-indent reference entry."""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Mm(-12.7)
    p.paragraph_format.left_indent = Mm(12.7)
    return p


def save(path):
    doc.save(path)
    print('saved', path)
