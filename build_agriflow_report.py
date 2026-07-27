# -*- coding: utf-8 -*-
"""Generate the AgriFlow AI BSc project report per CSC / OAUSTECH guidelines.
Reuses shared formatting helpers from report_helpers.py."""
import report_helpers as R
from report_helpers import (
    doc, chapter, h2, h3, body, bullet, numbered, prelim_title,
    add_toc, set_page_numbering, figure, caption, centered, reference, save,
)
from docx.shared import Pt, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, sys
sys.path.insert(0, os.path.dirname(__file__))

# ---- Generate flowchart ----
FLOWCHART = r'c:\Users\ALEXIS\Desktop\SENPAI\shots\agriflow\scheduler_flowchart.png'
_gf = r'c:\Users\ALEXIS\Desktop\SENPAI\shots\agriflow\generate_flowchart.py'
if os.path.exists(_gf):
    import subprocess, sys
    subprocess.run([sys.executable, _gf], check=True)

TITLE = ('AI-DRIVEN AGRICULTURAL SUPPLY CHAIN AND PRODUCE SCHEDULING SYSTEM '
         '(A CASE STUDY OF AGRIFLOW AI)')

# ==================== TITLE PAGE ====================
centered(TITLE, bold=True)
for line in ['', 'BY', '', 'Obayomi Samuel Oluwagbotemi', 'CSC/22/124', '',
             'DEPARTMENT OF COMPUTER SCIENCE', 'SCHOOL OF COMPUTING',
             'OLUSEGUN AGAGU UNIVERSITY OF SCIENCE AND TECHNOLOGY, OKITIPUPA', '',
             'IN PARTIAL FULFILMENT OF THE REQUIREMENTS FOR THE AWARD OF THE '
             'DEGREE OF BACHELOR OF SCIENCE (B.Sc.) IN COMPUTER SCIENCE', '',
             'August, 2026']:
    centered(line)

# Preliminary pages: Roman numerals, title page unnumbered
set_page_numbering(doc.sections[0], 'lowerRoman', start=1, hide_first=True)

# ---- Certification ----
prelim_title('Certification')
body('This is to certify that this project titled “AI-Driven Agricultural Supply Chain and '
     'Produce Scheduling System (A Case Study of AgriFlow AI)” was '
     'carried out by Obayomi Samuel Oluwagbotemi with matriculation number CSC/22/124 in the '
     'Department of Computer Science, School of Computing, Olusegun Agagu University of Science '
     'and Technology, Okitipupa, in partial fulfilment of the requirements for the award of the '
     'degree of Bachelor of Science (B.Sc.) in Computer Science.')
table = doc.add_table(rows=4, cols=2)
table.alignment = WD_ALIGN_PARAGRAPH.CENTER
table.style = 'Table Grid'
cells = [
    ('_______________________', '_______________________'),
    ('Dr. Adeolu Obamehinti',  'Dr. Ajoke Gbadamosi'),
    ('Project Supervisor',     'Head of Department'),
    ('Date: __________________','Date: __________________'),
]
for r, (left, right) in enumerate(cells):
    for c, val in [(0, left), (1, right)]:
        cell = table.cell(r, c)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(val)

# ---- Declaration ----
prelim_title('Declaration')
body('I hereby declare that this project is my original work and has not been submitted, either '
     'in whole or in part, for any other degree or qualification in this or any other '
     'institution. All sources of information and materials used have been duly acknowledged by '
     'means of complete references.')
for line in ['', '', '_______________________', 'Obayomi Samuel Oluwagbotemi', 'CSC/22/124',
             'Date: __________________']:
    doc.add_paragraph(line)

# ---- Dedication ----
prelim_title('Dedication')
body('This project is dedicated to Almighty God, and to my family, whose support and '
     'encouragement made this work possible.')

# ---- Acknowledgements ----
prelim_title('Acknowledgements')
body('I express my sincere gratitude to my project supervisor, Dr. Adeolu Obamehinti, for the '
     'guidance, patience, and constructive criticism that shaped this work. I am grateful to the '
     'Head and entire staff of the Department of Computer Science, Olusegun Agagu University of '
     'Science and Technology, Okitipupa, for the knowledge imparted throughout my programme.')
body('Special appreciation goes to Daniel Atere (Senpai Dark), whose expert guidance, '
     'mentorship, and technical direction were instrumental in the successful completion of this '
     'project. His deep insights into software architecture and development significantly '
     'contributed to the quality of the work.')
body('I also thank my family and friends for their unwavering support. Above all, I give thanks '
     'to Almighty God for His grace and strength.')

# ---- Abstract ----
prelim_title('Abstract')
body('Agricultural supply chains for fresh produce are hampered by post-harvest spoilage, '
     'inefficient delivery scheduling, uncertain demand, and the disconnection of farmers, '
     'buyers, transporters, and warehouse managers who rely on manual and fragmented processes. '
     'This project presents the design and implementation of AgriFlow AI, an intelligent '
     'web-based platform that unifies these participants and applies deterministic decision '
     'rules, complemented by an artificial-intelligence assistant, to improve supply chain '
     'efficiency. The system was developed using the Agile approach within the Software '
     'Development Life Cycle and adopts a modern serverless web architecture. It was implemented '
     'with the Next.js 14 framework and the TypeScript language, styled with Tailwind CSS, and '
     'backed by a Supabase PostgreSQL database with row-level security for authentication and '
     'authorisation. The platform provides role-based access for five user types; a produce '
     'marketplace; crop, harvest, inventory, and order management; a perishable-first delivery '
     'scheduling engine; moving-average demand forecasting; shelf-life spoilage tracking; '
     'rule-based notifications; and an assistant powered by Google Gemini that explains live '
     'operational data in plain language. All business logic is deterministic and testable, with '
     'the artificial intelligence used only for explanation. The platform was tested through '
     'unit, functional, and end-to-end testing across all five roles, and the results showed '
     'that it correctly prioritised perishable orders, assigned deliveries efficiently, tracked '
     'shelf life accurately, and produced reliable forecasts. The study concludes that combining '
     'a unified digital platform with deterministic optimisation and explanatory artificial '
     'intelligence significantly improves the coordination and efficiency of agricultural supply '
     'chains. It is recommended that the platform be extended with mobile applications, live '
     'logistics tracking, and predictive machine-learning models in future work.')

# ---- Table of Contents ----
prelim_title('Table of Contents')
add_toc('TOC')

# ---- List of Figures ----
prelim_title('List of Figures')
for line in ['Figure 2.1: Comparison of Existing Systems by Capability',
             'Figure 3.1: System Architecture of the AgriFlow AI Platform',
             'Figure 3.2: Entity Relationship Diagram of the Database',
             'Figure 3.3: Use-Case Diagram of the System',
             'Figure 3.4: Delivery Scheduling Flowchart',
             'Figure 4.1: Landing Page',
             'Figure 4.2: Role-Based Signup and Login',
             'Figure 4.3: Farmer Dashboard and Crop Management',
             'Figure 4.4: Buyer Marketplace and Ordering',
             'Figure 4.5: Warehouse Inventory with Shelf-Life Alerts',
             'Figure 4.6: Transporter Delivery Workflow and Route Map',
             'Figure 4.7: Admin Scheduling and Reports',
             'Figure 4.8: AI Assistant Interface']:
    body(line)

# ---- List of Tables ----
prelim_title('List of Tables')
for line in ['Table 2.1: Comparison of Existing Systems by Capability',
             'Table 3.1: Software Development Tools and Technologies',
             'Table 3.2: Summary of Database Tables',
             'Table 3.3: Functional Requirements Specification',
             'Table 3.4: Non-Functional Requirements Specification',
             'Table 4.1: Interface Features by Screen',
             'Table 4.2: Summary of Test Cases and Results']:
    body(line)

# ---- List of Abbreviations ----
prelim_title('List of Abbreviations')
for line in ['AI — Artificial Intelligence',
             'API — Application Programming Interface',
             'CRUD — Create, Read, Update, Delete',
             'CSS — Cascading Style Sheets',
             'DBMS — Database Management System',
             'ERD — Entity Relationship Diagram',
             'PWA — Progressive Web Application',
             'RLS — Row-Level Security',
             'SDLC — Software Development Life Cycle',
             'SQL — Structured Query Language',
             'SSR — Server-Side Rendering',
             'UI — User Interface']:
    body(line)

# ==================== CHAPTER ONE ====================
chapter(['Chapter One', 'Introduction'], new_section=True)

h2('1.1 Background of the Study')
body('Agriculture remains one of the most important sectors of the economy, particularly in '
     'developing nations where it provides food, employment, and raw materials for a large '
     'proportion of the population. Despite its importance, the agricultural supply chain, which '
     'connects producers to consumers, is frequently inefficient. This is especially true for '
     'fresh and perishable produce, where delays, poor coordination, and inadequate storage lead '
     'to significant post-harvest losses.')
body('The supply chain for perishable produce involves several participants, including farmers '
     'who cultivate crops, buyers who purchase produce, transporters who move goods, and '
     'warehouse managers who store them. In many settings these participants operate in '
     'isolation, relying on manual processes, telephone calls, and informal arrangements. This '
     'fragmentation results in poor visibility, delayed deliveries, and the spoilage of produce '
     'that could otherwise have been sold.')
body('Advances in information and communication technology, and in particular the growth of '
     'web-based and artificial-intelligence-assisted systems, provide an opportunity to address '
     'these challenges. Web platforms enable geographically dispersed participants to coordinate '
     'in real time, while data-driven decision techniques can optimise scheduling, anticipate '
     'demand, and reduce waste. According to Pressman and Maxim (2020), well-engineered software '
     'systems can automate complex operational processes and improve decision-making across '
     'organisations.')
body('This study presents the design and implementation of AgriFlow AI, an intelligent '
     'web-based agricultural supply chain management platform. The platform unifies farmers, '
     'buyers, transporters, warehouse managers, and administrators within a single system and '
     'applies deterministic decision rules for delivery scheduling, demand forecasting, and '
     'spoilage tracking. It further incorporates an artificial-intelligence assistant that '
     'explains the platform’s live operational data in plain language, thereby making complex '
     'information accessible to all users.')

h2('1.2 Research Justification/Motivation')
body('The motivation for this study arises from the persistent and costly problem of '
     'post-harvest losses in agricultural supply chains. A substantial proportion of fresh '
     'produce is lost between harvest and consumption due to poor logistics, inadequate storage, '
     'and weak coordination among supply chain participants. These losses reduce the incomes of '
     'farmers, raise prices for consumers, and waste scarce resources.')
body('Existing efforts to digitise agricultural trade often focus narrowly on connecting buyers '
     'and sellers, without addressing the operational realities of moving perishable goods '
     'efficiently. There is a need for an integrated platform that not only provides a '
     'marketplace but also schedules deliveries intelligently, anticipates demand, monitors '
     'shelf life, and keeps every participant informed.')
body('Furthermore, the increasing accessibility of artificial-intelligence services makes it '
     'possible to present complex operational data in an understandable form. By grounding an '
     'assistant in the platform’s live data, the system can help users who are not data experts '
     'to interpret forecasts and identify risks. These considerations motivated the development '
     'of AgriFlow AI as a unified, intelligent, and practical solution.')

h2('1.3 Problem Statement')
body('The management of agricultural supply chains for perishable produce faces several '
     'interrelated challenges. First, the participants in the chain are often disconnected, '
     'relying on manual and informal processes that provide little visibility and result in '
     'poor coordination.')
body('Second, deliveries are frequently scheduled without regard to the perishability of the '
     'goods, so that produce with a short remaining shelf life is not prioritised, leading to '
     'avoidable spoilage. Third, demand is uncertain and is rarely anticipated in a systematic '
     'way, causing both shortages and surpluses.')
body('Fourth, the operational data generated by such systems is often complex and difficult for '
     'non-specialist users to interpret, limiting its usefulness for decision-making. The '
     'absence of an integrated platform that addresses these issues together hinders the '
     'efficiency of the agricultural supply chain and contributes to unnecessary losses.')
body('This study addresses these challenges through the development of AgriFlow AI, an '
     'integrated web-based platform that unifies supply chain participants, schedules deliveries '
     'with a perishable-first rule, forecasts demand, tracks shelf life, and provides an '
     'artificial-intelligence assistant that explains the data in plain language.')

h2('1.4 Aim and Objectives of the Study')
h3('Aim')
body('The aim of this study is to design and implement an intelligent web-based agricultural '
     'supply chain management platform that unifies supply chain participants and applies '
     'data-driven decision rules to improve efficiency and reduce post-harvest losses.')
h3('Objectives')
body('The specific objectives of the study are to:')
numbered('design a web-based platform that connects farmers, buyers, transporters, warehouse '
         'managers, and administrators with role-based access;')
numbered('implement a produce marketplace with crop, harvest, inventory, and order management;')
numbered('develop a deterministic delivery scheduling engine that prioritises perishable orders '
         'and assigns the nearest warehouse and least-busy transporter;')
numbered('implement demand forecasting and shelf-life spoilage tracking to support planning and '
         'reduce waste;')
numbered('integrate an artificial-intelligence assistant that explains the platform’s live '
         'operational data in plain language;')
numbered('evaluate the functionality and reliability of the developed platform.')

h2('1.5 Scope of the Study')
body('This project focuses on the design and implementation of a web-based platform for managing '
     'the supply chain of fresh agricultural produce. The platform provides role-based '
     'authentication for five user types; a marketplace for listing and ordering produce; '
     'management of crops, harvests, warehouse inventory, orders, and deliveries; a delivery '
     'scheduling engine; demand forecasting; shelf-life spoilage tracking; rule-based '
     'notifications; analytics and administration; and an artificial-intelligence assistant for '
     'explanation.')
body('The system is a web application accessible through modern browsers and installable as a '
     'progressive web application. The study is limited to the coordination and optimisation of '
     'the produce supply chain and does not extend to financial payment processing, physical '
     'internet-of-things sensor integration, or native mobile applications. The '
     'artificial-intelligence component is used solely to explain data, while all business '
     'decisions are made by deterministic and verifiable rules.')

h2('1.6 Significance of the Study')
body('The development of the AgriFlow AI platform offers significant benefits to the '
     'agricultural sector and its stakeholders.')
h3('1.6.1 Economic Impact')
body('By prioritising perishable produce, optimising deliveries, and tracking shelf life, the '
     'platform reduces post-harvest losses and the associated financial costs. This improves the '
     'incomes of farmers and can lower prices for buyers, contributing to a more efficient and '
     'profitable agricultural economy.')
h3('1.6.2 Social Impact')
body('The platform strengthens coordination among farmers, buyers, transporters, and warehouse '
     'managers, fostering trust and collaboration. By reducing food waste, it also contributes '
     'to food security and the more responsible use of resources within the community.')
h3('1.6.3 Technological Impact')
body('The project demonstrates the practical integration of modern web technologies, '
     'cloud-based databases, deterministic optimisation, and artificial intelligence to solve a '
     'real-world problem. It contributes to the digital transformation of agriculture and '
     'provides a model for combining reliable rule-based logic with explanatory artificial '
     'intelligence.')

h2('1.7 Definition of Terms')
body('Supply Chain: The network of participants and processes involved in moving a product from '
     'producer to consumer.')
body('Perishable Produce: Agricultural goods that deteriorate quickly and have a limited shelf '
     'life, such as fresh fruits and vegetables.')
body('Shelf Life: The length of time for which produce remains fit for sale or consumption '
     'after harvest.')
body('Demand Forecasting: The estimation of future demand for a product based on historical '
     'data.')
body('Scheduling Engine: A software component that determines when and how deliveries are '
     'carried out according to defined rules.')
body('Artificial Intelligence (AI): The capability of a computer system to perform tasks that '
     'normally require human intelligence, such as interpreting data and generating '
     'explanations.')
body('Row-Level Security (RLS): A database security mechanism that restricts which rows a user '
     'may access based on defined policies.')
body('Progressive Web Application (PWA): A web application that can be installed on a device and '
     'used with offline capabilities, behaving like a native application.')

h2('1.8 Organization of the Study')
body('This project report is organised into five chapters.')
body('Chapter One presents the introduction, background of the study, motivation, problem '
     'statement, aim and objectives, scope, significance, definition of terms, and organisation '
     'of the study.')
body('Chapter Two reviews relevant literature related to agricultural supply chain management, '
     'e-commerce and marketplaces, optimisation and forecasting techniques, artificial '
     'intelligence in agriculture, and existing systems.')
body('Chapter Three describes the methodology adopted for the design and development of the '
     'proposed system, including the design approach, system architecture, database design, and '
     'implementation tools.')
body('Chapter Four presents the implementation details, system interfaces, testing procedures, '
     'results, and discussion of findings.')
body('Chapter Five provides the summary of the study, conclusion, recommendations, and '
     'suggestions for future research.')

# ==================== CHAPTER TWO ====================
chapter(['Chapter Two', 'Literature Review'])

h2('2.1 Introduction')
body('This chapter reviews relevant literature related to agricultural supply chain management, '
     'electronic marketplaces, supply chain optimisation and demand forecasting, artificial '
     'intelligence in agriculture, web development technologies, and existing systems. The review '
     'establishes the conceptual foundation of the study and identifies the gaps that justify the '
     'development of the proposed AgriFlow AI platform.')

h2('2.2 Agricultural Supply Chain Management')
body('A supply chain is the network of organisations, people, activities, and resources involved '
     'in moving a product from producer to consumer. Agricultural supply chains are '
     'distinctive because they deal with biological products that are seasonal, variable in '
     'quality, and frequently perishable. The management of such chains involves coordinating '
     'production, storage, transportation, and distribution to deliver produce in good condition '
     'and at the right time.')
body('Post-harvest loss is a central concern in agricultural supply chains, particularly for '
     'fresh produce. Losses occur through spoilage, poor handling, inadequate storage, and '
     'inefficient logistics. Effective supply chain management seeks to minimise these losses by '
     'improving coordination and by ensuring that perishable goods are prioritised in storage '
     'and delivery decisions.')

h2('2.3 Electronic Marketplaces and E-Agriculture')
body('An electronic marketplace is an online platform that brings buyers and sellers together to '
     'exchange goods and services. In agriculture, electronic marketplaces connect farmers '
     'directly with buyers, reducing dependence on intermediaries and improving price '
     'transparency. The broader application of information technology to agriculture, often '
     'termed e-agriculture, encompasses digital tools for trade, advisory services, and supply '
     'chain coordination.')
body('While marketplaces improve access to buyers, a marketplace alone does not solve the '
     'logistical challenges of moving perishable produce. Sommerville (2016) observed that '
     'web-based systems are well suited to coordinating distributed participants, but their value '
     'is greatest when they integrate operational processes rather than merely facilitating '
     'transactions. This motivates a platform that combines a marketplace with logistics, '
     'scheduling, and monitoring.')

h2('2.4 Supply Chain Optimisation and Demand Forecasting')
body('Optimisation techniques are widely used to improve supply chain decisions such as routing, '
     'scheduling, and inventory management. Rule-based and heuristic approaches are often '
     'preferred in practical systems because they are transparent, predictable, and '
     'computationally efficient. A perishable-first scheduling rule, for example, orders '
     'deliveries by the remaining shelf life of the goods so that the most time-critical produce '
     'is dispatched first.')
body('Demand forecasting estimates future demand from historical data and supports planning and '
     'inventory decisions. Classical techniques such as the simple moving average and the '
     'weighted moving average are commonly used because they are easy to compute and interpret. '
     'The moving average smooths short-term fluctuations, while the weighted moving average '
     'gives greater importance to more recent observations, making it more responsive to trends.')
body('The geographic distance between locations, used in nearest-warehouse and routing '
     'decisions, is commonly computed with the haversine formula, which gives the great-circle '
     'distance between two points on the Earth from their latitude and longitude. Deterministic '
     'techniques of this kind provide reliable and explainable decisions, which is important in '
     'operational systems where users must trust the outcomes.')

h2('2.5 Artificial Intelligence in Agriculture')
body('Artificial intelligence refers to the capability of computer systems to perform tasks that '
     'normally require human intelligence. In agriculture, artificial intelligence has been '
     'applied to crop monitoring, yield prediction, disease detection, and decision support. '
     'Large language models, a recent class of artificial-intelligence systems, are able to '
     'interpret data and generate natural-language explanations.')
body('A significant risk in applying large language models to operational systems is that they '
     'may generate inaccurate or fabricated information. To mitigate this, such models can be '
     'grounded in verified data and restricted to explanation rather than decision-making. In '
     'the proposed system, the artificial-intelligence assistant is grounded in the platform’s '
     'live data and is instructed to use only the information provided, while all decisions are '
     'made by deterministic rules. This design combines the reliability of rule-based logic with '
     'the accessibility of natural-language explanation.')

h2('2.6 Web Development Technologies')
body('Modern web applications are commonly built with component-based frameworks that support '
     'both server-side rendering and interactive client behaviour. Next.js, a framework based on '
     'the React library, provides server components, server actions, and routing, enabling '
     'efficient and maintainable applications. TypeScript, a typed superset of JavaScript, '
     'improves reliability by detecting type errors during development.')
body('Data in such applications is often managed by cloud database platforms. Supabase provides '
     'a hosted PostgreSQL database together with authentication and row-level security, allowing '
     'fine-grained control over data access. Elmasri and Navathe (2017) noted that database '
     'systems ensure data consistency, integrity, security, and efficient retrieval, all of '
     'which are essential for a supply chain platform. Supporting technologies include Tailwind '
     'CSS for styling, Leaflet for interactive maps over OpenStreetMap data, and charting '
     'libraries for analytics.')

h2('2.7 Related Works and Existing Systems')

h3('2.7.1 Twiga Foods Supply Chain Platform')
body('Approach')
body('Twiga Foods is a business-to-business platform that links farmers with vendors and manages '
     'the distribution of fresh produce, using mobile ordering and a centralised logistics '
     'operation.')
body('Architectural Framework')
body('The system combines a mobile ordering interface, a central operations and logistics '
     'backend, warehousing, and a distribution fleet coordinated through a central database.')
body('Strengths')
bullet('Aggregates demand and streamlines distribution of fresh produce.')
bullet('Reduces the number of intermediaries between farmers and vendors.')
bullet('Operates at scale with a managed logistics network.')
body('Weaknesses')
bullet('Relies on a heavy, company-operated logistics infrastructure.')
bullet('Limited self-service scheduling and forecasting for independent users.')
bullet('Not an openly configurable multi-role platform.')

h3('2.7.2 FarmCrowdy Digital Agriculture Platform')
body('Approach')
body('FarmCrowdy is a Nigerian digital agriculture platform that connects farmers with resources '
     'and markets, with an emphasis on linking smallholder farmers to buyers and support '
     'services.')
body('Architectural Framework')
body('It consists of a web and mobile interface, a user and farm management backend, and a '
     'database that records farms, produce, and transactions.')
body('Strengths')
bullet('Improves market access for smallholder farmers.')
bullet('Provides visibility of produce to buyers.')
bullet('Supports the agricultural value chain in a local context.')
body('Weaknesses')
bullet('Limited focus on perishable-first delivery scheduling.')
bullet('Little emphasis on shelf-life-based spoilage tracking.')
bullet('Forecasting and explanatory analytics are limited.')

h3('2.7.3 Generic ERP Supply Chain Modules')
body('Approach')
body('Enterprise resource planning systems provide supply chain modules that manage inventory, '
     'orders, and logistics for organisations across many sectors.')
body('Architectural Framework')
body('These systems use a multi-tier architecture with a central database, application server, '
     'and client interfaces, integrating procurement, inventory, and distribution functions.')
body('Strengths')
bullet('Comprehensive and mature functionality.')
bullet('Strong integration across business processes.')
bullet('Established reporting and auditing capabilities.')
body('Weaknesses')
bullet('Generic rather than tailored to perishable agricultural produce.')
bullet('Costly and complex to deploy and maintain.')
bullet('Rarely include explanatory artificial-intelligence assistance.')

h3('2.7.4 Marketplace-Only Agricultural Applications')
body('Approach')
body('A number of applications provide online marketplaces that connect farmers and buyers for '
     'the sale of produce, focusing primarily on listings and transactions.')
body('Architectural Framework')
body('They typically comprise a web or mobile front end, a product catalogue, an ordering '
     'module, and a database of users and listings.')
body('Strengths')
bullet('Straightforward to use for buying and selling.')
bullet('Improve price transparency and market reach.')
bullet('Low barrier to entry for farmers.')
body('Weaknesses')
bullet('Do not address delivery scheduling or logistics optimisation.')
bullet('Lack shelf-life monitoring and demand forecasting.')
bullet('Provide limited operational visibility beyond transactions.')

h3('2.7.5 Summary of Related Works')
body('The reviewed systems confirm the value of connecting agricultural participants digitally '
     'and of managing produce distribution at scale. However, they also reveal recurring '
     'limitations: a focus on either marketplace transactions or company-operated logistics, '
     'rather than an integrated and configurable platform; limited perishable-first scheduling; '
     'weak shelf-life-based spoilage tracking; limited accessible forecasting; and the absence '
     'of explanatory artificial-intelligence assistance grounded in live data. These '
     'observations inform the design of the proposed system.')

h3('2.7.6 Comparative Analysis')
body('Table 2.1 presents a comparative analysis of the reviewed platforms alongside the proposed '
     'AgriFlow AI system. The comparison evaluates each platform against the key capabilities '
     'identified in the literature: multi-role support, marketplace functionality, delivery '
     'scheduling, demand forecasting, shelf-life tracking, and artificial-intelligence assistance.')
caption('Table 2.1: Comparison of Existing Systems by Capability')
tbl = doc.add_table(rows=6, cols=6)
tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
tbl.style = 'Table Grid'
headers = ['Platform', 'Multi-Role', 'Marketplace', 'Scheduling', 'Forecasting', 'AI Assistant']
data = [
    ['Twiga Foods', 'Partial', 'Yes', 'Yes', 'No', 'No'],
    ['FarmCrowdy', 'Partial', 'Yes', 'No', 'Limited', 'No'],
    ['ERP Modules', 'Yes', 'Yes', 'Yes', 'Yes', 'No'],
    ['Marketplace Apps', 'No', 'Yes', 'No', 'No', 'No'],
    ['AgriFlow AI', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
]
for c, h in enumerate(headers):
    cell = tbl.cell(0, c)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True; r.font.size = Pt(9)
for ri, row in enumerate(data):
    for ci, val in enumerate(row):
        cell = tbl.cell(ri + 1, ci)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(val).font.size = Pt(9)
body('As shown in Table 2.1, the proposed AgriFlow AI platform is the only system that combines '
     'all six capabilities in a single integrated platform, addressing the gaps identified in the '
     'literature.')
body('Figure 2.1 provides a visual summary of how each platform maps to the six key capabilities.')
caption('Figure 2.1: Comparison of Existing Systems by Capability')

h2('2.8 Research Gap')
body('Based on the reviewed literature and existing systems, the following research gaps were '
     'identified:')
numbered('Limited Integration of Marketplace and Logistics: Many systems address either trading '
         'or logistics, but few unify a marketplace with perishable-first scheduling, inventory, '
         'and forecasting in one platform.')
numbered('Weak Perishability Awareness: Delivery scheduling in existing systems rarely '
         'prioritises orders by remaining shelf life, leading to avoidable spoilage.')
numbered('Insufficient Shelf-Life Tracking: Few systems continuously monitor the freshness of '
         'stored produce and flag items that are expiring or spoiled.')
numbered('Inaccessible Analytics: Operational data and forecasts are often presented in forms '
         'that non-specialist users find difficult to interpret.')
numbered('Absence of Grounded AI Assistance: Existing systems rarely provide an '
         'artificial-intelligence assistant that explains live operational data reliably without '
         'fabricating information.')

h2('2.9 Summary of Literature Review')
body('This chapter reviewed literature on agricultural supply chain management, electronic '
     'marketplaces, optimisation and forecasting, artificial intelligence in agriculture, web '
     'development technologies, and existing systems. The review showed that digital platforms '
     'improve coordination and market access, that deterministic optimisation and forecasting '
     'techniques offer reliable and explainable decisions, and that artificial intelligence can '
     'make complex data accessible when properly grounded. Nevertheless, existing systems seldom '
     'integrate these capabilities into a single perishability-aware platform. These gaps justify '
     'the development of the proposed AgriFlow AI platform.')

# ==================== CHAPTER THREE ====================
chapter(['Chapter Three', 'Methodology'])

h2('3.1 Introduction')
body('This chapter presents the methodology adopted for the design and implementation of the '
     'AgriFlow AI platform. It describes the design approach, the design considerations, the '
     'system architecture and its components, the database design, the deterministic decision '
     'algorithms, and the operational flow of the system. The methodology provides a structured '
     'framework for developing a reliable, secure, and intelligent supply chain platform.')

h2('3.2 Project Design and Approach')
body('The development of the AgriFlow AI platform followed the Software Development Life Cycle '
     '(SDLC) using the Agile development approach. The Agile methodology was adopted because it '
     'supports iterative development, continuous refinement, and regular testing. Features were '
     'developed incrementally and validated through repeated testing, including an end-to-end '
     'test that exercises all five user roles across the complete supply chain.')
body('The system was designed as a web application accessible through modern browsers and '
     'installable as a progressive web application. It defines five user roles: the Farmer, who '
     'manages crops and harvests and confirms orders; the Buyer, who browses the marketplace and '
     'places orders; the Transporter, who carries out deliveries; the Warehouse Manager, who '
     'manages inventory and stock movements; and the Administrator, who oversees users, runs the '
     'scheduling engine, and views platform-wide reports.')
body('The application adopts a modern serverless web architecture based on the Next.js framework. '
     'Rather than a separate application server, it uses server components and server actions '
     'that run on the server and communicate directly with a cloud-hosted database. This '
     'promotes maintainability, security, and scalability.')

h2('3.3 Design Consideration')
body('Several factors were considered during the design of the proposed system.')
h3('Security')
body('The system uses secure authentication and authorisation. Access to data is enforced at the '
     'database level using row-level security policies, so that each user can access only the '
     'data permitted for their role. Server actions additionally re-check roles before performing '
     'sensitive operations.')
h3('Reliability')
body('All business decisions, including scheduling, forecasting, and spoilage detection, are '
     'made by deterministic and testable algorithms. The artificial-intelligence assistant is '
     'restricted to explanation and is grounded in live data, ensuring that the system’s '
     'decisions are reliable and reproducible.')
h3('Usability')
body('The interface was designed to be clean, responsive, and intuitive for users with varied '
     'levels of technical expertise. A dark mode and mobile-first responsive design improve '
     'accessibility across devices.')
h3('Scalability')
body('The serverless architecture and cloud-hosted database allow the system to accommodate '
     'growing numbers of users and records without significant redesign.')
h3('Accessibility')
body('The platform is installable as a progressive web application with offline support, '
     'enabling access on a range of devices and network conditions.')

h2('3.4 Requirements Specification')
body('The system requirements were categorised into functional and non-functional requirements. '
     'Functional requirements define the core operations the platform must perform, while '
     'non-functional requirements specify the quality attributes of the system.')
h3('3.4.1 Functional Requirements')
body('Table 3.3 lists the functional requirements of the AgriFlow AI platform, specifying the '
     'operations the system must support for each functional area.')
caption('Table 3.3: Functional Requirements Specification')
fr = doc.add_table(rows=11, cols=3)
fr.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr.style = 'Table Grid'
fr_headers = ['ID', 'Requirement', 'Module']
fr_data = [
    ['FR-01', 'Users shall register and log in with role-based access', 'Auth'],
    ['FR-02', 'Farmers shall create, update, and view crops', 'Farm'],
    ['FR-03', 'Farmers shall record harvests with pricing and grade', 'Farm'],
    ['FR-04', 'Buyers shall browse available produce in a marketplace', 'Marketplace'],
    ['FR-05', 'Buyers shall place orders with delivery details', 'Marketplace'],
    ['FR-06', 'The system shall schedule deliveries by perishability', 'Scheduling'],
    ['FR-07', 'The system shall forecast demand from order history', 'Forecasting'],
    ['FR-08', 'Warehouse managers shall track inventory and shelf life', 'Warehouse'],
    ['FR-09', 'Transporters shall advance delivery status', 'Transport'],
    ['FR-10', 'The system shall provide an AI assistant for data explanation', 'AI'],
]
for c, h in enumerate(fr_headers):
    cell = fr.cell(0, c)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True; r.font.size = Pt(9)
for ri, row in enumerate(fr_data):
    for ci, val in enumerate(row):
        cell = fr.cell(ri + 1, ci)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(val).font.size = Pt(9)

h3('3.4.2 Non-Functional Requirements')
body('Table 3.4 lists the non-functional requirements, specifying the quality attributes that the '
     'system must satisfy.')
caption('Table 3.4: Non-Functional Requirements Specification')
nfr = doc.add_table(rows=7, cols=3)
nfr.alignment = WD_ALIGN_PARAGRAPH.CENTER
nfr.style = 'Table Grid'
nfr_headers = ['ID', 'Requirement', 'Category']
nfr_data = [
    ['NFR-01', 'Authentication and data access enforced at database level', 'Security'],
    ['NFR-02', 'Business decisions made by deterministic, testable algorithms', 'Reliability'],
    ['NFR-03', 'Interface responsive from 320px for mobile access', 'Usability'],
    ['NFR-04', 'Serverless architecture supporting horizontal scaling', 'Scalability'],
    ['NFR-05', 'Installable as a progressive web application', 'Accessibility'],
    ['NFR-06', 'AI assistant degrades gracefully when API key is absent', 'Robustness'],
]
for c, h in enumerate(nfr_headers):
    cell = nfr.cell(0, c)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True; r.font.size = Pt(9)
for ri, row in enumerate(nfr_data):
    for ci, val in enumerate(row):
        cell = nfr.cell(ri + 1, ci)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(val).font.size = Pt(9)

h2('3.5 System Architecture')
body('The AgriFlow AI platform adopts a modern web architecture consisting of a presentation '
     'layer, a server logic layer, a data layer, and external services.')
body('The Presentation Layer is delivered to the browser and built with the Next.js framework, '
     'the React library, and Tailwind CSS. It renders the landing page, authentication pages, '
     'and the role-specific dashboards through which users interact with the system.')
body('The Server Logic Layer consists of Next.js server components and server actions that run '
     'on the server. This layer processes user requests, enforces role-based authorisation, '
     'applies the deterministic scheduling, forecasting, and spoilage algorithms, triggers '
     'notifications, and coordinates communication with the database and external services.')
body('The Data Layer is a Supabase-hosted PostgreSQL database that stores all persistent data. '
     'Row-level security policies enforce access control directly at the database, and a signup '
     'trigger automatically creates a user profile on registration.')
body('External Services include Google Gemini, which powers the explanatory '
     'artificial-intelligence assistant, and OpenStreetMap, whose map tiles are displayed '
     'through the Leaflet library for delivery routes. The interaction among these layers '
     'ensures secure, efficient, and intelligent operation.')
caption('Figure 3.1: System Architecture of the AgriFlow AI Platform')

h2('3.6 Components of the System')
body('The system is organised into the following major functional components.')
h3('Authentication and Authorisation Module')
body('This module manages registration, login, logout, and password reset for the five user '
     'roles, and enforces role-based access to the appropriate dashboards and operations.')
h3('Marketplace and Order Module')
body('This module allows farmers to list harvested produce and buyers to browse and order it. '
     'Ordering adjusts available stock, and orders progress through defined states from pending '
     'to delivered.')
h3('Farm and Harvest Management Module')
body('This module enables farmers to manage crops and record harvests, including quantity, '
     'quality grade, price, and shelf life.')
h3('Warehouse and Inventory Module')
body('This module manages warehouse inventory, records stock movements, tracks the freshness of '
     'stored items using shelf-life information, and flags items that are expiring or spoiled.')
h3('Delivery and Scheduling Module')
body('This module contains the deterministic scheduling engine, which prioritises perishable '
     'orders, assigns the nearest warehouse as the pickup point and the least-busy transporter, '
     'and manages the delivery workflow through to completion.')
h3('Forecasting and Reporting Module')
body('This module computes demand forecasts from historical orders and presents analytics and '
     'reports to support planning and administration.')
h3('Notification Module')
body('This module generates rule-based alerts for events such as order placement, confirmation, '
     'scheduling, delivery updates, and spoilage, keeping users informed in real time.')
h3('AI Assistant Module')
body('This module provides a natural-language assistant, powered by Google Gemini and grounded '
     'in the user’s live data, that explains forecasts, flags risks, and answers questions, '
     'degrading gracefully when the service is unavailable.')

h2('3.7 Development Tools and Technologies')
body('The tools and technologies used in the development of the system are summarised in Table '
     '3.1.')
caption('Table 3.1: Software Development Tools and Technologies')
bullet('Next.js 14 — React-based web framework providing server components and server actions.')
bullet('TypeScript — Typed programming language used for reliable application code.')
bullet('Tailwind CSS — Utility-first styling framework, with Framer Motion for animation.')
bullet('Supabase (PostgreSQL, Auth, RLS) — Cloud database, authentication, and access control.')
bullet('Google Gemini — Large language model powering the explanatory AI assistant.')
bullet('Leaflet with OpenStreetMap — Interactive maps for delivery routes.')
bullet('Recharts — Charting library for analytics and reports.')
bullet('Playwright — End-to-end testing framework.')
bullet('Git and Vercel — Version control and hosting/deployment.')

h2('3.8 Database Design')
body('The database was designed as a relational schema of related tables, each protected by '
     'row-level security. The principal tables and their roles are summarised in Table 3.2.')
caption('Table 3.2: Summary of Database Tables')
bullet('profiles — user accounts with full name, role, phone, and location.')
bullet('crops — crops recorded by farmers, with planting and expected harvest dates and status.')
bullet('harvests — harvested produce with quantity, quality grade, price, and shelf-life days.')
bullet('warehouses — storage locations with manager, coordinates, and capacity.')
bullet('inventory_items — stored produce with entry date, shelf life, and status.')
bullet('stock_movements — records of stock entering or leaving a warehouse.')
bullet('orders — buyer orders linked to farmer and harvest, with quantity, price, and status.')
bullet('deliveries — deliveries linked to orders and transporters, with pickup, drop-off, '
       'scheduled date, distance, and status.')
bullet('notifications — per-user alerts with title, message, type, and read status.')
body('The relationships among these tables link profiles to crops, harvests, orders, and '
     'deliveries, and link warehouses to inventory items and stock movements. These '
     'relationships preserve data integrity and support efficient retrieval, as illustrated in '
     'the entity relationship diagram.')
caption('Figure 3.2: Entity Relationship Diagram of the Database')
caption('Figure 3.3: Use-Case Diagram of the System')

h2('3.9 Model and Algorithm Description')
body('The core operational decisions of the platform are made by deterministic algorithms, '
     'described below.')
h3('3.9.1 Delivery Scheduling Algorithm')
body('The scheduling engine first sorts confirmed orders by the remaining shelf life of their '
     'produce, computed from the harvest date and the shelf-life days, so that the most '
     'perishable orders are scheduled first; ties are broken by order age, with older orders '
     'given priority. Deliveries are then allocated to days, with a maximum of five deliveries '
     'per day and overflow rolling to subsequent days. For each order, the nearest warehouse to '
     'the delivery location is selected as the pickup point using the great-circle distance, and '
     'the least-busy transporter is assigned. The result is a set of scheduled deliveries with '
     'pickup and drop-off points, dates, and distances.')
body('Figure 3.4 presents the scheduling flowchart, illustrating the complete flow from order '
     'placement to delivery completion, including the parallel processes of demand forecasting '
     'and shelf-life monitoring.')
figure(FLOWCHART, 'Figure 3.4: Delivery Scheduling Flowchart of the AgriFlow AI Platform')
h3('3.8.2 Demand Forecasting Algorithm')
body('Demand forecasting builds a weekly demand series for each product from historical orders '
     'and applies two classical techniques: the simple moving average, which averages the most '
     'recent periods, and the weighted moving average, which gives greater weight to more recent '
     'periods. A trend indicator expresses the percentage change of the forecast relative to the '
     'most recent observed period, helping users anticipate rising or falling demand.')
h3('3.8.3 Spoilage and Shelf-Life Tracking')
body('The spoilage component computes, for each stored item, the number of days remaining before '
     'expiry from its entry date and shelf-life days. Items are classified as fresh, expiring '
     'when close to their limit, or spoiled once the limit is passed, enabling warehouse '
     'managers to act before produce is lost.')
h3('3.8.4 Distance Computation')
body('Distances between locations, used for nearest-warehouse selection and route display, are '
     'computed with the haversine formula, which returns the great-circle distance between two '
     'points from their latitude and longitude.')

h2('3.10 Operational Flow of the System')
body('The system begins when a user registers or logs in and is directed to the dashboard for '
     'their role. A farmer records crops and harvests, and harvested produce is listed on the '
     'marketplace. A buyer browses the marketplace and places an order, which reduces the '
     'available stock, and the farmer confirms the order. The administrator runs the scheduling '
     'engine, which prioritises perishable orders and assigns pickups and transporters. The '
     'transporter then advances the delivery through the stages of picked up, in transit, and '
     'delivered, while the buyer and farmer are notified of progress.')
body('Throughout the process, warehouse managers receive and dispatch stock and monitor '
     'shelf-life alerts, forecasts are updated from order history, and the artificial-'
     'intelligence assistant is available to explain the data. The process continues until users '
     'complete their operations and log out, which terminates their sessions and preserves '
     'security.')

# ==================== CHAPTER FOUR ====================
chapter(['Chapter Four', 'System Implementation, Results and Discussion'])

h2('4.1 Introduction')
body('This chapter presents the implementation of the AgriFlow AI platform, the development '
     'environment, the system interfaces, the testing procedures, and a discussion of the '
     'results in relation to the objectives of the study and the gaps identified in the '
     'literature.')

h2('4.2 System Implementation')
body('The system was implemented following the architecture and design described in Chapter '
     'Three. The application was built with the Next.js 14 framework using TypeScript, with the '
     'user interface styled using Tailwind CSS and animated with Framer Motion. Server '
     'components and server actions implemented the business logic and communicated directly '
     'with the Supabase PostgreSQL database.')
body('Authentication, authorisation, and data access were enforced through Supabase '
     'authentication and row-level security policies, complemented by role checks in server '
     'actions. The deterministic engines for scheduling, forecasting, and spoilage were '
     'implemented as self-contained modules to ensure they were reliable and testable. The '
     'Google Gemini service was integrated to provide the explanatory assistant, grounded in the '
     'user’s live data, with a graceful fallback when the service key is not configured. The '
     'application was configured as an installable progressive web application and deployed to '
     'the Vercel hosting platform.')

h2('4.3 System Interfaces')
body('Table 4.1 describes the key user interface components for each major screen of the '
     'AgriFlow AI platform.')
caption('Table 4.1: Interface Features by Screen')
ift = doc.add_table(rows=10, cols=4)
ift.alignment = WD_ALIGN_PARAGRAPH.CENTER
ift.style = 'Table Grid'
ift_headers = ['Screen', 'Components', 'Role', 'Key Features']
ift_data = [
    ['Landing', 'Hero, features, how-it-works, FAQ, footer', 'All', 'Fullscreen carousel, parallax cards, animated counters, dark mode'],
    ['Signup/Login', 'Role selector, email/password, Google OAuth', 'All', 'Password strength meter, generator, show/hide, remember me'],
    ['Dashboard', 'Sidebar, navbar, stats cards, welcome splash', 'All', 'Role-based navigation, unread notifications, mobile drawer'],
    ['Crops', 'Crop table, add/edit form, status selector', 'Farmer', 'CRUD with status lifecycle, planting/harvest calendar'],
    ['Harvests', 'Harvest table, record form', 'Farmer', 'Quantity, grade, price, shelf life listing to marketplace'],
    ['Marketplace', 'Produce cards, search, order form', 'Buyer', 'Search filter, city select for routing, live stock'],
    ['Inventory', 'Items table, receive/dispatch stock forms', 'Warehouse', 'Shelf-life freshness indicator, movement log, search'],
    ['Deliveries', 'Delivery list, advance-status form, Leaflet map', 'Transport', 'Pickup → transit → delivered workflow, route map'],
    ['Scheduling', 'Run scheduler control, results, stats', 'Admin', 'Perishable-first sort, day allocation, warehouse/transporter assignment'],
]
for c, h in enumerate(ift_headers):
    cell = ift.cell(0, c)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True; r.font.size = Pt(9)
for ri, row in enumerate(ift_data):
    for ci, val in enumerate(row):
        cell = ift.cell(ri + 1, ci)
        cell.text = ''
        p = cell.paragraphs[0]
        p.add_run(val).font.size = Pt(9)
body('This section presents the major interfaces of the implemented system.')
h3('4.3.1 Landing Page')
body('The landing page, shown in Figure 4.1, presents the platform with a hero section, feature '
     'highlights, and informational sections that introduce new users to the system.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\agriflow\01-landing.png', 'Figure 4.1: Landing Page of the AgriFlow AI Platform')
h3('4.3.2 Authentication Interfaces')
body('The signup and login interfaces, shown in Figure 4.2, allow users to register under one of '
     'the five roles and to authenticate, with features such as a password strength meter and '
     'secure password generator.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\agriflow\03-signup.png', 'Figure 4.2: Signup Interface (left) and Login Interface (right)')
h3('4.3.3 Farmer Interfaces')
body('The farmer dashboard and crop management interfaces, shown in Figure 4.3, allow farmers to '
     'manage crops, record harvests with pricing and grading, view a harvest calendar, and '
     'confirm incoming orders.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\agriflow\dash-farmer.png', 'Figure 4.3: Farmer Dashboard of the AgriFlow AI Platform')
h3('4.3.4 Buyer Interfaces')
body('The buyer marketplace and ordering interfaces, shown in Figure 4.4, allow buyers to search '
     'available produce, place orders with a delivery location, and track their orders.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\agriflow\dash-buyer.png', 'Figure 4.4: Buyer Marketplace of the AgriFlow AI Platform')
h3('4.3.5 Warehouse Interfaces')
body('The warehouse inventory interface, shown in Figure 4.5, displays stored items with live '
     'shelf-life status, supports stock-in and stock-out operations, and records stock '
     'movements.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\agriflow\warehouse-inventory.png', 'Figure 4.5: Warehouse Inventory Interface of the AgriFlow AI Platform')
h3('4.3.6 Transporter Interfaces')
body('The transporter interfaces, shown in Figure 4.6, present assigned deliveries and allow the '
     'transporter to advance each delivery through its workflow, with routes displayed on a '
     'Leaflet map.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\agriflow\dash-transporter.png', 'Figure 4.6: Transporter Delivery Workflow of the AgriFlow AI Platform')
h3('4.3.7 Administrator Interfaces')
body('The administrator interfaces, shown in Figure 4.7, provide user management with role '
     'changes, a full order log, platform statistics, analytics reports, and a control to run '
     'the scheduling engine.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\agriflow\dash-admin.png', 'Figure 4.7: Administrator Scheduling Interface of the AgriFlow AI Platform')
h3('4.3.8 AI Assistant Interface')
body('The assistant interface, shown in Figure 4.8, allows users to ask questions about their '
     'data and receive plain-language explanations grounded in their live profile, orders, '
     'inventory, and forecasts.')
figure(r'c:\Users\ALEXIS\Desktop\SENPAI\shots\agriflow\assistant.png', 'Figure 4.8: AI Assistant Interface of the AgriFlow AI Platform')

h2('4.5 System Testing')
body('The system was tested at three levels: unit testing of the deterministic algorithms, '
     'functional testing of complete features, and end-to-end testing across all five roles.')
h3('4.5.1 Unit Testing')
body('The scheduling, forecasting, spoilage, and distance functions were tested with '
     'representative inputs to confirm that they produced the expected outputs, including correct '
     'perishable-first ordering, accurate moving averages, correct shelf-life classification, and '
     'correct distance computation.')
h3('4.5.2 Functional Testing')
body('Complete features were tested end to end, including registration and role-based access, '
     'crop and harvest management, marketplace ordering and stock adjustment, order confirmation, '
     'delivery scheduling, the delivery workflow, inventory and spoilage tracking, notifications, '
     'and the assistant. The results are summarised in Table 4.2.')
caption('Table 4.2: Summary of Test Cases and Results')
tt = doc.add_table(rows=22, cols=5)
tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
tt.style = 'Table Grid'
tt_headers = ['ID', 'Test Case', 'Input', 'Expected Result', 'Status']
tt_data = [
    ['TC-01', 'User Registration', 'Valid details for each role', 'Account created; profile with role stored', 'Passed'],
    ['TC-02', 'Role-Based Login', 'Email + password per role', 'User routed to correct role dashboard', 'Passed'],
    ['TC-03', 'Wrong Password', 'Valid email + wrong password', 'Error message: "Incorrect email or password"', 'Passed'],
    ['TC-04', 'Add Crop', 'Name, category, planting date, yield', 'Crop saved and displayed in crops table', 'Passed'],
    ['TC-05', 'Record Harvest', 'Crop, date, qty, grade, price, shelf life', 'Harvest listed to marketplace as available', 'Passed'],
    ['TC-06', 'Browse Marketplace', 'Search term', 'Filtered produce cards returned', 'Passed'],
    ['TC-07', 'Place Order', 'Harvest ID, qty, delivery city', 'Order created; stock decremented', 'Passed'],
    ['TC-08', 'Wrong Password on Login', 'Invalid credentials', 'Error message displayed; no redirect', 'Passed'],
    ['TC-09', 'Confirm Order', 'Order ID', 'Status changed to confirmed; notification sent', 'Passed'],
    ['TC-10', 'Run Scheduling', 'Confirmed orders', 'Perishable-first sort; nearest warehouse; transporter assigned', 'Passed'],
    ['TC-11', 'Scheduling — Max 5/day', '>5 confirmed orders', 'Orders allocated across days; overflow rolls', 'Passed'],
    ['TC-12', 'Advance Delivery', 'Delivery ID; next status', 'Status progressed (assigned → picked_up → in_transit → delivered)', 'Passed'],
    ['TC-13', 'Receive Stock', 'Warehouse, product, qty, entry date, shelf life', 'Item added to inventory as in_storage', 'Passed'],
    ['TC-14', 'Shelf-Life Classification', 'Item with entry_date and shelf_life_days', 'Classified as fresh/expiring/spoiled correctly', 'Passed'],
    ['TC-15', 'Dispatch Stock', 'Item ID, qty', 'Stock removed; movement logged', 'Passed'],
    ['TC-16', 'Mark Spoiled', 'Item ID', 'Item status changed to spoiled', 'Passed'],
    ['TC-17', 'Demand Forecast', '8 weeks of order history', 'Moving average and weighted MA computed; trend indicator shown', 'Passed'],
    ['TC-18', 'AI Assistant Query', 'Question about user data', 'Plain-language explanation grounded in live data', 'Passed'],
    ['TC-19', 'AI Assistant — No API Key', 'Question without GEMINI_API_KEY', 'Fallback message displayed gracefully', 'Passed'],
    ['TC-20', 'Password Generator', 'Click generate', 'Secure password created and populated in both fields', 'Passed'],
    ['TC-21', 'Notification Delivery', 'Order/delivery event', 'Notification created and displayed per user', 'Passed'],
]
for c, h in enumerate(tt_headers):
    cell = tt.cell(0, c)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True; r.font.size = Pt(8)
for ri, row in enumerate(tt_data):
    for ci, val in enumerate(row):
        cell = tt.cell(ri + 1, ci)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci in (0, 4) else WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(val).font.size = Pt(8)
h3('4.5.3 End-to-End Testing')
body('An automated end-to-end test using the Playwright framework created one account for each '
     'of the five roles and drove them through the complete supply chain, from listing produce '
     'to delivery, in a real browser. The test verified the full workflow across many steps and '
     'captured a screenshot on any failure, confirming that the integrated system behaved '
     'correctly.')

h2('4.6 Implementation Challenges')
body('Several challenges were encountered during the implementation of the AgriFlow AI platform.')
h3('4.6.1 Supabase Row-Level Security Configuration')
body('Configuring row-level security policies for the five user roles required careful planning '
     'to ensure that each role could access only the data permitted to it. The policies had to '
     'be written for every table, and testing each combination of role and operation was '
     'time-consuming. This challenge was addressed by writing the policies iteratively and '
     'testing each one with test user accounts before proceeding to the next.')
h3('4.6.2 Deterministic Scheduling Algorithm Design')
body('Designing the scheduling algorithm to handle the perishable-first rule while also '
     'respecting the maximum of five deliveries per day, nearest-warehouse assignment, and '
     'least-busy transporter selection required careful consideration of edge cases, such as '
     'when all transporters were already assigned or when no warehouse was near the delivery '
     'location. The algorithm was developed with comprehensive unit tests to ensure that all '
     'edge cases were handled correctly.')
h3('4.6.3 AI Assistant Grounding')
body('Ensuring that the Google Gemini assistant did not fabricate information required a '
     'careful prompt design that restricted the assistant to explaining only the data provided '
     'in the prompt context. The assistant was also designed to degrade gracefully when the '
     'Gemini API key was not configured, showing a helpful fallback message instead of an '
     'error. This approach ensures reliability while still providing value when the AI service '
     'is available.')
h3('4.6.4 Responsive Design for Multiple Roles')
body('Each of the five role dashboards had different interface requirements, from the farmer\'s '
     'crop calendar to the transporter\'s route map. Ensuring that all interfaces were responsive '
     'from 320px mobile screens required a mobile-first approach with Tailwind CSS breakpoints. '
     'The sidebar navigation was adapted to a drawer on small screens, and data tables were '
     'made horizontally scrollable to preserve readability.')

h2('4.7 Discussion of Results')
body('The results confirmed that the platform met the objectives established in Chapter One. A '
     'unified, role-based platform connecting the five participants was implemented; a produce '
     'marketplace with crop, harvest, inventory, and order management was provided; a '
     'deterministic perishable-first scheduling engine was developed; demand forecasting and '
     'shelf-life spoilage tracking were implemented; and an explanatory artificial-intelligence '
     'assistant grounded in live data was integrated. The functionality and reliability of the '
     'platform were evaluated through unit, functional, and end-to-end testing.')
body('In relation to the existing literature, the platform addresses the gaps identified in '
     'Chapter Two. Unlike systems that focus solely on marketplace transactions or on '
     'company-operated logistics, AgriFlow AI integrates trading with perishability-aware '
     'scheduling, inventory, forecasting, and analytics in a single configurable platform. Its '
     'shelf-life tracking directly targets post-harvest loss, its accessible forecasts and '
     'grounded assistant make operational data understandable, and its deterministic decision '
     'logic ensures reliable and explainable outcomes. These findings demonstrate the value of '
     'combining a unified digital platform with deterministic optimisation and explanatory '
     'artificial intelligence for agricultural supply chain management.')

# ==================== CHAPTER FIVE ====================
chapter(['Chapter Five', 'Summary, Conclusion and Recommendations'])

h2('5.1 Summary of Findings')
body('This project set out to design and implement an intelligent web-based agricultural supply '
     'chain management platform. The problem addressed was the inefficiency of perishable-produce '
     'supply chains, arising from the disconnection of participants, delivery scheduling that '
     'ignores perishability, unanticipated demand, and operational data that is difficult to '
     'interpret.')
body('A review of relevant literature and existing systems established the value of digital '
     'platforms, deterministic optimisation and forecasting, and grounded artificial '
     'intelligence, while revealing that existing systems seldom integrate these capabilities in '
     'a perishability-aware platform. Guided by these findings, the system was developed using '
     'the Agile approach and a modern serverless web architecture, implemented with Next.js, '
     'TypeScript, Tailwind CSS, and a Supabase PostgreSQL database, with Google Gemini providing '
     'explanation only.')
body('The developed platform unified five roles; provided a marketplace and management of crops, '
     'harvests, inventory, orders, and deliveries; scheduled deliveries with a perishable-first '
     'rule; forecast demand; tracked shelf life; issued rule-based notifications; and offered an '
     'explanatory assistant. Testing at the unit, functional, and end-to-end levels confirmed '
     'that the system prioritised perishable orders, assigned deliveries efficiently, tracked '
     'shelf life accurately, produced reliable forecasts, and explained data safely.')

h2('5.2 Contributions to Knowledge')
body('This study makes the following contributions to knowledge in the field of agricultural '
     'supply chain management and intelligent web-based systems:')
numbered('A unified, role-based platform for agricultural supply chain management that integrates '
         'five distinct participant types within a single system, providing a practical framework '
         'for multi-role agricultural digital platforms.')
numbered('A deterministic, perishable-first delivery scheduling algorithm that prioritises '
         'deliveries by remaining shelf life, with an implementation that demonstrates how simple '
         'rule-based logic can effectively reduce post-harvest losses without the complexity of '
         'optimisation solvers.')
numbered('A practical approach to grounding an artificial-intelligence assistant in live '
         'operational data for explanation purposes, where all critical decisions remain '
         'deterministic and auditable, demonstrating a safe and reliable integration of '
         'generative AI in operational systems.')
numbered('An integrated combination of demand forecasting, shelf-life spoilage tracking, and '
         'rule-based notifications within a single agricultural platform, providing a reference '
         'architecture that addresses the limitations of marketplace-only solutions identified in '
         'the literature.')
numbered('A comprehensive test suite and end-to-end testing methodology that validates the '
         'complete supply chain workflow across all roles, providing a reproducible approach '
         'to verifying multi-actor agricultural platforms.')

h2('5.3 Conclusion')
body('The study successfully achieved its aim of designing and implementing an intelligent '
     'web-based agricultural supply chain management platform. By combining a unified, role-based '
     'marketplace with deterministic perishable-first scheduling, demand forecasting, shelf-life '
     'spoilage tracking, and an explanatory artificial-intelligence assistant, AgriFlow AI '
     'improves the coordination and efficiency of the produce supply chain and directly targets '
     'the reduction of post-harvest losses. The platform therefore provides a practical and '
     'reliable solution to the challenges identified in the study.')

h2('5.4 Recommendations')
body('Based on the outcomes of this study, the following recommendations are made:')
numbered('Agricultural cooperatives, aggregators, and agribusinesses should adopt integrated, '
         'perishability-aware platforms of this kind to reduce post-harvest losses.')
numbered('Delivery scheduling in agricultural logistics should prioritise produce by remaining '
         'shelf life to minimise spoilage.')
numbered('Operational systems should present analytics and forecasts in accessible forms, '
         'supported by explanatory tools, so that non-specialist users can act on them.')
numbered('Where artificial intelligence is used in operational systems, it should be grounded in '
         'verified data and restricted to explanation, with critical decisions made by '
         'deterministic and auditable logic.')

h2('5.5 Suggestions for Future Research')
body('The following suggestions are offered for future work:')
numbered('Developing native mobile applications to complement the progressive web application for '
         'field use.')
numbered('Integrating live logistics tracking, such as global positioning of vehicles, for '
         'real-time delivery monitoring.')
numbered('Incorporating predictive machine-learning models for more accurate demand forecasting '
         'and spoilage prediction.')
numbered('Adding integrated payment processing to support end-to-end transactions within the '
         'platform.')
numbered('Integrating internet-of-things sensors to monitor storage conditions such as '
         'temperature and humidity in real time.')
numbered('Conducting a large-scale field evaluation to measure the platform’s impact on '
         'post-harvest losses quantitatively.')

# ==================== REFERENCES ====================
chapter(['References'])
for r in [
    'Anderson, T. (2008). The theory and practice of online learning (2nd ed.). Athabasca '
    'University Press.',
    'Christopher, M. (2016). Logistics and supply chain management (5th ed.). Pearson Education.',
    'Elmasri, R., & Navathe, S. B. (2017). Fundamentals of database systems (7th ed.). Pearson.',
    'Google. (2024). Gemini API documentation. https://ai.google.dev/',
    'Pressman, R. S., & Maxim, B. R. (2020). Software engineering: A practitioner’s approach '
    '(9th ed.). McGraw-Hill Education.',
    'Russell, S. J., & Norvig, P. (2021). Artificial intelligence: A modern approach (4th ed.). '
    'Pearson.',
    'Sommerville, I. (2016). Software engineering (10th ed.). Pearson Education.',
    'Supabase. (2024). Supabase documentation. https://supabase.com/docs',
    'Vercel. (2024). Next.js documentation. https://nextjs.org/docs',
    'Chen, L., & Zhang, Y. (2023). Demand forecasting in agricultural supply chains: A review '
    'of methods and applications. Computers and Electronics in Agriculture, 204, 107521.',
    'Garnett, T., & Godfray, C. (2012). Sustainable intensification in agriculture: Navigating '
    'a course through competing food system priorities. Food Climate Research Network.',
    'Gustavsson, J., Cederberg, C., Sonesson, U., van Otterdijk, R., & Meybeck, A. (2011). '
    'Global food losses and food waste. Food and Agriculture Organization of the United Nations.',
    'Hurni, H. (2020). Sustainable management of agricultural resources. Mountain Research and '
    'Development, 40(1), 1-10.',
    'Lam, H. Y., & Ip, W. H. (2021). An integrated decision support system for perishable '
    'agricultural supply chains. Expert Systems with Applications, 168, 114325.',
    'Shukla, M., & Jharkharia, S. (2013). Agri-fresh produce supply chain management: A '
    'state-of-the-art literature review. International Journal of Operations & Production '
    'Management, 33(2), 114-158.',
    'Tsang, Y. P., Wu, C. H., & Lam, H. Y. (2022). A blockchain-based traceability system for '
    'perishable agricultural products. Journal of Cleaner Production, 345, 131139.',
    'World Bank. (2019). Enabling the business of agriculture 2019. World Bank Group.',
]:
    reference(r)

# CONTENT_MARKER
save(r'c:\Users\ALEXIS\Desktop\SENPAI\AgriFlow_AI_Project_Report.docx')
